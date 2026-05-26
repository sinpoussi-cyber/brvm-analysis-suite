# ==============================================================================
# MODULE: REPORT GENERATOR V30.0 FINAL - HISTORISATION + RISQUE CHIFFRÉ
# Toutes fonctionnalités V29 + Historisation DB + Calcul risque + Décisions
# ==============================================================================

import os
import logging
import psycopg2
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
import time
import json
from collections import defaultdict, Counter
import io
import base64
try:
    import matplotlib
    matplotlib.use('Agg')          # backend non-interactif, safe en CI/GitHub Actions
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    from matplotlib.gridspec import GridSpec
    MATPLOTLIB_OK = True
except ImportError:
    MATPLOTLIB_OK = False
    logging.warning("⚠️  matplotlib non disponible — graphiques désactivés")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s: %(message)s')

# --- Configuration & Secrets ---
DB_NAME = os.environ.get('DB_NAME')
DB_USER = os.environ.get('DB_USER')
DB_PASSWORD = os.environ.get('DB_PASSWORD')
DB_HOST = os.environ.get('DB_HOST')
DB_PORT = os.environ.get('DB_PORT')

# ✅ CONFIGURATION MULTI-AI (Rotation: DeepSeek → Claude → Gemini → Mistral)
DEEPSEEK_API_KEY = os.environ.get('DEEPSEEK_API_KEY')
DEEPSEEK_MODEL = "deepseek-chat"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
GEMINI_MODEL = "gemini-1.5-flash"
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"

MISTRAL_API_KEY = os.environ.get('MISTRAL_API_KEY')
MISTRAL_MODEL = "mistral-large-latest"
MISTRAL_API_URL = "https://api.mistral.ai/v1/chat/completions"

# Claude (Anthropic) — fallback final si toutes les autres IA échouent
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')
ANTHROPIC_MODEL = "claude-haiku-4-5-20251001"   # modèle rapide et économique
ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"
ANTHROPIC_API_VERSION = "2023-06-01"


class BRVMReportGenerator:
    def __init__(self):
        self.db_conn = None
        self.request_count = {'deepseek': 0, 'gemini': 0, 'mistral': 0, 'claude': 0, 'total': 0}
        self.all_recommendations = {}
        
        try:
            self.db_conn = psycopg2.connect(
                dbname=DB_NAME, user=DB_USER, password=DB_PASSWORD,
                host=DB_HOST, port=DB_PORT,
                connect_timeout=10,
                options='-c statement_timeout=60000'
            )
            logging.info("✅ Connexion DB établie")
        except Exception as e:
            logging.error(f"❌ Erreur connexion DB: {e}")
            raise

    def _get_market_events(self):
        """Récupère les événements récents du marché"""
        logging.info("📰 Récupération des événements marquants...")
        
        query = """
        SELECT event_date, event_summary
        FROM new_market_events
        ORDER BY event_date DESC
        LIMIT 5;
        """
        
        try:
            df = pd.read_sql(query, self.db_conn)
            if not df.empty:
                events = []
                for _, row in df.iterrows():
                    events.append(f"• {row['event_date'].strftime('%d/%m/%Y')}: {row['event_summary']}")
                return "\n".join(events)
            return "Aucun événement récent enregistré."
        except Exception as e:
            logging.error(f"❌ Erreur récupération événements: {e}")
            return "Données indisponibles."

    def _get_market_indicators(self):
        """Récupère les derniers indicateurs du marché + historique 100j pour commentaire"""
        # Dernière ligne valide (brvm_composite non null), ordonnée par id décroissant
        query = """
        SELECT 
            id,
            brvm_composite, 
            brvm_30, 
            brvm_prestige, 
            capitalisation_globale,
            extraction_date
        FROM new_market_indicators
        WHERE brvm_composite IS NOT NULL
          AND brvm_composite > 0
        ORDER BY id DESC
        LIMIT 1;
        """
        
        try:
            df = pd.read_sql(query, self.db_conn)
            if not df.empty:
                row = df.iloc[0]
                composite = row.get('brvm_composite')
                if pd.notna(composite):
                    indicators = {
                        'composite': float(composite),
                        'capitalisation': float(row.get('capitalisation_globale')) if pd.notna(row.get('capitalisation_globale')) else None
                    }
                    
                    # Variation journalière — comparaison avec la veille valide (id précédent)
                    try:
                        current_id = int(row.get('id'))
                        query_prev = """
                        SELECT brvm_composite
                        FROM new_market_indicators 
                        WHERE brvm_composite IS NOT NULL
                          AND brvm_composite > 0
                          AND id < %s
                        ORDER BY id DESC 
                        LIMIT 1;
                        """
                        df_prev = pd.read_sql(query_prev, self.db_conn, params=(current_id,))
                        
                        if not df_prev.empty and pd.notna(df_prev.iloc[0]['brvm_composite']):
                            prev_composite = float(df_prev.iloc[0]['brvm_composite'])
                            current_composite = float(composite)
                            var_day = ((current_composite - prev_composite) / prev_composite) * 100
                            indicators['composite_var_day'] = round(var_day, 2)
                        else:
                            indicators['composite_var_day'] = None
                    except Exception as e:
                        logging.warning(f"⚠️ Calcul variation journalière échoué: {e}")
                        indicators['composite_var_day'] = None
                    
                    # Historique 100 derniers jours valides, ordonné par id
                    try:
                        query_hist = """
                        SELECT id, brvm_composite, capitalisation_globale, extraction_date
                        FROM new_market_indicators
                        WHERE brvm_composite IS NOT NULL
                          AND brvm_composite > 0
                          AND extraction_date >= CURRENT_DATE - INTERVAL '150 days'
                        ORDER BY extraction_date DESC
                        LIMIT 100;
                        """
                        df_hist = pd.read_sql(query_hist, self.db_conn)
                        if not df_hist.empty:
                            # Remettre dans l'ordre chronologique (id croissant = du plus ancien au plus récent)
                            df_hist = df_hist.sort_values('id').reset_index(drop=True)
                            indicators['history_100d'] = df_hist
                            logging.info(f"   📊 Historique BRVM: {len(df_hist)} jours valides chargés "
                                        f"({df_hist.iloc[0]['extraction_date']} → {df_hist.iloc[-1]['extraction_date']})")
                        else:
                            indicators['history_100d'] = None
                    except Exception as e:
                        logging.warning(f"⚠️ Récupération historique 100j échouée: {e}")
                        indicators['history_100d'] = None
                    
                    return indicators
            return None
        except Exception as e:
            logging.error(f"❌ Erreur récupération indicateurs: {e}")
            return None

    def _get_historical_data_100days(self, company_id):
        """Récupère les 100 derniers jours de données historiques"""
        query = f"""
        SELECT trade_date, price, volume, company_capitalization
        FROM historical_data
        WHERE company_id = {company_id}
        ORDER BY trade_date DESC
        LIMIT 100;
        """
        
        try:
            df = pd.read_sql(query, self.db_conn)
            if not df.empty:
                df = df.sort_values('trade_date')
            return df
        except Exception as e:
            logging.error(f"❌ Erreur récupération historique: {e}")
            return pd.DataFrame()

    def _get_all_data_from_db(self):
        """
        ✅ V30.2: Récupération des données en 4 requêtes séparées + fusion Python
        Garantit que TOUTES les analyses fondamentales remontent, indépendamment
        de la présence de données de marché récentes.
        """
        logging.info("📂 Récupération des données...")
        
        # 1. Récupérer toutes les sociétés
        companies_query = """
        SELECT id, symbol, name, sector
        FROM companies
        ORDER BY symbol;
        """
        companies_df = pd.read_sql(companies_query, self.db_conn)
        logging.info(f"   ✅ {len(companies_df)} société(s) trouvée(s)")
        
        # 2. Récupérer les dernières données historiques (30 derniers jours)
        date_limite = (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d')
        hist_query = f"""
        WITH ranked_data AS (
            SELECT 
                company_id,
                id as historical_data_id,
                trade_date,
                price,
                volume,
                ROW_NUMBER() OVER(PARTITION BY company_id ORDER BY trade_date DESC) as rn
            FROM historical_data
            WHERE trade_date >= '{date_limite}'
        )
        SELECT 
            company_id,
            historical_data_id,
            trade_date,
            price,
            volume
        FROM ranked_data
        WHERE rn = 1;
        """
        hist_df = pd.read_sql(hist_query, self.db_conn)
        logging.info(f"   ✅ {len(hist_df)} société(s) avec données historiques récentes")
        
        # 3. Récupérer les analyses techniques
        tech_query = """
        SELECT 
            historical_data_id,
            mm20, mm50, mm_decision,
            bollinger_superior, bollinger_inferior, bollinger_decision,
            macd_line, signal_line, macd_decision,
            rsi, rsi_decision,
            stochastic_k, stochastic_d, stochastic_decision
        FROM technical_analysis;
        """
        tech_df = pd.read_sql(tech_query, self.db_conn)
        logging.info(f"   ✅ {len(tech_df)} enregistrements techniques")
        
        # 4. ✅ Récupérer TOUTES les analyses fondamentales (sans filtre de date)
        fund_query = """
        SELECT 
            company_id,
            report_title,
            report_date,
            analysis_summary
        FROM fundamental_analysis
        WHERE analysis_summary IS NOT NULL
          AND analysis_summary <> ''
        ORDER BY company_id, report_date DESC NULLS LAST;
        """
        fund_df = pd.read_sql(fund_query, self.db_conn)
        logging.info(f"   ✅ {len(fund_df)} analyses fondamentales trouvées au total")
        
        if not fund_df.empty:
            fund_counts = fund_df.groupby('company_id').size().to_dict()
            companies_with_fund = len(fund_counts)
            logging.info(f"   📊 {companies_with_fund} société(s) ont des analyses fondamentales")
            for cid, count in fund_counts.items():
                sym = companies_df[companies_df['id'] == cid]['symbol'].values
                if len(sym) > 0:
                    logging.info(f"      - {sym[0]}: {count} analyse(s)")
        
        # ── Fusion Python ─────────────────────────────────────────────────────────
        result_rows = []
        
        for _, company in companies_df.iterrows():
            company_id   = company['id']
            symbol       = company['symbol']
            company_name = company['name']
            sector       = company['sector']
            
            # Données historiques
            hist_data = hist_df[hist_df['company_id'] == company_id]
            if not hist_data.empty:
                hist_row          = hist_data.iloc[0]
                historical_data_id = hist_row['historical_data_id']
                trade_date        = hist_row['trade_date']
                price             = hist_row['price']
                volume            = hist_row['volume']
                
                # Données techniques
                tech_data = tech_df[tech_df['historical_data_id'] == historical_data_id]
                if not tech_data.empty:
                    tech_row            = tech_data.iloc[0]
                    mm20                = tech_row['mm20']
                    mm50                = tech_row['mm50']
                    mm_decision         = tech_row['mm_decision']
                    bollinger_superior  = tech_row['bollinger_superior']
                    bollinger_inferior  = tech_row['bollinger_inferior']
                    bollinger_decision  = tech_row['bollinger_decision']
                    macd_line           = tech_row['macd_line']
                    signal_line         = tech_row['signal_line']
                    macd_decision       = tech_row['macd_decision']
                    rsi                 = tech_row['rsi']
                    rsi_decision        = tech_row['rsi_decision']
                    stochastic_k        = tech_row['stochastic_k']
                    stochastic_d        = tech_row['stochastic_d']
                    stochastic_decision = tech_row['stochastic_decision']
                else:
                    mm20 = mm50 = mm_decision = None
                    bollinger_superior = bollinger_inferior = bollinger_decision = None
                    macd_line = signal_line = macd_decision = None
                    rsi = rsi_decision = None
                    stochastic_k = stochastic_d = stochastic_decision = None
            else:
                historical_data_id = trade_date = price = volume = None
                mm20 = mm50 = mm_decision = None
                bollinger_superior = bollinger_inferior = bollinger_decision = None
                macd_line = signal_line = macd_decision = None
                rsi = rsi_decision = None
                stochastic_k = stochastic_d = stochastic_decision = None
            
            # ✅ Construire fundamental_summaries avec séparateurs compatibles avec le parsing existant
            company_fund_data = fund_df[fund_df['company_id'] == company_id]
            fundamental_summaries     = None
            nb_rapports_fondamentaux  = 0
            
            if not company_fund_data.empty:
                parts_list = []
                for _, fund_row in company_fund_data.iterrows():
                    title   = fund_row['report_title'] or 'Sans titre'
                    date    = fund_row['report_date']
                    date_s  = date.strftime('%Y-%m-%d') if hasattr(date, 'strftime') and date else 'Date inconnue'
                    summary = fund_row['analysis_summary'] or ''
                    if summary:
                        parts_list.append(f"{title}###SEP_FIELD###{date_s}###SEP_FIELD###{summary}")
                
                if parts_list:
                    fundamental_summaries    = '###SEP_REPORT###'.join(parts_list)
                    nb_rapports_fondamentaux = len(parts_list)
                    logging.info(f"   📄 {symbol}: {nb_rapports_fondamentaux} rapport(s) fondamental/aux chargé(s)")
            
            result_rows.append({
                'company_id':              company_id,
                'symbol':                  symbol,
                'company_name':            company_name,
                'sector':                  sector,
                'trade_date':              trade_date,
                'price':                   price,
                'volume':                  volume,
                'mm20':                    mm20,
                'mm50':                    mm50,
                'mm_decision':             mm_decision,
                'bollinger_superior':      bollinger_superior,
                'bollinger_inferior':      bollinger_inferior,
                'bollinger_decision':      bollinger_decision,
                'macd_line':               macd_line,
                'signal_line':             signal_line,
                'macd_decision':           macd_decision,
                'rsi':                     rsi,
                'rsi_decision':            rsi_decision,
                'stochastic_k':            stochastic_k,
                'stochastic_d':            stochastic_d,
                'stochastic_decision':     stochastic_decision,
                'fundamental_summaries':   fundamental_summaries,
                'nb_rapports_fondamentaux': nb_rapports_fondamentaux,
            })
        
        result_df = pd.DataFrame(result_rows)
        
        # Statistiques finales
        has_fundamental = result_df['fundamental_summaries'].notna().sum()
        logging.info(f"   ✅ Fusion terminée: {len(result_df)} société(s)")
        logging.info(f"   📊 {has_fundamental}/{len(result_df)} ont des analyses fondamentales dans le résultat final")
        
        if has_fundamental > 0:
            syms_with_fund = result_df[result_df['fundamental_summaries'].notna()]['symbol'].tolist()
            logging.info(f"   📋 Sociétés avec analyses: {', '.join(syms_with_fund)}")
        
        return result_df

    def _get_predictions_from_db(self):
        """Récupération des prédictions avec bornes IC et niveau de confiance"""
        logging.info("🔮 Récupération des prédictions (avec IC)...")

        query = """
        WITH latest_run AS (
            SELECT company_id, MAX(run_date) AS last_run
            FROM predictions
            GROUP BY company_id
        ),
        latest_predictions AS (
            SELECT
                p.company_id,
                p.prediction_date,
                p.predicted_price,
                p.lower_bound,
                p.upper_bound,
                p.confidence_level
            FROM predictions p
            JOIN latest_run lr
              ON p.company_id = lr.company_id
             AND p.run_date   = lr.last_run
            WHERE p.prediction_date >= CURRENT_DATE
        )
        SELECT
            c.symbol,
            lp.prediction_date,
            lp.predicted_price,
            lp.lower_bound,
            lp.upper_bound,
            lp.confidence_level
        FROM companies c
        LEFT JOIN latest_predictions lp ON c.id = lp.company_id
        ORDER BY c.symbol, lp.prediction_date;
        """

        try:
            df = pd.read_sql(query, self.db_conn)
            logging.info(f"   ✅ {len(df)} prédiction(s) chargées (avec IC)")
            return df
        except Exception as e:
            logging.error(f"❌ Erreur prédictions: {e}")
            return pd.DataFrame()


    # =========================================================================
    # PHASE 1 — A : Graphique cours 100j + volumes
    # =========================================================================

    def _generate_price_chart(self, symbol, hist_df):
        """
        Génère courbe de cours + volumes sans GridSpec ni tight_layout
        pour éviter les warnings et crashs silencieux.
        """
        if not MATPLOTLIB_OK or hist_df is None or hist_df.empty or len(hist_df) < 5:
            return None
        try:
            fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(9, 4.5),
                                            gridspec_kw={'height_ratios': [3, 1]})
            fig.patch.set_facecolor('white')

            dates  = pd.to_datetime(hist_df['trade_date'])
            prices = hist_df['price'].astype(float)
            vols   = hist_df['volume'].astype(float) if 'volume' in hist_df.columns else pd.Series([0]*len(hist_df))

            ax1.plot(dates, prices, color='#1a5276', linewidth=1.8, zorder=3)
            ax1.fill_between(dates, prices, prices.min(), alpha=0.08, color='#1a5276')

            idx_max = prices.idxmax(); idx_min = prices.idxmin()
            ax1.annotate(f"{prices[idx_max]:,.0f}",
                         xy=(dates[idx_max], prices[idx_max]),
                         fontsize=7, color='#27ae60', fontweight='bold',
                         xytext=(0, 6), textcoords='offset points', ha='center')
            ax1.annotate(f"{prices[idx_min]:,.0f}",
                         xy=(dates[idx_min], prices[idx_min]),
                         fontsize=7, color='#c0392b', fontweight='bold',
                         xytext=(0, -12), textcoords='offset points', ha='center')

            evol = ((prices.iloc[-1] - prices.iloc[0]) / prices.iloc[0] * 100) if prices.iloc[0] else 0
            col  = '#27ae60' if evol >= 0 else '#c0392b'
            sign = '+' if evol >= 0 else ''
            ax1.set_title(f"{symbol} — Cours 100 derniers jours  ({sign}{evol:.1f}%)",
                          fontsize=10, fontweight='bold', color=col, pad=6)
            ax1.set_ylabel("Prix (FCFA)", fontsize=8)
            ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:,.0f}"))
            ax1.grid(True, linestyle='--', alpha=0.4, color='#aaaaaa')
            ax1.tick_params(axis='both', labelsize=7)
            plt.setp(ax1.get_xticklabels(), visible=False)
            ax1.spines[['top','right']].set_visible(False)

            bar_colors = ['#27ae60' if p >= prices.iloc[max(0,i-1)] else '#c0392b'
                          for i, p in enumerate(prices)]
            ax2.bar(dates, vols, color=bar_colors, alpha=0.65, width=0.8)
            ax2.set_ylabel("Volume", fontsize=7)
            ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x/1000:.0f}k" if x >= 1000 else f"{x:.0f}"))
            ax2.xaxis.set_major_formatter(mdates.DateFormatter('%d/%m'))
            ax2.xaxis.set_major_locator(mdates.WeekdayLocator(byweekday=0, interval=2))
            plt.setp(ax2.get_xticklabels(), rotation=30, fontsize=7)
            ax2.grid(True, linestyle='--', alpha=0.3, color='#aaaaaa', axis='y')
            ax2.spines[['top','right']].set_visible(False)
            ax2.set_xlabel("Date", fontsize=8)

            fig.subplots_adjust(left=0.09, right=0.97, top=0.88, bottom=0.18, hspace=0.12)
            buf = io.BytesIO()
            fig.savefig(buf, format='png', dpi=130, bbox_inches='tight', facecolor='white')
            buf.seek(0)
            plt.close(fig)
            return buf
        except Exception as e:
            logging.warning(f"⚠️  Graphique {symbol}: {e}")
            try: plt.close('all')
            except Exception: pass
            return None

    def _generate_composite_chart(self, df_hist):
        """
        Génère DEUX courbes d'évolution séparées (BytesIO PNG) :
        - buf_comp : courbe lisse de l'indice BRVM Composite
        - buf_cap  : courbe lisse de la capitalisation boursière (Mds FCFA)
        Utilise subplots() simple sans GridSpec pour éviter les warnings tight_layout.
        """
        if not MATPLOTLIB_OK or df_hist is None or df_hist.empty or len(df_hist) < 5:
            return None, None

        # ── Données ───────────────────────────────────────────────────────
        df = df_hist.copy().reset_index(drop=True)
        dates = pd.to_datetime(df['extraction_date'])
        comp  = df['brvm_composite'].astype(float)

        # Capitalisation : nettoyer les nulls et normaliser en Mds FCFA
        cap_raw = df['capitalisation_globale'].copy()
        cap_raw = pd.to_numeric(cap_raw, errors='coerce')
        cap_raw = cap_raw.ffill().bfill()
        # Détecter l'unité : si médiane > 1e9 → valeurs en FCFA brut, diviser par 1e9
        if cap_raw.median() > 1e9:
            cap = cap_raw / 1e9
        else:
            cap = cap_raw   # déjà en Mds FCFA

        def _style_ax(ax, title, ylabel, color_title):
            """Style commun pour les deux axes."""
            ax.set_title(title, fontsize=11, fontweight='bold', color=color_title, pad=8)
            ax.set_ylabel(ylabel, fontsize=9)
            ax.set_xlabel("", fontsize=9)
            # Format court : 01/25 pour janvier 2025
            ax.xaxis.set_major_formatter(mdates.DateFormatter('%m/%y'))
            ax.xaxis.set_major_locator(mdates.WeekdayLocator(byweekday=0, interval=3))
            plt.setp(ax.get_xticklabels(), rotation=45, fontsize=8, ha='right')
            ax.grid(True, linestyle='--', alpha=0.35, color='#cccccc')
            ax.spines[['top', 'right']].set_visible(False)
            ax.tick_params(axis='y', labelsize=8)

        # ── Graphique 1 : Courbe Indice BRVM Composite ────────────────────
        buf_comp = None
        try:
            evol  = ((comp.iloc[-1] - comp.iloc[0]) / comp.iloc[0] * 100) if comp.iloc[0] else 0
            sign  = '+' if evol >= 0 else ''
            col_t = '#27ae60' if evol >= 0 else '#c0392b'

            fig1, ax1 = plt.subplots(figsize=(11, 4), dpi=110)
            fig1.patch.set_facecolor('white')

            # Fond gris uniforme (évite le triangle du fill_between)
            ax1.set_facecolor('#EEF2F7')

            # Courbe principale
            ax1.plot(dates, comp, color='#154360', linewidth=2.0, zorder=3)

            # Annotations min / max
            i_max = comp.idxmax(); i_min = comp.idxmin()
            ax1.annotate(f"{comp[i_max]:,.2f} pts",
                         xy=(dates[i_max], comp[i_max]),
                         fontsize=8, color='#27ae60', fontweight='bold',
                         xytext=(0, 10), textcoords='offset points', ha='center',
                         arrowprops=dict(arrowstyle='->', color='#27ae60', lw=1.0))
            ax1.annotate(f"{comp[i_min]:,.2f} pts",
                         xy=(dates[i_min], comp[i_min]),
                         fontsize=8, color='#c0392b', fontweight='bold',
                         xytext=(0, -15), textcoords='offset points', ha='center',
                         arrowprops=dict(arrowstyle='->', color='#c0392b', lw=1.0))

            # Dernière valeur
            ax1.annotate(f" {comp.iloc[-1]:,.2f}",
                         xy=(dates.iloc[-1], comp.iloc[-1]),
                         fontsize=8.5, color='#154360', fontweight='bold',
                         xytext=(6, 0), textcoords='offset points', va='center')

            ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:,.1f}"))
            y_pad = (comp.max() - comp.min()) * 0.20 if comp.max() != comp.min() else comp.max() * 0.05
            ax1.set_ylim(comp.min() - y_pad * 0.3, comp.max() + y_pad)
            ax1.set_xlim(dates.iloc[0], dates.iloc[-1])
            _style_ax(ax1,
                      f"Indice BRVM Composite — {len(comp)} jours ({sign}{evol:.2f}%)",
                      "Points",
                      col_t)

            fig1.subplots_adjust(left=0.09, right=0.97, top=0.88, bottom=0.18)
            buf_comp = io.BytesIO()
            fig1.savefig(buf_comp, format='png', bbox_inches='tight',
                         facecolor='white', edgecolor='none')
            buf_comp.seek(0)
            plt.close(fig1)
        except Exception as e:
            logging.warning(f"⚠️  Graphique composite: {e}")
            try: plt.close('all')
            except Exception: pass

        # ── Graphique 2 : Courbe Capitalisation Boursière ─────────────────
        buf_cap = None
        try:
            evol_c  = ((cap.iloc[-1] - cap.iloc[0]) / cap.iloc[0] * 100) if cap.iloc[0] else 0
            sign_c  = '+' if evol_c >= 0 else ''
            col_c   = '#27ae60' if evol_c >= 0 else '#c0392b'

            fig2, ax2 = plt.subplots(figsize=(11, 4), dpi=110)
            fig2.patch.set_facecolor('white')

            # Fond gris uniforme (évite le triangle du fill_between)
            ax2.set_facecolor('#EEF2F7')

            # Courbe principale
            ax2.plot(dates, cap, color='#1a5276', linewidth=2.0, zorder=3)

            # Annotations min / max
            i_max2 = cap.idxmax(); i_min2 = cap.idxmin()
            ax2.annotate(f"{cap[i_max2]:,.0f} Mds",
                         xy=(dates[i_max2], cap[i_max2]),
                         fontsize=8, color='#27ae60', fontweight='bold',
                         xytext=(0, 10), textcoords='offset points', ha='center',
                         arrowprops=dict(arrowstyle='->', color='#27ae60', lw=1.0))
            ax2.annotate(f"{cap[i_min2]:,.0f} Mds",
                         xy=(dates[i_min2], cap[i_min2]),
                         fontsize=8, color='#c0392b', fontweight='bold',
                         xytext=(0, -15), textcoords='offset points', ha='center',
                         arrowprops=dict(arrowstyle='->', color='#c0392b', lw=1.0))

            # Dernière valeur
            ax2.annotate(f" {cap.iloc[-1]:,.0f}",
                         xy=(dates.iloc[-1], cap.iloc[-1]),
                         fontsize=8.5, color='#1a5276', fontweight='bold',
                         xytext=(6, 0), textcoords='offset points', va='center')

            ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:,.0f}"))
            y_pad2 = (cap.max() - cap.min()) * 0.20 if cap.max() != cap.min() else cap.max() * 0.05
            ax2.set_ylim(cap.min() - y_pad2 * 0.3, cap.max() + y_pad2)
            ax2.set_xlim(dates.iloc[0], dates.iloc[-1])
            _style_ax(ax2,
                      f"Capitalisation Boursière BRVM — {len(cap)} jours ({sign_c}{evol_c:.2f}%)",
                      "Mds FCFA",
                      col_c)

            fig2.subplots_adjust(left=0.09, right=0.97, top=0.88, bottom=0.18)
            buf_cap = io.BytesIO()
            fig2.savefig(buf_cap, format='png', bbox_inches='tight',
                         facecolor='white', edgecolor='none')
            buf_cap.seek(0)
            plt.close(fig2)
        except Exception as e:
            logging.warning(f"⚠️  Graphique capitalisation: {e}")
            try: plt.close('all')
            except Exception: pass

        return buf_comp, buf_cap

    def _generate_price_chart_with_predictions(self, symbol, hist_df, predictions):
        """
        Graphique cours reels (bleu) + cours predits (orange pointille) + IC + volumes.
        Distinction visuelle claire entre historique et previsions.
        """
        if not MATPLOTLIB_OK or hist_df is None or hist_df.empty or len(hist_df) < 5:
            return None
        try:
            fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(9, 4.5),
                                            gridspec_kw={'height_ratios': [3, 1]})
            dates_r  = pd.to_datetime(hist_df['trade_date'])
            prices_r = hist_df['price'].astype(float)
            vols     = hist_df['volume'].astype(float) if 'volume' in hist_df.columns else pd.Series([0]*len(hist_df))
            # Cours reels
            ax1.plot(dates_r, prices_r, color='#1a5276', linewidth=1.8, zorder=3, label='Cours r\u00e9els')
            ax1.fill_between(dates_r, prices_r, prices_r.min(), alpha=0.07, color='#1a5276')
            idx_max = prices_r.idxmax(); idx_min = prices_r.idxmin()
            ax1.annotate(f"{prices_r[idx_max]:,.0f}", xy=(dates_r[idx_max], prices_r[idx_max]),
                         fontsize=7, color='#27ae60', fontweight='bold', xytext=(0, 6), textcoords='offset points', ha='center')
            ax1.annotate(f"{prices_r[idx_min]:,.0f}", xy=(dates_r[idx_min], prices_r[idx_min]),
                         fontsize=7, color='#c0392b', fontweight='bold', xytext=(0, -12), textcoords='offset points', ha='center')
            # Cours predits
            has_preds = False
            if predictions:
                pred_dates  = []
                pred_prices = []
                pred_lower  = []
                pred_upper  = []
                for p in predictions:
                    pd_d = pd.to_datetime(p.get('date'))
                    pp   = p.get('price')
                    pl   = p.get('lower_bound')
                    pu   = p.get('upper_bound')
                    if pd_d is not None and pp is not None:
                        pred_dates.append(pd_d)
                        pred_prices.append(float(pp))
                        pred_lower.append(float(pl) if pl else float(pp)*0.98)
                        pred_upper.append(float(pu) if pu else float(pp)*1.02)
                if pred_dates:
                    has_preds = True
                    # Trait de jonction dernier reel -> premier predit
                    ax1.plot([dates_r.iloc[-1], pred_dates[0]],
                             [prices_r.iloc[-1], pred_prices[0]],
                             color='#e67e22', linewidth=1.2, linestyle='--', alpha=0.5)
                    # Courbe predite
                    ax1.plot(pred_dates, pred_prices, color='#e67e22', linewidth=2.0,
                             linestyle='--', marker='o', markersize=4, zorder=4,
                             label='Cours pr\u00e9dits (ML)')
                    # Zone intervalle de confiance
                    ax1.fill_between(pred_dates, pred_lower, pred_upper,
                                     alpha=0.18, color='#e67e22', label='IC 90%')
                    # Annotations J+1 et dernier jour
                    ax1.annotate(f"J+1\n{pred_prices[0]:,.0f}",
                                 xy=(pred_dates[0], pred_prices[0]),
                                 fontsize=6.5, color='#e67e22', fontweight='bold',
                                 xytext=(4, 8), textcoords='offset points')
                    ax1.annotate(f"J+{len(pred_dates)}\n{pred_prices[-1]:,.0f}",
                                 xy=(pred_dates[-1], pred_prices[-1]),
                                 fontsize=6.5, color='#e67e22', fontweight='bold',
                                 xytext=(4, 8), textcoords='offset points')
                    # Ligne verticale separatrice
                    ax1.axvline(x=dates_r.iloc[-1], color='#7f8c8d', linewidth=0.9,
                                linestyle=':', alpha=0.8)
                    y_top = prices_r.max() * 1.01
                    ax1.text(dates_r.iloc[-1], y_top, "  Aujourd'hui",
                             fontsize=6.5, color='#7f8c8d', va='top')
            evol = ((prices_r.iloc[-1] - prices_r.iloc[0]) / prices_r.iloc[0] * 100) if prices_r.iloc[0] else 0
            sign = '+' if evol >= 0 else ''
            color_t = '#27ae60' if evol >= 0 else '#c0392b'
            pred_lbl = "  +  pr\u00e9dictions J+1\u2192J+10" if has_preds else ""
            ax1.set_title(f"{symbol} \u2014 Cours r\u00e9els ({sign}{evol:.1f}%){pred_lbl}",
                          fontsize=10, fontweight='bold', color=color_t, pad=6)
            ax1.set_ylabel("Prix (FCFA)", fontsize=8)
            ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:,.0f}"))
            ax1.grid(True, linestyle='--', alpha=0.4, color='#aaaaaa')
            ax1.tick_params(axis='both', labelsize=7)
            ax1.legend(fontsize=7, loc='upper left')
            plt.setp(ax1.get_xticklabels(), visible=False)
            ax1.spines[['top','right']].set_visible(False)
            # Volumes
            bar_colors = ['#27ae60' if p >= prices_r.iloc[max(0,i-1)] else '#c0392b' for i, p in enumerate(prices_r)]
            ax2.bar(dates_r, vols, color=bar_colors, alpha=0.65, width=0.8)
            ax2.set_ylabel("Volume", fontsize=7)
            ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x/1000:.0f}k" if x >= 1000 else f"{x:.0f}"))
            ax2.xaxis.set_major_formatter(mdates.DateFormatter('%d/%m'))
            ax2.xaxis.set_major_locator(mdates.WeekdayLocator(byweekday=0, interval=2))
            plt.setp(ax2.get_xticklabels(), rotation=30, fontsize=7)
            ax2.grid(True, linestyle='--', alpha=0.3, color='#aaaaaa', axis='y')
            ax2.spines[['top','right']].set_visible(False)
            ax2.set_xlabel("Date", fontsize=8)
            fig.patch.set_facecolor('white')
            fig.subplots_adjust(left=0.09, right=0.97, top=0.88, bottom=0.18, hspace=0.12)
            buf = io.BytesIO()
            fig.savefig(buf, format='png', dpi=130, bbox_inches='tight', facecolor='white')
            buf.seek(0)
            plt.close(fig)
            return buf
        except Exception as e:
            logging.warning(f"\u26a0\ufe0f  Graphique predit {symbol}: {e}")
            try: plt.close('all')
            except Exception: pass
            return None

    # =========================================================================
    # PHASE 1 — C : Score composite d'investissement 0–100
    # =========================================================================


    # =========================================================================
    # PHASE 3 — I : Portefeuilles modèles
    # =========================================================================

    def _build_model_portfolios(self, all_company_data):
        """
        Construit 3 portefeuilles modèles à partir des scores, risques et liquidités.

        Défensif  : score ≥ 45 + risque Faible + liquidité Bonne/Excellente
        Équilibré : score ≥ 50 + risque Faible ou Moyen
        Offensif  : score ≥ 60 (tous risques) — les meilleures opportunités
        """
        defensif   = []
        equilibre  = []
        offensif   = []

        for sym, d in all_company_data.items():
            sc  = d.get('investment_score', 0)
            rl  = str(d.get('risk_level','')).lower()
            rec = str(d.get('recommendation','')).upper()
            if 'VENTE' in rec:
                continue     # aucun portefeuille ne prend un signal vente

            try:
                rd = json.loads(d.get('risk_details','{}')) if isinstance(d.get('risk_details'), str) else (d.get('risk_details') or {})
            except Exception:
                rd = {}
            liq_str = str(rd.get('liquidite','')).lower()
            liq_ok  = any(k in liq_str for k in ['excellente','bonne'])

            name  = d.get('company_name','')
            price = d.get('current_price') or 0
            entry = {'sym': sym, 'name': name, 'price': price, 'score': sc,
                     'risk': d.get('risk_level','—'), 'rec': rec}

            if sc >= 45 and 'faible' in rl and liq_ok:
                defensif.append(entry)
            if sc >= 50 and rl in ('faible','moyen'):
                equilibre.append(entry)
            if sc >= 60:
                offensif.append(entry)

        # Trier par score décroissant, limiter à 8 titres par portefeuille
        for lst in (defensif, equilibre, offensif):
            lst.sort(key=lambda x: x['score'], reverse=True)

        # Pondérations équipondérées (sauf cash 15 % dans défensif)
        def _weights(lst, cash_pct=0):
            n = min(len(lst), 8)
            if n == 0: return []
            w = round((100 - cash_pct) / n, 1)
            return [(e, w) for e in lst[:n]]

        return {
            'defensif':  (_weights(defensif, cash_pct=15),  15),
            'equilibre': (_weights(equilibre, cash_pct=10), 10),
            'offensif':  (_weights(offensif,  cash_pct=5),   5),
        }

    def _compute_investment_score(self, company_data):
        """
        Score composite 0–100 :
          Technique    30 %  (signaux achat/vente)
          Fondamental  40 %  (présence et qualité des données fondamentales)
          Risque       20 %  (risk_score inversé)
          Liquidité    10 %  (volume moyen)

        Retourne un entier 0–100 et un label (Excellent/Bon/Moyen/Faible).
        """
        score = 0.0

        # — Technique (30 %) —
        tech_keys = ['mm_decision','bollinger_decision','macd_decision',
                     'rsi_decision','stochastic_decision']
        signals = [str(company_data.get(k,'')).upper() for k in tech_keys
                   if company_data.get(k)]
        n = len(signals) or 1
        buy  = sum(1 for s in signals if 'ACHAT' in s)
        sell = sum(1 for s in signals if 'VENTE' in s)
        tech_pct = buy / n                  # 0 → 1
        score += tech_pct * 30

        # — Fondamental (40 %) —
        fund_text = company_data.get('fundamental_analysis','')
        if fund_text and len(fund_text) > 200:
            rec = str(company_data.get('recommendation','')).upper()
            if   'ACHAT FORT' in rec: fund_pts = 40
            elif 'ACHAT'      in rec: fund_pts = 32
            elif 'CONSERVER'  in rec: fund_pts = 20
            elif 'VENTE'      in rec: fund_pts = 8
            else:                     fund_pts = 16
            # Bonus si brvm_rapports disponibles
            if company_data.get('brvm_rapports_raw'):
                fund_pts = min(40, fund_pts + 5)
        else:
            fund_pts = 10   # pénalité données absentes
        score += fund_pts

        # — Risque (20 %) — risk_score est déjà 0–100 (haut = risqué)
        risk_raw = float(company_data.get('risk_score', 50))
        score += (1 - risk_raw / 100) * 20

        # — Liquidité (10 %) —
        risk_details_raw = company_data.get('risk_details','{}')
        try:
            rd = json.loads(risk_details_raw) if isinstance(risk_details_raw, str) else risk_details_raw
            liq_str = str(rd.get('liquidite','')).lower()
            if   'excellente' in liq_str: liq_pts = 10
            elif 'bonne'      in liq_str: liq_pts = 7
            elif 'faible'     in liq_str: liq_pts = 2
            else:                         liq_pts = 5
        except Exception:
            liq_pts = 5
        score += liq_pts

        score = max(0, min(100, round(score)))

        if   score >= 70: label = 'Excellent'
        elif score >= 55: label = 'Bon'
        elif score >= 40: label = 'Moyen'
        else:             label = 'Faible'

        return score, label

    # =========================================================================
    # PHASE 1 — D : Résumé Exécutif automatique
    # =========================================================================

    def _build_executive_summary(self, all_company_data, market_indicators):
        """
        Génère les données du résumé exécutif depuis les données agrégées.
        Retourne un dict avec toutes les stats nécessaires.
        """
        total   = len(all_company_data)
        achats  = sum(1 for d in all_company_data.values() if 'ACHAT' in str(d.get('recommendation','')).upper())
        ventes  = sum(1 for d in all_company_data.values() if 'VENTE' in str(d.get('recommendation','')).upper())
        neutres = total - achats - ventes

        # Meilleure société par score composite
        scored = [(sym, d.get('investment_score', 0)) for sym, d in all_company_data.items()]
        scored.sort(key=lambda x: x[1], reverse=True)
        top_sym    = scored[0][0]  if scored else '—'
        top_score  = scored[0][1]  if scored else 0

        # Secteur le plus performant
        sec_perf = {}
        for d in all_company_data.values():
            sec = d.get('sector','Autre') or 'Autre'
            p   = d.get('price_evolution_100d') or 0
            sec_perf.setdefault(sec, []).append(p)
        sec_avg = {s: sum(v)/len(v) for s,v in sec_perf.items() if v}
        best_sector  = max(sec_avg, key=sec_avg.get) if sec_avg else '—'
        best_sec_pct = sec_avg.get(best_sector, 0)
        worst_sector = min(sec_avg, key=sec_avg.get) if sec_avg else '—'

        # Titres faible liquidité
        low_liq = sum(
            1 for d in all_company_data.values()
            if 'faible' in str(d.get('risk_details','{}')).lower()
            and 'liquidite' in str(d.get('risk_details','{}')).lower()
        )

        # Marché haussier / neutre / baissier
        composite = market_indicators.get('composite') if market_indicators else None
        var_day    = market_indicators.get('composite_var_day') if market_indicators else None
        if   var_day and var_day >  1.0:  marche = 'HAUSSIER 📈'
        elif var_day and var_day < -1.0:  marche = 'BAISSIER 📉'
        else:                              marche = 'NEUTRE / STABLE ➡️'

        # Divergences majeures (tech≠fond)
        divergences = sum(
            1 for d in all_company_data.values()
            if d.get('technical_decision','') != d.get('fundamental_decision','')
            and d.get('technical_decision','') not in ('','NEUTRE')
        )

        return {
            'total': total, 'achats': achats, 'ventes': ventes, 'neutres': neutres,
            'top_sym': top_sym, 'top_score': top_score,
            'best_sector': best_sector, 'best_sec_pct': best_sec_pct,
            'worst_sector': worst_sector,
            'low_liq': low_liq, 'marche': marche,
            'divergences': divergences,
            'composite': composite, 'var_day': var_day,
        }

    # =========================================================================
    # PHASE 1 — E : Faits marquants depuis google_alerts_rapports
    # =========================================================================

    def _get_google_alerts_events(self):
        """
        Récupère les alertes depuis google_alerts_rapports.
        Exploite toutes les colonnes disponibles : mail_subject, alert_keyword,
        points_cles, mot_cle, rapport_type, pertinence, sentiment.
        Fallback sur new_market_events si vide.
        """
        logging.info("📰 Récupération google_alerts_rapports (enrichi)...")
        query = """
        SELECT
            mail_date,
            mail_subject,
            titre,
            resume,
            points_cles,
            sentiment,
            pertinence,
            categorie,
            rapport_type,
            alert_keyword,
            mot_cle,
            source_url
        FROM google_alerts_rapports
        WHERE (resume IS NOT NULL AND resume <> '')
           OR (mail_subject IS NOT NULL AND mail_subject <> '')
        ORDER BY mail_date DESC NULLS LAST
        LIMIT 20;
        """
        try:
            df = pd.read_sql(query, self.db_conn)
            if not df.empty:
                logging.info(f"   ✅ {len(df)} alerte(s) google chargée(s)")
                return df
        except Exception as e:
            logging.warning(f"⚠️  google_alerts_rapports non disponible: {e}")

        # Fallback : new_market_events
        logging.info("   ↩️  Fallback → new_market_events")
        try:
            df2 = pd.read_sql(
                "SELECT event_date AS mail_date, event_summary AS resume "
                "FROM new_market_events ORDER BY event_date DESC LIMIT 10;",
                self.db_conn
            )
            return df2
        except Exception:
            return pd.DataFrame()


    # =========================================================================
    # ACTUALITÉ INTERNATIONALE / AFRICAINE / UEMOA
    # =========================================================================

    def _get_macro_news(self):
        """
        Charge les actualités macro depuis google_alerts_rapports.
        Retourne un dict structuré sur 3 types × 5 zones géographiques.

        Structure retournée :
          {
            'macroeconomique': {'international','afrique','afrique_ouest','uemoa','brvm'},
            'politique':       {'international','afrique','afrique_ouest','uemoa','brvm'},
            'financiere':      {'international','afrique','afrique_ouest','uemoa','brvm'},
            '_flat':           DataFrame complet (tous types/zones, pour comptage badge)
          }
        """
        logging.info("🌍 Chargement actualités macro (3 types × 5 zones)...")

        COLS = """
            mail_date,
            COALESCE(mail_subject, titre, '') AS mail_subject,
            COALESCE(titre, mail_subject, '') AS titre,
            COALESCE(resume, '') AS resume,
            points_cles,
            COALESCE(sentiment, 'neutre') AS sentiment,
            COALESCE(pertinence, 50) AS pertinence,
            COALESCE(categorie, '') AS categorie,
            COALESCE(source_url, '') AS source_url,
            COALESCE(source_rss, '') AS source_rss,
            COALESCE(zone, '') AS zone,
            COALESCE(type_actualite, 'macroeconomique') AS type_actualite,
            COALESCE(score_importance, 50) AS score_importance,
            COALESCE(impact_brvm, 'neutre') AS impact_brvm,
            COALESCE(impact_bourses_mondiales, '') AS impact_bourses_mondiales
        """

        TYPES  = ['macroeconomique', 'politique', 'financiere']
        ZONES  = ['international', 'afrique', 'afrique_ouest', 'uemoa', 'brvm']

        # Mots-clés de fallback par zone (si colonne zone absente/vide)
        KW_ZONE = {
            'brvm':        ["brvm","bourse abidjan","sociétés cotées","bourse régionale"],
            'uemoa':       ["uemoa","bceao","fcfa","zone franc","côte d'ivoire","sénégal","mali","burkina","niger","togo","bénin"],
            'afrique_ouest':["afrique de l'ouest","cedeao","ecowas","afrique ouest","nigeria","ghana"],
            'afrique':     ["afrique","africa","union africaine"],
            'international':["fed","bce","banque centrale","dollar","euro","wall street","pétrole","inflation","récession","g7","g20"],
        }
        KW_TYPE = {
            'macroeconomique': ["inflation","pib","croissance","taux directeur","bceao","fed","bce","matières premières","cacao","pétrole","or","récession","fmi"],
            'politique':       ["politique","gouvernement","élection","coup","sécurité","géopolitique","conflit","guerre","sanctions","transition","président"],
            'financiere':      ["bourse","marché financier","action","obligation","indice","change","forex","dollar","euro","brvm","cotation","introduction"],
        }

        def _kw_filter(kw_list, field="LOWER(COALESCE(titre,'') || ' ' || COALESCE(resume,'') || ' ' || COALESCE(zone,'') || ' ' || COALESCE(type_actualite,''))"):
            return " OR ".join([f"{field} LIKE '%{kw.replace(chr(39), chr(39)*2)}%'" for kw in kw_list])

        result = {t: {z: pd.DataFrame() for z in ZONES} for t in TYPES}
        result['_flat'] = pd.DataFrame()

        try:
            # ── Requête principale : zone ET type_actualite explicites ───────
            for t in TYPES:
                for z in ZONES:
                    df = pd.read_sql(f"""
                        SELECT {COLS} FROM google_alerts_rapports
                        WHERE zone = '{z}' AND type_actualite = '{t}'
                          AND resume IS NOT NULL AND resume <> ''
                        ORDER BY COALESCE(collecte_date, mail_date) DESC NULLS LAST
                        LIMIT 8;
                    """, self.db_conn)
                    result[t][z] = df

            # ── Fallback mots-clés pour les combinaisons vides ───────────────
            for t in TYPES:
                for z in ZONES:
                    if result[t][z].empty:
                        kw_combined = KW_ZONE.get(z, []) + KW_TYPE.get(t, [])
                        if kw_combined:
                            df = pd.read_sql(f"""
                                SELECT {COLS} FROM google_alerts_rapports
                                WHERE ({_kw_filter(kw_combined)})
                                  AND resume IS NOT NULL AND resume <> ''
                                ORDER BY COALESCE(collecte_date, mail_date) DESC NULLS LAST
                                LIMIT 5;
                            """, self.db_conn)
                            result[t][z] = df

            # ── DataFrame plat pour le badge ─────────────────────────────────
            result['_flat'] = pd.read_sql(f"""
                SELECT {COLS} FROM google_alerts_rapports
                WHERE resume IS NOT NULL AND resume <> ''
                ORDER BY COALESCE(collecte_date, mail_date) DESC NULLS LAST
                LIMIT 100;
            """, self.db_conn)

            total = sum(len(result[t][z]) for t in TYPES for z in ZONES)
            logging.info(f"   ✅ {total} actualités chargées (3 types × 5 zones)")

        except Exception as e:
            logging.error(f"❌ Erreur _get_macro_news: {e}")
            import traceback as _tb; logging.error(_tb.format_exc())

        return result

    def _generate_macro_analysis(self, macro_news, all_company_data, market_indicators):
        """
        Génère l analyse macro complète via IA.
        Nouvelle structure : 3 types × 5 zones + impact BRVM par actualité.
        """
        logging.info("🤖 Génération analyse macro (3 types × 5 zones)...")

        def _df_to_text(df, max_items=3):
            if df is None or df.empty:
                return "Aucune actualité."
            lines = []
            for _, row in df.head(max_items).iterrows():
                date_s = str(row.get('mail_date',''))[:10]
                titre  = str(row.get('titre') or row.get('mail_subject',''))[:100]
                resume = str(row.get('resume',''))[:150]
                sent   = str(row.get('sentiment','') or '').capitalize()
                lines.append(f"[{date_s}] {titre} — {resume} (Sentiment: {sent})")
            return " | ".join(lines)

        ZONES_LABELS = {
            'international': 'MONDIAL / INTERNATIONAL',
            'afrique':       'AFRICAIN',
            'afrique_ouest': "AFRIQUE DE L'OUEST",
            'uemoa':         'UEMOA / ZONE FRANC',
            'brvm':          'BRVM (BOURSE RÉGIONALE)',
        }
        TYPES_LABELS = {
            'macroeconomique': 'MACRO-ÉCONOMIQUE',
            'politique':       'POLITIQUE & GÉOPOLITIQUE',
            'financiere':      'FINANCIÈRE & MARCHÉS',
        }

        sectors = list(set(d.get('sector','') for d in all_company_data.values() if d.get('sector')))
        top_cos = sorted([(s,d) for s,d in all_company_data.items()],
                         key=lambda x: x[1].get('investment_score',0), reverse=True)[:15]
        top_str = ", ".join([f"{s}({d.get('sector','')})" for s,d in top_cos])
        composite_val = market_indicators.get('composite','N/A') if market_indicators else 'N/A'

        # Construire le contexte complet pour l'IA
        context_lines = []
        TYPES  = ['macroeconomique', 'politique', 'financiere']
        ZONES  = ['international', 'afrique', 'afrique_ouest', 'uemoa', 'brvm']
        for t in TYPES:
            t_data = macro_news.get(t, {})
            if not isinstance(t_data, dict):
                continue
            has_content = any(not df.empty for df in t_data.values() if isinstance(df, pd.DataFrame))
            if not has_content:
                continue
            context_lines.append(f"\n=== ACTUALITÉS {TYPES_LABELS[t]} ===")
            for z in ZONES:
                df = t_data.get(z, pd.DataFrame())
                if not isinstance(df, pd.DataFrame) or df.empty:
                    continue
                context_lines.append(f"-- {ZONES_LABELS[z]} --")
                context_lines.append(_df_to_text(df))

        context_str = "\n".join(context_lines) if context_lines else "Aucune actualité disponible."

        prompt = f"""Tu es un analyste financier senior spécialisé dans les marchés africains et la BRVM.

CONTEXTE BRVM :
- Indice BRVM Composite : {composite_val}
- Secteurs cotés : {', '.join(sectors)}
- Principales sociétés : {top_str}

ACTUALITÉS COLLECTÉES :
{context_str}

═══════════════════════════════════════════════════════════════
MISSION : Rédige un rapport d'analyse macro complet en FRANÇAIS.
Structure OBLIGATOIRE (respecte exactement les marqueurs ## et ###) :

## TYPE_1. ACTUALITÉS MACRO-ÉCONOMIQUES
### PLAN MONDIAL
[3-4 événements macro internationaux majeurs : Fed, BCE, PIB, inflation, matières premières]
#### Impact estimé sur la BRVM
[Analyse précise : quels secteurs/sociétés BRVM sont touchés et pourquoi]

### PLAN AFRICAIN
[2-3 faits macro africains importants : croissance, dette, FMI, investissements]
#### Impact estimé sur la BRVM
[Impact sur les sociétés cotées opérant à l'échelle africaine]

### PLAN AFRIQUE DE L'OUEST
[2-3 faits économiques pour la sous-région CEDEAO]
#### Impact estimé sur la BRVM
[Lien direct avec les entreprises opérant en Afrique de l'Ouest]

### PLAN UEMOA
[Décisions BCEAO, inflation zone, FCFA, liquidité bancaire, matières premières clés (cacao, caoutchouc)]
#### Impact estimé sur la BRVM
[Impact direct sur chaque secteur BRVM : banques, agro-industrie, distribution, énergie]

### PLAN BRVM
[Tendances de marché, volumes, sociétés en mouvement, IPO, dividendes]
#### Impact estimé sur la BRVM
[Recommandation par société cotée : SGBC, BICB, PALC, SOGC, ETIT, NTLC, SIVC, ONTBF, CFAC, etc.]

## TYPE_2. ACTUALITÉS POLITIQUES & GÉOPOLITIQUES
### PLAN MONDIAL
[Conflits, élections majeures, sanctions, relations commerciales]
#### Impact estimé sur la BRVM
[Comment l'instabilité géopolitique mondiale affecte les investisseurs de la BRVM]

### PLAN AFRICAIN
[Stabilité politique africaine, transitions, coups d'État, élections]
#### Impact estimé sur la BRVM
[Flux de capitaux, confiance des investisseurs]

### PLAN AFRIQUE DE L'OUEST
[Sécurité CEDEAO, transitions politiques, AES, relations régionales]
#### Impact estimé sur la BRVM
[Risques spécifiques pour les sociétés cotées opérant dans ces pays]

### PLAN UEMOA
[Situation politique dans les pays membres, réformes institutionnelles]
#### Impact estimé sur la BRVM
[Risques réglementaires et politiques directs pour la BRVM]

### PLAN BRVM
[Réglementation CREPMF, nouvelles cotations, gouvernance sociétés]
#### Impact estimé sur la BRVM
[Décisions réglementaires et leurs effets sur le marché]

## TYPE_3. ACTUALITÉS FINANCIÈRES & MARCHÉS
### PLAN MONDIAL
[Wall Street, indices, taux, change, flux capitaux, crypto]
#### Impact estimé sur la BRVM
[Corrélations avec les marchés émergents et effets sur les investisseurs BRVM]

### PLAN AFRICAIN
[Marchés obligataires africains, eurobonds, notation souveraine]
#### Impact estimé sur la BRVM
[Concurrence avec d'autres marchés africains pour les capitaux]

### PLAN AFRIQUE DE L'OUEST
[Bourses régionales concurrentes, flux de capitaux sous-régionaux]
#### Impact estimé sur la BRVM
[Positionnement de la BRVM face aux autres places financières]

### PLAN UEMOA
[Marché des titres publics UEMOA, taux interbancaire, liquidité]
#### Impact estimé sur la BRVM
[Arbitrage entre titres publics et actions BRVM]

### PLAN BRVM
[Performance des indices BRVM, top/flop, volumes, perspectives]
#### Impact estimé sur la BRVM
[Analyse par valeur : opportunités d'achat, titres sous pression, recommandations]

## SYNTHESE_FINALE
Niveau d'alerte global pour un investisseur BRVM : VERT / ORANGE / ROUGE
[Synthèse en 3-4 phrases. Recommandation concrète.]

═══════════════════════════════════════════════════════════════
RÈGLES IMPÉRATIVES :
- Sois factuel, base-toi UNIQUEMENT sur les actualités fournies
- Si une section manque de données : "Données insuffisantes."
- Cite les symboles BRVM (SGBC, BICB, PALC, ETIT, ONTBF, SOGC, CFAC, NTLC, SIVC, SLBC)
- Sois concis : 2-3 phrases par sous-section "Impact estimé"
- Maximum 1500 mots au total — sois synthétique
"""

        analysis_text = None
        ai_provider   = None

        # Appels IA séquentiels — sans lambda pour éviter le bug de late binding
        # Ordre : DeepSeek → Gemini → Mistral → Claude
        # max_tokens limité à 2500 pour rester dans les limites de tous les plans
        _ai_attempts = [
            ('deepseek', self._generate_analysis_with_deepseek),
            ('gemini',   self._generate_analysis_with_gemini),
            ('mistral',  self._generate_analysis_with_mistral),
            ('claude',   self._generate_analysis_with_claude),
        ]
        for name, fn in _ai_attempts:
            try:
                logging.info(f"   🤖 Tentative analyse macro via {name}...")
                text, prov = fn('MACRO', {}, prompt)
                if text and len(text.strip()) > 200:
                    analysis_text = text
                    ai_provider   = prov or name
                    logging.info(f"   ✅ Analyse macro générée via {ai_provider} ({len(text)} chars)")
                    break
                else:
                    logging.warning(f"   ⚠️ {name}: réponse trop courte ou vide ({len(text) if text else 0} chars)")
            except Exception as e:
                logging.warning(f"   ⚠️ {name} macro exception: {e}")
                import traceback as _tb; logging.debug(_tb.format_exc())
                continue

        if not analysis_text:
            logging.warning("   ⚠️ Toutes les IAs ont échoué — fallback structuré à partir des données brutes")
            analysis_text = self._build_fallback_macro_analysis(macro_news, all_company_data)
            ai_provider   = 'fallback_structuré'

        return {'analysis_text': analysis_text, 'ai_provider': ai_provider}

    def _build_fallback_macro_analysis(self, macro_news, all_company_data):
        """
        Génère une analyse macro structurée sans IA,
        à partir des données brutes (titres, résumés, sentiments, impacts).
        """
        sections = []

        # ── Fonction d'analyse automatique d'impact ───────────────────────
        def _auto_impact(df, zone):
            """Génère un texte d'impact automatique à partir des sentiments et résumés."""
            if df is None or df.empty:
                return "- Données insuffisantes pour cette zone."

            pos = [r for _, r in df.iterrows()
                   if 'positif' in str(r.get('impact_brvm','')).lower()
                   or 'positif' in str(r.get('sentiment','')).lower()]
            neg = [r for _, r in df.iterrows()
                   if 'negatif' in str(r.get('impact_brvm','')).lower()
                   or 'negatif' in str(r.get('sentiment','')).lower()]
            neu = [r for _, r in df.iterrows() if r not in pos and r not in neg]

            lines = []

            # Tendance générale
            if len(pos) > len(neg):
                lines.append(
                    f"- Tendance globalement positive sur {len(df)} actualité(s). "
                    f"{len(pos)} signal(aux) haussier(s) détecté(s) pour la BRVM."
                )
            elif len(neg) > len(pos):
                lines.append(
                    f"- Tendance globalement négative sur {len(df)} actualité(s). "
                    f"{len(neg)} signal(aux) baissier(s) détecté(s) pour la BRVM."
                )
            else:
                lines.append(
                    f"- Tendance neutre/mixte sur {len(df)} actualité(s) — "
                    f"{len(pos)} positif(s), {len(neg)} négatif(s)."
                )

            # Détailler les impacts positifs
            for row in pos[:2]:
                titre = str(row.get('titre') or row.get('mail_subject') or '')[:80]
                if titre:
                    lines.append(f"- 🟢 Impact positif : {titre}")

            # Détailler les impacts négatifs
            for row in neg[:2]:
                titre = str(row.get('titre') or row.get('mail_subject') or '')[:80]
                if titre:
                    lines.append(f"- 🔴 Impact négatif : {titre}")

            # Sociétés potentiellement concernées selon la zone
            zone_societes = {
                'uemoa':        "SGBC, BICB, PALC, SOGC, ETIT, ONTBF, CFAC, SLBC",
                'brvm':         "Toutes sociétés cotées BRVM",
                'afrique_ouest':"BOAB, BOABF, BOAC, BOAM, BOAN, BOAS, SMBC, STBC",
                'afrique':      "ORAC, SNTS, ORGT, SGBC, BICB",
                'international':"PALC (cacao/huile de palme), SOGC (caoutchouc), ETIT, NTLC",
            }
            societes = zone_societes.get(zone, "Sociétés cotées BRVM")
            lines.append(f"- Sociétés potentiellement impactées : {societes}")

            return "\n".join(lines) if lines else "- Impact non déterminable sur la base des données disponibles."

        # Nouvelle structure : 3 types × 5 zones
        TYPES_MAP = {
            'macroeconomique': '## TYPE_1. ACTUALITÉS MACRO-ÉCONOMIQUES',
            'politique':       '## TYPE_2. ACTUALITÉS POLITIQUES & GÉOPOLITIQUES',
            'financiere':      '## TYPE_3. ACTUALITÉS FINANCIÈRES & MARCHÉS',
        }
        ZONES_MAP = {
            'international': ('### PLAN MONDIAL',          '🌐'),
            'afrique':       ('### PLAN AFRICAIN',         '🌍'),
            'afrique_ouest': ("### PLAN AFRIQUE DE L'OUEST", '0001F30D'),
            'uemoa':         ('### PLAN UEMOA',            '🏦'),
            'brvm':          ('### PLAN BRVM',             '📈'),
        }

        for t_key, t_label in TYPES_MAP.items():
            t_data = macro_news.get(t_key, {})
            if not isinstance(t_data, dict):
                continue
            has_any = any(
                isinstance(df, pd.DataFrame) and not df.empty
                for df in t_data.values()
            )
            if not has_any:
                continue
            sections.append(t_label)
            for z_key, (z_header, z_emoji) in ZONES_MAP.items():
                df = t_data.get(z_key, pd.DataFrame())
                if not isinstance(df, pd.DataFrame) or df.empty:
                    continue
                items = []
                for _, row in df.head(5).iterrows():
                    date_s = str(row.get('mail_date',''))[:10]
                    titre  = str(row.get('titre') or row.get('mail_subject') or '')[:120]
                    resume = str(row.get('resume',''))[:300]
                    sent   = str(row.get('sentiment','') or '').lower()
                    impact = str(row.get('impact_brvm','') or '').lower()
                    s_lbl  = '🟢 Positif' if 'positif' in sent else ('🔴 Négatif' if 'negatif' in sent else '⚪ Neutre')
                    i_lbl  = '🟢 Positif' if 'positif' in impact else ('🔴 Négatif' if 'negatif' in impact else '⚪ Neutre')
                    items.append(
                        f"- [{date_s}] {titre}\n"
                        f"  {resume}\n"
                        f"  Sentiment: {s_lbl} | Impact BRVM estimé: {i_lbl}"
                    )
                sections.append(f"{z_header}\n" + "\n\n".join(items))

                # ── Impact automatique basé sur sentiments + mots-clés ────
                impact_lines = _auto_impact(df, z_key)
                sections.append("#### Impact estimé sur la BRVM\n" + impact_lines)

        if not sections:
            sections.append(
                "## TYPE_1. ACTUALITÉS MACRO-ÉCONOMIQUES\n"
                "### PLAN MONDIAL\nAucune actualité disponible pour cette période."
            )

        # ── Synthèse automatique globale ─────────────────────────────────
        all_dfs = [
            macro_news.get(t, {}).get(z, pd.DataFrame())
            for t in TYPES_MAP
            for z in ZONES_MAP
            if isinstance(macro_news.get(t, {}), dict)
        ]
        all_dfs = [d for d in all_dfs if isinstance(d, pd.DataFrame) and not d.empty]
        total_arts = sum(len(d) for d in all_dfs)

        pos_count = sum(
            1 for d in all_dfs
            for _, r in d.iterrows()
            if 'positif' in str(r.get('sentiment','')).lower()
               or 'positif' in str(r.get('impact_brvm','')).lower()
        )
        neg_count = sum(
            1 for d in all_dfs
            for _, r in d.iterrows()
            if 'negatif' in str(r.get('sentiment','')).lower()
               or 'negatif' in str(r.get('impact_brvm','')).lower()
        )
        if total_arts > 0:
            pct_pos = pos_count / total_arts * 100
            pct_neg = neg_count / total_arts * 100
        else:
            pct_pos = pct_neg = 0

        if pct_pos > 50:
            alerte = "VERT"
            alerte_txt = (
                f"L'environnement macro est globalement favorable à la BRVM. "
                f"Sur {total_arts} actualités analysées, {pos_count} présentent "
                f"un sentiment positif ({pct_pos:.0f}%). Les flux d'investissement "
                f"pourraient soutenir la tendance haussière des indices."
            )
        elif pct_neg > 40:
            alerte = "ROUGE"
            alerte_txt = (
                f"L'environnement macro présente des risques significatifs. "
                f"Sur {total_arts} actualités analysées, {neg_count} signalent "
                f"des facteurs défavorables ({pct_neg:.0f}%). Prudence recommandée "
                f"sur les positions longues à court terme."
            )
        else:
            alerte = "ORANGE"
            alerte_txt = (
                f"L'environnement macro est mitigé. Sur {total_arts} actualités analysées, "
                f"{pos_count} sont positives et {neg_count} négatives. "
                f"Une vigilance accrue est conseillée, particulièrement sur les valeurs "
                f"exposées aux matières premières et aux changes."
            )

        sections.append(
            f"## SYNTHESE_FINALE\n"
            f"Niveau d'alerte global pour un investisseur BRVM : {alerte}\n"
            f"- {alerte_txt}\n"
            f"- Analyse automatique basée sur {total_arts} actualités collectées "
            f"(sentiment positif: {pos_count}, négatif: {neg_count}, neutre: {total_arts-pos_count-neg_count}).\n"
            f"- Consulter les tableaux de sources en fin de section pour le détail."
        )

        return "\n\n".join(sections)

    def _get_brvm_actualites(self):
        """
        Charge les documents officiels récents (AG, dividendes, convocations,
        résultats) depuis brvm_documents — toutes sociétés confondues.
        Retourne un DataFrame trié par date décroissante.
        """
        logging.info("📄 Chargement actualités brvm_documents (AG/dividendes/convocations)...")
        query = """
        SELECT
            societe_confirmee,
            titre,
            date_doc,
            date_publication,
            categorie,
            type_document,
            rapport_type,
            resume,
            points_cles,
            impact,
            doc_url
        FROM brvm_documents
        WHERE resume IS NOT NULL AND resume <> ''
        ORDER BY date_doc DESC NULLS LAST, date_publication DESC NULLS LAST
        LIMIT 30;
        """
        try:
            df = pd.read_sql(query, self.db_conn)
            logging.info(f"   ✅ {len(df)} document(s) brvm_documents chargés (vue globale)")
            return df
        except Exception as e:
            logging.warning(f"⚠️  brvm_documents actualités: {e}")
            return pd.DataFrame()

    # =========================================================================
    # TABLE DE CORRESPONDANCE : nom long BDD → symbole BRVM officiel
    # Basée sur les valeurs réelles de societe_confirmee / societe dans Supabase
    # Cross-referencée avec fundamental_analyzer.py (source de vérité symboles)
    # =========================================================================
    _SOCIETE_TO_SYMBOL = {
        # brvm_rapports_societes.societe → symbole BRVM
        "AGL":             "SDSC",
        "AIR LIQUIDE CI":  "SIVC",
        "BERNABE CI":      "BNBC",
        "BICI CI":         "BICC",
        "BIIC":            "BICB",
        "BOA BF":          "BOABF",
        "BOA BN":          "BOAB",
        "BOA CI":          "BOAC",
        "BOA ML":          "BOAM",
        "BOA NG":          "BOAN",
        "BOA SN":          "BOAS",
        "CFAO MOTORS CI":  "CFAC",
        "CIE CI":          "CIEC",
        "CORIS BANK":      "CBIBF",  # Coris Bank International Burkina Faso
        "ECOBANK CI":      "ECOC",
        "ECOBANK TG":      "ETIT",
        "FILTISAC CI":     "FTSC",
        "LNB":             "LNBB",
        "NEI CEDA CI":     "NEIC",
        "NESTLE CI":       "NTLC",
        "NSBC":            "NSBC",
        "ONATEL BF":       "ONTBF",
        "ORAGROUP":        "ORGT",
        "ORANGE CI":       "ORAC",
        "PALM CI":         "PALC",
        "SAFCA CI":        "SAFC",
        "SAPH CI":         "SPHC",
        "SERVAIR CI":      "ABJC",
        "SETAO CI":        "STAC",
        "SIB":             "SIBC",
        "SICABLE":         "CABC",
        "SITAB":           "STBC",
        "SMB":             "SMBC",
        "SODECI":          "SDCC",
        "SOGB":            "SOGC",
        "SOLIBRA":         "SLBC",
        "SONATEL":         "SNTS",
        "SUCRIVOIRE":      "SCRC",
        "TOTAL CI":        "TTLC",
        "TOTAL SENEGAL":   "TTLS",
        "TRACTAFRIC CI":   "PRSC",
        "UNIWAX CI":       "UNXC",
        "VIVO ENERGY CI":  "SHEC",
        # Aliases supplémentaires (variantes fréquentes et accents)
        "TOTAL CÔTE D\'IVOIRE":   "TTLC",
        "TOTAL COTE D\'IVOIRE":   "TTLC",
        "ECOBANK":                  "ECOC",
        "NESTLE":                   "NTLC",
        "NESTLÉ CI":                "NTLC",
        "NESTLÉ":                   "NTLC",
        "ORANGE":                   "ORAC",
        "SONATEL SENEGAL":          "SNTS",
        "AIR LIQUIDE":              "SIVC",
        "PALM":                     "PALC",
        "SAPH":                     "SPHC",
        "SODECI":                   "SDCC",
        "SUCRIVOIRE":               "SCRC",
        "SOLIBRA":                  "SLBC",
        "SERVAIR":                  "ABJC",
        "BERNABÉ CI":               "BNBC",
        "TRACTAFRIC":               "PRSC",
        "VIVO ENERGY":              "SHEC",
        "ECOBANK CÔTE D\'IVOIRE":  "ECOC",
        "ECOBANK COTE D\'IVOIRE":  "ECOC",
        "ECOBANK TOGO":             "ETIT",
        "ONATEL":                   "ONTBF",
        "ORANGE CÔTE D\'IVOIRE":   "ORAC",
        "BOA BURKINA":              "BOABF",
        "BOA SENEGAL":              "BOAS",
        "BOA BÉNIN":                "BOAB",
        "BOA NIGER":                "BOAN",
        "BOA MALI":                 "BOAM",
        "CORIS BANK INTERNATIONAL":   "CBIBF",
        "CORIS BANK BURKINA":         "CBIBF",
        "CORIS BANK BF":              "CBIBF",
        "FILTISAC":                   "FTSC",
        "FILTISAC COTE D\'IVOIRE":   "FTSC",
        "FILTISAC CI":                "FTSC",
        "SICOR":                    "SICC",
        "SGB CI":                   "SGBC",
        "SGBCI":                    "SGBC",
    }

    def _normalize_societe_name(self, raw_name: str) -> str:
        """
        Convertit un nom long de société (tel que stocké en BDD) en symbole BRVM.

        Stratégie en 4 passes :
          1. Correspondance exacte dans _SOCIETE_TO_SYMBOL
          2. Correspondance insensible à la casse
          3. Le nom est déjà un symbole BRVM valide (≤ 6 chars, uppercase)
          4. Recherche partielle sur les premières lettres significatives
        Retourne le symbole normalisé ou le nom original en majuscules si aucune
        correspondance n'est trouvée (avec warning dans les logs).
        """
        if not raw_name:
            return ""
        name = str(raw_name).strip()

        # Passe 1 : correspondance exacte
        if name in self._SOCIETE_TO_SYMBOL:
            return self._SOCIETE_TO_SYMBOL[name]

        # Passe 2 : insensible à la casse
        name_upper = name.upper()
        for key, sym in self._SOCIETE_TO_SYMBOL.items():
            if key.upper() == name_upper:
                return sym

        # Passe 3 : déjà un symbole BRVM (court, tout en majuscules)
        if len(name_upper) <= 6 and name_upper.isalpha():
            return name_upper

        # Passe 4 : correspondance partielle (le nom contient la clé ou vice-versa)
        for key, sym in self._SOCIETE_TO_SYMBOL.items():
            if key.upper() in name_upper or name_upper in key.upper():
                return sym

        logging.debug(f"_normalize_societe_name: aucun symbole trouvé pour '{name}'")
        return name_upper

    def _get_brvm_documents(self):
        """
        Charge tous les documents de brvm_documents groupés par société confirmée.
        Retourne un dict: { symbol -> [doc, doc, ...] }
        Chaque doc = { titre, date_doc, resume, points_cles, impact, categorie, doc_url }
        """
        logging.info("📂 Chargement brvm_documents...")

        query = """
        SELECT
            societe_confirmee,
            titre,
            date_doc,
            resume,
            points_cles,
            impact,
            categorie,
            doc_url
        FROM brvm_documents
        WHERE societe_confirmee IS NOT NULL
          AND societe_confirmee <> ''
          AND resume IS NOT NULL
          AND resume <> ''
        ORDER BY societe_confirmee, date_doc DESC NULLS LAST;
        """

        try:
            df = pd.read_sql(query, self.db_conn)
            logging.info(f"   ✅ {len(df)} document(s) chargé(s) depuis brvm_documents")
        except Exception as e:
            logging.error(f"❌ Erreur brvm_documents: {e}")
            return {}

        docs_by_symbol = {}
        for _, row in df.iterrows():
            raw  = str(row['societe_confirmee'] or '').strip()
            if not raw:
                continue
            sym  = self._normalize_societe_name(raw)
            if not sym:
                continue
            doc = {
                'titre':      row.get('titre') or 'Sans titre',
                'date_doc':   str(row.get('date_doc') or 'Date inconnue'),
                'resume':     row.get('resume') or '',
                'points_cles': row.get('points_cles'),   # jsonb → peut être list ou None
                'impact':     row.get('impact') or 'neutre',
                'categorie':  row.get('categorie') or '',
                'doc_url':    row.get('doc_url') or '',
            }
            docs_by_symbol.setdefault(sym, []).append(doc)

        logging.info(f"   📊 {len(docs_by_symbol)} société(s) avec documents brvm_documents")
        return docs_by_symbol


    # =========================================================================
    # brvm_rapports_societes — annonces & communiqués officiels par société
    # =========================================================================

    def _get_brvm_rapports_societes(self):
        """
        Charge tous les rapports/communiqués depuis brvm_rapports_societes,
        groupés par société.
        Retourne un dict: { symbol_upper -> [rapport, ...] }
        Chaque rapport = {
            annee, type_rapport, doc_titre, doc_url,
            resume, points_cles, indicateurs,
            recommandation, risques, perspectives,
            date_rapport, created_at
        }
        """
        logging.info("📋 Chargement brvm_rapports_societes...")

        query = """
        SELECT
            societe,
            annee,
            type_rapport,
            doc_titre,
            doc_url,
            resume,
            points_cles,
            indicateurs,
            recommandation,
            risques,
            perspectives,
            date_rapport,
            created_at
        FROM brvm_rapports_societes
        WHERE societe IS NOT NULL
          AND societe <> ''
          AND resume IS NOT NULL
          AND resume <> ''
        ORDER BY societe, annee DESC NULLS LAST, created_at DESC NULLS LAST;
        """

        try:
            df = pd.read_sql(query, self.db_conn)
            logging.info(f"   ✅ {len(df)} entrée(s) dans brvm_rapports_societes")
        except Exception as e:
            logging.error(f"❌ Erreur brvm_rapports_societes: {e}")
            return {}

        rapports_by_symbol = {}
        for _, row in df.iterrows():
            raw  = str(row['societe'] or '').strip()
            if not raw:
                continue
            sym  = self._normalize_societe_name(raw)
            if not sym:
                continue
            rapport = {
                'annee':          row.get('annee'),
                'type_rapport':   row.get('type_rapport') or '',
                'doc_titre':      row.get('doc_titre') or 'Sans titre',
                'doc_url':        row.get('doc_url') or '',
                'resume':         row.get('resume') or '',
                'points_cles':    row.get('points_cles'),    # jsonb
                'indicateurs':    row.get('indicateurs'),    # jsonb
                'recommandation': row.get('recommandation') or '',
                'risques':        row.get('risques') or '',
                'perspectives':   row.get('perspectives') or '',
                'date_rapport':   str(row.get('date_rapport') or ''),
                'created_at':     str(row.get('created_at') or ''),
            }
            rapports_by_symbol.setdefault(sym, []).append(rapport)

        logging.info(
            f"   📊 {len(rapports_by_symbol)} société(s) avec rapports brvm_rapports_societes"
        )
        return rapports_by_symbol

    def _format_rapports_societes_for_ai(self, rapports):
        """
        Formate les entrées brvm_rapports_societes pour le prompt IA.
        Max 5 entrées, résumé tronqué à 500 chars.
        """
        if not rapports:
            return ""

        parts = []
        for i, r in enumerate(rapports[:5]):
            # Points clés
            pk = r.get('points_cles')
            if isinstance(pk, list):
                pk_text = " | ".join(str(p) for p in pk[:4])
            elif isinstance(pk, str) and pk.strip():
                pk_text = pk[:250]
            else:
                pk_text = ""

            # Indicateurs financiers (jsonb)
            ind = r.get('indicateurs')
            ind_text = ""
            if isinstance(ind, dict) and ind:
                pairs = [f"{k}: {v}" for k, v in list(ind.items())[:6]]
                ind_text = " | ".join(pairs)
            elif isinstance(ind, str) and ind.strip():
                ind_text = ind[:200]

            label = f"[RAPPORT {i+1}] {r['doc_titre']}"
            if r.get('annee'):
                label += f" ({r['annee']})"
            if r.get('type_rapport'):
                label += f" — {r['type_rapport']}"

            block = label + "\n"
            block += f"Résumé: {r['resume'][:500]}\n"
            if pk_text:
                block += f"Points clés: {pk_text}\n"
            if ind_text:
                block += f"Indicateurs: {ind_text}\n"
            if r.get('recommandation'):
                block += f"Recommandation source: {r['recommandation']}\n"
            if r.get('risques'):
                block += f"Risques identifiés: {r['risques'][:200]}\n"
            if r.get('perspectives'):
                block += f"Perspectives: {r['perspectives'][:200]}\n"
            parts.append(block)

        return "\n\n".join(parts)

    def _format_rapports_societes_for_word(self, rapport, idx):
        """Prépare un rapport brvm_rapports_societes pour l'insertion Word."""
        pk = rapport.get('points_cles')
        if isinstance(pk, list):
            points = [str(p) for p in pk[:6]]
        elif isinstance(pk, str) and pk.strip():
            try:
                import json as _j
                parsed = _j.loads(pk)
                points = [str(p) for p in parsed[:6]] if isinstance(parsed, list) else [pk[:200]]
            except Exception:
                points = [pk[:200]]
        else:
            points = []

        ind = rapport.get('indicateurs')
        indicateurs_kv = {}
        if isinstance(ind, dict):
            indicateurs_kv = {k: v for k, v in list(ind.items())[:8]}
        elif isinstance(ind, str) and ind.strip():
            try:
                import json as _j
                parsed = _j.loads(ind)
                if isinstance(parsed, dict):
                    indicateurs_kv = {k: v for k, v in list(parsed.items())[:8]}
            except Exception:
                pass

        rec = str(rapport.get('recommandation', '')).upper()
        if 'ACHAT' in rec:
            rec_color = RGBColor(0, 128, 0)
        elif 'VENTE' in rec:
            rec_color = RGBColor(192, 0, 0)
        else:
            rec_color = RGBColor(100, 100, 100)

        return {
            'idx':            idx,
            'annee':          rapport.get('annee') or '—',
            'type_rapport':   rapport.get('type_rapport') or '',
            'doc_titre':      rapport.get('doc_titre') or 'Sans titre',
            'doc_url':        rapport.get('doc_url') or '',
            'resume':         rapport.get('resume') or '',
            'points':         points,
            'indicateurs':    indicateurs_kv,
            'recommandation': rapport.get('recommandation') or '',
            'rec_color':      rec_color,
            'risques':        rapport.get('risques') or '',
            'perspectives':   rapport.get('perspectives') or '',
            'date_rapport':   rapport.get('date_rapport') or '',
        }

    def _format_brvm_documents_for_ai(self, docs):
        """
        Formate les documents brvm_documents en texte structuré pour l'IA.
        Limite à 5 documents max pour ne pas dépasser le contexte.
        """
        if not docs:
            return ""

        parts = []
        for i, doc in enumerate(docs[:5]):
            # Points clés : peut être une list Python (jsonb décodé par psycopg2)
            pk = doc.get('points_cles')
            if isinstance(pk, list):
                pk_text = " | ".join(str(p) for p in pk[:5]) if pk else ""
            elif isinstance(pk, str):
                pk_text = pk[:300]
            else:
                pk_text = ""

            impact_icon = {"positif": "🟢", "negatif": "🔴", "neutre": "⚪"}.get(
                str(doc.get('impact', 'neutre')).lower(), "⚪"
            )

            parts.append(
                f"[DOC {i+1}] {doc['titre']} | {doc['date_doc']} | "
                f"Catégorie: {doc['categorie']} | Impact: {impact_icon} {doc['impact']}\n"
                f"Résumé: {doc['resume'][:600]}\n"
                + (f"Points clés: {pk_text}\n" if pk_text else "")
            )

        return "\n\n".join(parts)

    def _format_brvm_documents_for_word(self, doc, doc_index):
        """
        Retourne un dict structuré prêt pour l'insertion Word.
        """
        pk = doc.get('points_cles')
        if isinstance(pk, list):
            points = [str(p) for p in pk[:6]]
        elif isinstance(pk, str) and pk.strip():
            # Essai de parse JSON si c'est une chaîne
            try:
                import json as _json
                parsed = _json.loads(pk)
                points = [str(p) for p in parsed[:6]] if isinstance(parsed, list) else [pk[:200]]
            except Exception:
                points = [pk[:200]]
        else:
            points = []

        impact = str(doc.get('impact', 'neutre')).lower()
        impact_icon = {"positif": "🟢 POSITIF", "negatif": "🔴 NÉGATIF", "neutre": "⚪ NEUTRE"}.get(impact, "⚪ NEUTRE")
        impact_color = {"positif": RGBColor(0, 128, 0), "negatif": RGBColor(192, 0, 0), "neutre": RGBColor(100, 100, 100)}.get(impact, RGBColor(100, 100, 100))

        # Badge fraîcheur pour brvm_documents
        date_doc_raw = str(doc.get('date_doc','') or '')
        try:
            from datetime import date as _d2
            import re as _re2
            m = _re2.search(r'(\d{4})[-/](\d{1,2})', date_doc_raw)
            if m:
                dy, dm = int(m.group(1)), int(m.group(2))
                td2    = _d2.today()
                mago2  = (td2.year - dy)*12 + (td2.month - dm)
                fresh2 = "⭐ RÉCENT" if mago2 <= 3 else ("📅 À JOUR" if mago2 <= 12 else "⚠️ ANCIEN")
                date_doc_display = f"{date_doc_raw} {fresh2}"
            else:
                date_doc_display = date_doc_raw
        except Exception:
            date_doc_display = date_doc_raw

        return {
            'num':          doc_index,
            'titre':        doc.get('titre', 'Sans titre'),
            'date_doc':     date_doc_display,
            'categorie':    doc.get('categorie', ''),
            'resume':       doc.get('resume', ''),
            'points':       points,
            'impact':       impact_icon,
            'impact_color': impact_color,
            'doc_url':      doc.get('doc_url', ''),
        }

    def _extract_recommendation_from_analysis(self, analysis_text,
                                                tech_decision=None,
                                                fund_decision=None):
        """
        Extrait et CORRIGE la recommandation du texte d'analyse IA.

        Règle de cohérence :
          • Si tech=VENTE et fond=ACHAT (ou inverse) → ajuste à NEUTRE/CONSERVER
            avec note de divergence, sauf si l'IA a explicitement justifié
            le choix final avec les mots 'malgré' ou 'divergence'.
          • Si tech et fond convergent → confirme la recommandation IA.
          • Sinon → recommandation IA brute.
        """
        al = analysis_text.lower()

        # 1. Lire la recommandation brute de l'IA
        if 'achat fort' in al:
            raw_rec, raw_score = 'ACHAT FORT', 5
        elif 'achat' in al:
            raw_rec, raw_score = 'ACHAT', 4
        elif 'conserver' in al:
            raw_rec, raw_score = 'CONSERVER', 3
        elif 'vente forte' in al:
            raw_rec, raw_score = 'VENTE FORTE', 1
        elif 'vente' in al:
            raw_rec, raw_score = 'VENTE', 2
        else:
            raw_rec, raw_score = 'CONSERVER', 3

        # 2. Correction de cohérence tech / fondamental
        if tech_decision and fund_decision:
            td = str(tech_decision).upper()
            fd = str(fund_decision).upper()
            divergence_forte = (
                ('ACHAT' in td and 'VENTE' in fd) or
                ('VENTE' in td and 'ACHAT' in fd)
            )
            # L'IA a-t-elle elle-même reconnu la divergence ?
            ia_acknowledged = any(kw in al for kw in
                                  ['malgré', 'divergence', 'contradiction',
                                   'incohérence', 'signal mixte'])

            if divergence_forte and not ia_acknowledged:
                # Recommandation ajustée
                if raw_rec in ('ACHAT FORT', 'ACHAT'):
                    return 'CONSERVER ⚠️', 3   # score 3 = neutre
                elif raw_rec in ('VENTE FORTE', 'VENTE'):
                    return 'CONSERVER ⚠️', 3
                # else: déjà CONSERVER, on laisse

        return raw_rec, raw_score

    # ============================================================================
    # NOUVELLES FONCTIONS MULTI-AI (DeepSeek → Gemini → Mistral)
    # ============================================================================
    
    def _generate_analysis_with_deepseek(self, symbol, data_dict, prompt):
        """Génération d'analyse avec DeepSeek"""
        if not DEEPSEEK_API_KEY:
            return None, None
        
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": DEEPSEEK_MODEL,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 4000,
            "temperature": 0.4
        }
        
        try:
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    text = result['choices'][0]['message']['content']
                    self.request_count['deepseek'] += 1
                    self.request_count['total'] += 1
                    return text, "deepseek"
            
            return None, None
            
        except Exception as e:
            logging.error(f"❌ DeepSeek exception: {e}")
            return None, None

    def _generate_analysis_with_gemini(self, symbol, data_dict, prompt):
        """Génération d'analyse avec Gemini"""
        if not GEMINI_API_KEY:
            return None, None
        
        url = f"{GEMINI_API_URL}?key={GEMINI_API_KEY}"
        
        data = {
            "contents": [{
                "parts": [{"text": prompt}]
            }],
            "generationConfig": {
                "temperature": 0.4,
                "maxOutputTokens": 3000
            }
        }
        
        try:
            response = requests.post(url, json=data, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and len(result['candidates']) > 0:
                    text = result['candidates'][0]['content']['parts'][0]['text']
                    self.request_count['gemini'] += 1
                    self.request_count['total'] += 1
                    return text, "gemini"
            
            return None, None
            
        except Exception as e:
            logging.error(f"❌ Gemini exception: {e}")
            return None, None

    def _generate_analysis_with_mistral(self, symbol, data_dict, prompt):
        """Génération d'analyse avec Mistral — avec gestion rate limit 429"""
        if not MISTRAL_API_KEY:
            return None, None

        headers = {
            "Authorization": f"Bearer {MISTRAL_API_KEY}",
            "Content-Type": "application/json"
        }

        # max_tokens selon longueur du prompt (macro = plus long)
        prompt_len   = len(prompt)
        max_tok      = 2500 if prompt_len > 2000 else 1500
        request_body = {
            "model": MISTRAL_MODEL,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tok,
            "temperature": 0.4
        }

        for _attempt in range(3):
            try:
                response = requests.post(
                    MISTRAL_API_URL, headers=headers,
                    json=request_body, timeout=60
                )

                if response.status_code == 200:
                    data = response.json()
                    if 'choices' in data and len(data['choices']) > 0:
                        text = data['choices'][0]['message']['content']
                        self.request_count['mistral'] += 1
                        self.request_count['total'] += 1
                        return text, "mistral"
                    return None, None

                elif response.status_code == 429:
                    # Rate limit — attendre selon Retry-After ou délai exponentiel
                    retry_after = int(response.headers.get('Retry-After', 10 * (2 ** _attempt)))
                    logging.warning(
                        f"    ⏳ Mistral rate limit (429) pour {symbol} — "
                        f"attente {retry_after}s (tentative {_attempt+1}/3)"
                    )
                    time.sleep(retry_after)
                    continue

                elif response.status_code in (401, 403):
                    logging.error(
                        f"    🔑 Mistral clé API invalide ({response.status_code}) "
                        f"— vérifier MISTRAL_API_KEY dans GitHub Secrets"
                    )
                    return None, None

                else:
                    logging.warning(
                        f"    ⚠️  Mistral HTTP {response.status_code} pour {symbol}: "
                        f"{response.text[:150]}"
                    )
                    return None, None

            except Exception as e:
                logging.error(f"❌ Mistral exception {symbol}: {e}")
                if _attempt < 2:
                    time.sleep(5)
                    continue
                return None, None

        return None, None

    def _generate_analysis_with_claude(self, symbol, data_dict, prompt):
        """
        Génération d'analyse avec Claude (Anthropic) — fallback final.
        Utilise claude-haiku-4-5 (rapide, économique) via l'API Anthropic v1/messages.
        Variable d'environnement requise : ANTHROPIC_API_KEY dans GitHub Secrets.
        """
        if not ANTHROPIC_API_KEY:
            return None, None

        headers = {
            "x-api-key":         ANTHROPIC_API_KEY,
            "anthropic-version": ANTHROPIC_API_VERSION,
            "content-type":      "application/json",
        }

        request_body = {
            "model":      ANTHROPIC_MODEL,
            "max_tokens": 1500,
            "messages":   [{"role": "user", "content": prompt}],
        }

        for _attempt in range(3):
            try:
                response = requests.post(
                    ANTHROPIC_API_URL,
                    headers=headers,
                    json=request_body,
                    timeout=60,
                )

                if response.status_code == 200:
                    data = response.json()
                    # Réponse Anthropic : {"content": [{"type": "text", "text": "..."}]}
                    content = data.get("content", [])
                    text = " ".join(
                        block.get("text", "")
                        for block in content
                        if block.get("type") == "text"
                    ).strip()
                    if text:
                        self.request_count["claude"] += 1
                        self.request_count["total"]  += 1
                        logging.info(f"    ✅ {symbol}: Analyse générée via CLAUDE")
                        return text, "claude"
                    return None, None

                elif response.status_code == 429:
                    retry_after = int(response.headers.get("Retry-After", 15 * (2 ** _attempt)))
                    logging.warning(
                        f"    ⏳ Claude rate limit (429) pour {symbol} — "
                        f"attente {retry_after}s (tentative {_attempt+1}/3)"
                    )
                    time.sleep(retry_after)
                    continue

                elif response.status_code in (401, 403):
                    logging.error(
                        f"    🔑 Claude clé API invalide ({response.status_code}) "
                        f"— vérifier ANTHROPIC_API_KEY dans GitHub Secrets"
                    )
                    return None, None

                else:
                    logging.warning(
                        f"    ⚠️  Claude HTTP {response.status_code} pour {symbol}: "
                        f"{response.text[:150]}"
                    )
                    return None, None

            except Exception as e:
                logging.error(f"❌ Claude exception {symbol}: {e}")
                if _attempt < 2:
                    time.sleep(5)
                    continue
                return None, None

        return None, None


    def _get_donnees_financieres(self, symbol):
        """
        Charge les données structurées de brvm_donnees_financieres pour un symbole.
        Retourne la ligne la plus récente (annee max) ou None si absente.
        """
        try:
            query = f"""
                SELECT *
                FROM public.brvm_donnees_financieres
                WHERE symbol = '{symbol}'
                ORDER BY annee DESC
                LIMIT 1;
            """
            df = pd.read_sql(query, self.db_conn)
            if df.empty:
                return None
            return df.iloc[0].to_dict()
        except Exception as e:
            logging.warning(f"   ⚠️ brvm_donnees_financieres non chargé pour {symbol}: {e}")
            return None

    def _format_donnees_financieres(self, fin, symbol):
        """
        Formate TOUTES les données financières structurées pour le prompt IA.
        - Affiche toutes les variables disponibles (non nulles, non zéro)
        - Détecte le secteur (bancaire vs non-bancaire)
        - Ajoute des annotations contextuelles pour l'IA
        """
        if not fin:
            return ""

        def v(val, pct=False, milliards=True):
            """Formate une valeur numérique. Retourne None si 0 ou null."""
            if val is None or val == 0 or str(val).strip() == "":
                return None
            try:
                f = float(val)
                if f == 0.0:
                    return None
                if pct:
                    return f"{f*100:.4f}%"
                if milliards and abs(f) >= 1_000_000_000:
                    return f"{f/1_000_000_000:.3f} Mds FCFA"
                if milliards and abs(f) >= 1_000_000:
                    return f"{f/1_000_000:.2f} M FCFA"
                return f"{f:,.2f} FCFA"
            except (TypeError, ValueError):
                return None

        annee = fin.get('annee', 'N/A')

        # ── Détection secteur bancaire ──────────────────────────────────────────
        is_bank = any([
            fin.get('caisse_banque_centrale') and float(fin.get('caisse_banque_centrale') or 0) != 0,
            fin.get('produit_net_bancaire')   and float(fin.get('produit_net_bancaire')   or 0) != 0,
            fin.get('dettes_clientele')       and float(fin.get('dettes_clientele')       or 0) != 0,
            fin.get('creances_interbancaires')and float(fin.get('creances_interbancaires')or 0) != 0,
        ])
        secteur_label = "BANQUE" if is_bank else "ENTREPRISE NON BANCAIRE"

        lines = []
        lines.append(f"\n{'═'*70}")
        lines.append(f"📊 DONNÉES FINANCIÈRES STRUCTURÉES COMPLÈTES — {symbol} (Exercice {annee}) [{secteur_label}]")
        lines.append(f"{'═'*70}")
        lines.append("⚠️  Règle : toute valeur absente ou = 0 signifie donnée manquante ou non pertinente pour ce secteur.\n")

        # ══════════════════════════════════════════════════════════════════════
        # SECTION 1 : BILAN — ACTIF
        # ══════════════════════════════════════════════════════════════════════
        ba = []
        # Variables spécifiques banques
        if v(fin.get('caisse_banque_centrale')):
            ba.append(f"  • Caisse & Banque Centrale              : {v(fin.get('caisse_banque_centrale'))}"
                      + (" [🏦 BANQUE: liquidité primaire, réserves obligatoires]" if is_bank else ""))
        if v(fin.get('effets_publics')):
            ba.append(f"  • Effets publics & valeurs assimilées   : {v(fin.get('effets_publics'))}"
                      + (" [🏦 BANQUE: portefeuille de titres d'État]" if is_bank else ""))
        if v(fin.get('creances_interbancaires')):
            ba.append(f"  • Créances interbancaires               : {v(fin.get('creances_interbancaires'))}"
                      + (" [🏦 BANQUE: prêts aux autres établissements bancaires]" if is_bank else ""))
        if v(fin.get('creances_clientele')):
            ba.append(f"  • Créances sur la clientèle             : {v(fin.get('creances_clientele'))}"
                      + (" [🏦 BANQUE: portefeuille total de crédits accordés aux clients]" if is_bank else ""))
        # Variables spécifiques entreprises
        if v(fin.get('creances_clients')):
            ba.append(f"  • Créances clients                      : {v(fin.get('creances_clients'))}"
                      + (" [🏢 ENTREPRISE: montants dus par les clients]" if not is_bank else ""))
        if v(fin.get('stocks')):
            ba.append(f"  • Stocks                                : {v(fin.get('stocks'))}"
                      + (" [🏢 ENTREPRISE: marchandises, matières premières, produits finis]" if not is_bank else ""))
        if v(fin.get('actif_circulant')):
            ba.append(f"  • Actif circulant                       : {v(fin.get('actif_circulant'))}")
        # Variables communes
        if v(fin.get('immobilisations_incorporelles')):
            ba.append(f"  • Immobilisations incorporelles         : {v(fin.get('immobilisations_incorporelles'))}")
        if v(fin.get('immobilisations_corporelles')):
            ba.append(f"  • Immobilisations corporelles           : {v(fin.get('immobilisations_corporelles'))}")
        if v(fin.get('actif_immobilise_net')):
            ba.append(f"  • Actif immobilisé net                  : {v(fin.get('actif_immobilise_net'))}")
        if v(fin.get('tresorerie_actif')):
            ba.append(f"  • Trésorerie Actif                      : {v(fin.get('tresorerie_actif'))}")
        if v(fin.get('total_actif')):
            ba.append(f"  • Total Actif                           : {v(fin.get('total_actif'))}")
        if v(fin.get('total_bilan')):
            ba.append(f"  • Total Bilan                           : {v(fin.get('total_bilan'))}")
        if ba:
            lines.append("📌 1. BILAN — ACTIF")
            lines.extend(ba)

        # ══════════════════════════════════════════════════════════════════════
        # SECTION 2 : BILAN — PASSIF
        # ══════════════════════════════════════════════════════════════════════
        bp = []
        if v(fin.get('capital_souscrit')):
            bp.append(f"  • Capital souscrit                      : {v(fin.get('capital_souscrit'))}")
        if v(fin.get('reserves')):
            bp.append(f"  • Réserves                              : {v(fin.get('reserves'))}")
        if v(fin.get('capitaux_propres')):
            bp.append(f"  • Capitaux propres                      : {v(fin.get('capitaux_propres'))}")
        if v(fin.get('capitaux_permanents')):
            bp.append(f"  • Capitaux permanents                   : {v(fin.get('capitaux_permanents'))}")
        if v(fin.get('dettes_interbancaires')):
            bp.append(f"  • Dettes interbancaires                 : {v(fin.get('dettes_interbancaires'))}"
                      + (" [🏦 BANQUE: emprunts auprès d'autres banques]" if is_bank else ""))
        if v(fin.get('dettes_clientele')):
            bp.append(f"  • Dettes clientèle (dépôts)             : {v(fin.get('dettes_clientele'))}"
                      + (" [🏦 BANQUE: dépôts collectés auprès des clients → principale ressource]" if is_bank else ""))
        if v(fin.get('dettes_fournisseurs')):
            bp.append(f"  • Dettes fournisseurs                   : {v(fin.get('dettes_fournisseurs'))}"
                      + (" [🏢 ENTREPRISE]" if not is_bank else ""))
        if v(fin.get('dettes_financieres_lt_mt')):
            bp.append(f"  • Dettes financières LT/MT              : {v(fin.get('dettes_financieres_lt_mt'))}")
        if v(fin.get('dettes_financieres_totales')):
            bp.append(f"  • Dettes financières totales            : {v(fin.get('dettes_financieres_totales'))}")
        if v(fin.get('dettes_totales')):
            bp.append(f"  • Dettes totales                        : {v(fin.get('dettes_totales'))}")
        if v(fin.get('passif_circulant')):
            bp.append(f"  • Passif circulant                      : {v(fin.get('passif_circulant'))}")
        if bp:
            lines.append("\n📌 2. BILAN — PASSIF")
            lines.extend(bp)

        # ══════════════════════════════════════════════════════════════════════
        # SECTION 3 : COMPTE DE RÉSULTAT
        # ══════════════════════════════════════════════════════════════════════
        cr = []
        if v(fin.get('ca_pnb')):
            cr.append(f"  • CA / PNB (agrégé)                     : {v(fin.get('ca_pnb'))}")
        if v(fin.get('produit_net_bancaire')):
            cr.append(f"  • Produit Net Bancaire (PNB)             : {v(fin.get('produit_net_bancaire'))}"
                      + " [🏦 BANQUE: équivalent du chiffre d'affaires pour une banque]")
        if v(fin.get('chiffre_affaires')):
            cr.append(f"  • Chiffre d'affaires (CA)               : {v(fin.get('chiffre_affaires'))}"
                      + (" [🏢 ENTREPRISE]" if not is_bank else ""))
        if v(fin.get('interets_produits')):
            cr.append(f"  • Intérêts & produits assimilés         : {v(fin.get('interets_produits'))}"
                      + (" [🏦 BANQUE: revenus des crédits accordés]" if is_bank else ""))
        if v(fin.get('interets_charges')):
            cr.append(f"  • Intérêts & charges assimilées         : {v(fin.get('interets_charges'))}"
                      + (" [🏦 BANQUE: coût de la collecte de dépôts]" if is_bank else ""))
        if v(fin.get('commissions_produits')):
            cr.append(f"  • Commissions (produits)                : {v(fin.get('commissions_produits'))}"
                      + (" [🏦 BANQUE: frais bancaires perçus]" if is_bank else ""))
        if v(fin.get('commissions_charges')):
            cr.append(f"  • Commissions (charges)                 : {v(fin.get('commissions_charges'))}")
        if v(fin.get('charges_generales_exploitation')):
            cr.append(f"  • Charges générales d'exploitation      : {v(fin.get('charges_generales_exploitation'))}")
        if v(fin.get('dap_immobilisations')):
            cr.append(f"  • DAP immobilisations                   : {v(fin.get('dap_immobilisations'))}")
        if v(fin.get('charges_personnel')):
            cr.append(f"  • Charges du personnel                  : {v(fin.get('charges_personnel'))}")
        if v(fin.get('charges_financieres')):
            cr.append(f"  • Charges financières                   : {v(fin.get('charges_financieres'))}")
        if v(fin.get('valeur_ajoutee')):
            cr.append(f"  • Valeur Ajoutée (VA)                   : {v(fin.get('valeur_ajoutee'))}")
        if v(fin.get('ebe')):
            cr.append(f"  • Excédent Brut d'Exploitation (EBE)    : {v(fin.get('ebe'))}")
        if v(fin.get('rbe')):
            cr.append(f"  • Résultat Brut d'Exploitation (RBE)    : {v(fin.get('rbe'))}"
                      + (" [🏦 BANQUE: PNB - charges générales - DAP]" if is_bank else ""))
        if v(fin.get('resultat_exploitation')):
            cr.append(f"  • Résultat d'exploitation               : {v(fin.get('resultat_exploitation'))}")
        if v(fin.get('provisions')):
            cr.append(f"  • Provisions                            : {v(fin.get('provisions'))}")
        if v(fin.get('impot_benefices')):
            cr.append(f"  • Impôt sur les bénéfices               : {v(fin.get('impot_benefices'))}")
        if v(fin.get('resultat_avant_impot')):
            cr.append(f"  • Résultat avant impôt                  : {v(fin.get('resultat_avant_impot'))}")
        if v(fin.get('resultat_net')):
            cr.append(f"  • Résultat net                          : {v(fin.get('resultat_net'))}")
        if cr:
            lines.append("\n📌 3. COMPTE DE RÉSULTAT")
            lines.extend(cr)

        # ══════════════════════════════════════════════════════════════════════
        # SECTION 4 : CASH-FLOWS & TRÉSORERIE
        # ══════════════════════════════════════════════════════════════════════
        cf = []
        if v(fin.get('caf')):
            cf.append(f"  • CAF (Capacité d'autofinancement)      : {v(fin.get('caf'))}")
        if v(fin.get('cafg')):
            cf.append(f"  • CAFG                                  : {v(fin.get('cafg'))}")
        if v(fin.get('flux_operationnel')):
            cf.append(f"  • Flux opérationnel                     : {v(fin.get('flux_operationnel'))}")
        if v(fin.get('flux_investissement')):
            cf.append(f"  • Flux d'investissement                 : {v(fin.get('flux_investissement'))}")
        if v(fin.get('flux_financement')):
            cf.append(f"  • Flux de financement                   : {v(fin.get('flux_financement'))}")
        if v(fin.get('cashflow_operationnel')):
            cf.append(f"  • Cash-flow opérationnel                : {v(fin.get('cashflow_operationnel'))}")
        if v(fin.get('free_cash_flow')):
            cf.append(f"  • Free Cash Flow                        : {v(fin.get('free_cash_flow'))}")
        if v(fin.get('bfr')):
            cf.append(f"  • BFR (Besoin en Fonds de Roulement)    : {v(fin.get('bfr'))}"
                      + (" [🏢 ENTREPRISE: BFR négatif = favorable]" if not is_bank else ""))
        if v(fin.get('fonds_roulement')):
            cf.append(f"  • Fonds de Roulement                    : {v(fin.get('fonds_roulement'))}")
        if v(fin.get('tresorerie_nette')):
            cf.append(f"  • Trésorerie nette                      : {v(fin.get('tresorerie_nette'))}")
        if cf:
            lines.append("\n📌 4. CASH-FLOWS & TRÉSORERIE")
            lines.extend(cf)

        # ══════════════════════════════════════════════════════════════════════
        # SECTION 5 : RATIOS DE RENTABILITÉ
        # ══════════════════════════════════════════════════════════════════════
        rr = []
        if v(fin.get('marge_brute'), pct=True):
            rr.append(f"  • Marge brute                           : {v(fin.get('marge_brute'), pct=True)}"
                      + (" [🏢 ENTREPRISE: (CA - coût des ventes) / CA]" if not is_bank else ""))
        if v(fin.get('marge_nette'), pct=True):
            rr.append(f"  • Marge nette                           : {v(fin.get('marge_nette'), pct=True)}"
                      " [Résultat net / CA ou PNB]")
        if v(fin.get('marge_operationnelle'), pct=True):
            rr.append(f"  • Marge opérationnelle                  : {v(fin.get('marge_operationnelle'), pct=True)}")
        if v(fin.get('roe'), pct=True):
            rr.append(f"  • ROE (Rentabilité fonds propres)       : {v(fin.get('roe'), pct=True)}"
                      " [Résultat net / Capitaux propres — >15% = excellent]")
        if v(fin.get('roa'), pct=True):
            rr.append(f"  • ROA (Rentabilité des actifs)          : {v(fin.get('roa'), pct=True)}"
                      " [Résultat net / Total actif]")
        if v(fin.get('rotation_actifs'), pct=True):
            rr.append(f"  • Rotation des actifs                   : {v(fin.get('rotation_actifs'), pct=True)}"
                      " [CA / Total actif]")
        if v(fin.get('coefficient_exploitation'), pct=True):
            rr.append(f"  • Coefficient d'exploitation            : {v(fin.get('coefficient_exploitation'), pct=True)}"
                      + (" [🏦 BANQUE: charges / PNB — <60% = banque efficace, <50% = excellente]" if is_bank else ""))
        if v(fin.get('taux_croissance_ca'), pct=True):
            lbl = "PNB" if is_bank else "CA"
            rr.append(f"  • Taux de croissance {lbl}               : {v(fin.get('taux_croissance_ca'), pct=True)}")
        # Coût du risque — banques uniquement
        if v(fin.get('cout_risque'), pct=True):
            cr_val = float(fin.get('cout_risque') or 0)
            if cr_val != 0:
                if abs(cr_val) < 0.01:
                    interp_cr = "faible → portefeuille sain"
                elif abs(cr_val) < 0.03:
                    interp_cr = "modéré → surveillance requise"
                else:
                    interp_cr = "élevé → risque de crédit significatif"
                rr.append(f"  • Coût du risque / charges financières  : {v(fin.get('cout_risque'), pct=True)}"
                          f" [{interp_cr}]"
                          + (" [🏦 BANQUE: provisions / charges financières — PLUS c'est bas, MIEUX c'est]" if is_bank else ""))
        if rr:
            lines.append("\n📌 5. RATIOS DE RENTABILITÉ")
            lines.extend(rr)

        # ══════════════════════════════════════════════════════════════════════
        # SECTION 6 : RATIOS DE STRUCTURE FINANCIÈRE & LIQUIDITÉ
        # ══════════════════════════════════════════════════════════════════════
        rs = []
        if v(fin.get('autonomie_financiere'), pct=True):
            rs.append(f"  • Autonomie financière                  : {v(fin.get('autonomie_financiere'), pct=True)}"
                      " [Capitaux propres / Total bilan — >30% = sain]")
        if v(fin.get('dependance_financiere'), pct=True):
            rs.append(f"  • Dépendance financière                 : {v(fin.get('dependance_financiere'), pct=True)}"
                      " [Dettes financières / Total bilan]")
        if v(fin.get('ratio_endettement'), pct=True):
            rs.append(f"  • Ratio d'endettement                   : {v(fin.get('ratio_endettement'), pct=True)}"
                      " [Dettes totales / Capitaux propres]")
        if v(fin.get('solvabilite_generale'), pct=True):
            rs.append(f"  • Solvabilité générale                  : {v(fin.get('solvabilite_generale'), pct=True)}"
                      " [Total actif / Dettes totales — >1 = solvable]")
        if v(fin.get('liquidite_generale'), pct=True):
            rs.append(f"  • Liquidité générale                    : {v(fin.get('liquidite_generale'), pct=True)}"
                      " [Actif circulant / Passif circulant — >1 = bon]")
        if v(fin.get('liquidite_immediate'), pct=True):
            rs.append(f"  • Liquidité immédiate                   : {v(fin.get('liquidite_immediate'), pct=True)}"
                      " [Trésorerie / Passif circulant]")
        if v(fin.get('liquidite_reduite'), pct=True):
            rs.append(f"  • Liquidité réduite                     : {v(fin.get('liquidite_reduite'), pct=True)}"
                      " [(Actif circulant - Stocks) / Passif circulant]")
        if v(fin.get('financement_immobilisations'), pct=True):
            rs.append(f"  • Financement des immobilisations       : {v(fin.get('financement_immobilisations'), pct=True)}"
                      " [Capitaux permanents / Actif immobilisé]")
        if v(fin.get('capacite_remboursement')):
            rs.append(f"  • Capacité de remboursement             : {v(fin.get('capacite_remboursement'), milliards=False)}"
                      " [Dettes financières / CAF — en années]")
        if v(fin.get('couverture_interets'), pct=True):
            rs.append(f"  • Couverture des intérêts               : {v(fin.get('couverture_interets'), pct=True)}"
                      " [Résultat exploitation / Charges financières]")
        if v(fin.get('couverture_investissements_caf'), pct=True):
            rs.append(f"  • Couverture investissements par CAF    : {v(fin.get('couverture_investissements_caf'), pct=True)}"
                      " [CAF / Flux d'investissement]")
        if rs:
            lines.append("\n📌 6. RATIOS DE STRUCTURE FINANCIÈRE & LIQUIDITÉ")
            lines.extend(rs)

        # ══════════════════════════════════════════════════════════════════════
        # SECTION 7 : DÉLAIS D'EXPLOITATION (entreprises non-bancaires)
        # ══════════════════════════════════════════════════════════════════════
        dl = []
        if v(fin.get('delai_clients'), milliards=False):
            dl.append(f"  • Délai clients (jours)                 : {v(fin.get('delai_clients'), milliards=False)}"
                      " [🏢 Créances clients × 360 / CA — délai élevé = risque trésorerie]")
        if v(fin.get('delai_fournisseurs'), milliards=False):
            dl.append(f"  • Délai fournisseurs (jours)            : {v(fin.get('delai_fournisseurs'), milliards=False)}"
                      " [🏢 Dettes fourn. × 360 / Achats — délai élevé = bon pour la tréso]")
        if v(fin.get('duree_stockage'), milliards=False):
            dl.append(f"  • Durée de stockage (jours)             : {v(fin.get('duree_stockage'), milliards=False)}"
                      " [🏢 Stocks × 360 / CA — durée élevée = immobilisation de capital]")
        if dl:
            lines.append("\n📌 7. DÉLAIS D'EXPLOITATION")
            lines.extend(dl)

        lines.append(f"\n{'═'*70}")
        return "\n".join(lines)


    def _format_val_ratios_for_prompt(self, val_ratios):
        """Formate les ratios de valorisation pour le prompt IA."""
        if not val_ratios:
            return "Non calculables (données financières structurées manquantes)."
        lines = []
        if val_ratios.get('mkt_cap_txt'):
            lines.append(f"  • Capitalisation boursière : {val_ratios['mkt_cap_txt']}")
        if val_ratios.get('bpa_txt'):
            lines.append(f"  • BPA (Bénéfice par action): {val_ratios['bpa_txt']}")
        if val_ratios.get('per_txt'):
            lines.append(f"  • PER (Price/Earnings)     : {val_ratios['per_txt']}")
        if val_ratios.get('pb_txt'):
            lines.append(f"  • P/B (Price/Book)         : {val_ratios['pb_txt']}")
        if val_ratios.get('ev_ebitda_txt'):
            lines.append(f"  • EV/EBITDA                : {val_ratios['ev_ebitda_txt']}")
        if not lines:
            return "Ratios non calculables (données insuffisantes)."
        return "\n".join(lines)

    def _generate_professional_analysis(self, symbol, data_dict, attempt=1, max_attempts=3):
        """
        Génération analyse professionnelle avec rotation Multi-AI.
        ✅ V30: instruction fondamentale dynamique + vérification post-génération
        """
        
        if attempt > 1:
            logging.info(f"    🔄 {symbol}: Tentative {attempt}/{max_attempts}")
        
        # ── Préparer l'instruction fondamentale selon la disponibilité des données ──
        fundamental_text = data_dict.get('fundamental_analyses', '')
        has_fundamental = (
            fundamental_text
            and fundamental_text.strip()
            and "Aucun rapport" not in fundamental_text
        )

        # ── Charger et formatter les données structurées brvm_donnees_financieres ──
        fin_data    = self._get_donnees_financieres(symbol)
        fin_text    = self._format_donnees_financieres(fin_data, symbol) if fin_data else ""
        has_fin_data = bool(fin_text)
        if has_fin_data:
            annee_fin = fin_data.get('annee', 'N/A')
            logging.info(f"    💰 {symbol}: données financières structurées chargées ({annee_fin})")
        else:
            logging.info(f"    ℹ️ {symbol}: aucune donnée dans brvm_donnees_financieres")
        
        if has_fundamental:
            nb_rapports = fundamental_text.count('--- RAPPORT:')
            if nb_rapports == 0 and fundamental_text.strip():
                nb_rapports = 1
            preview = fundamental_text[:300].replace('\n', ' ') + "..."
            logging.info(f"    📊 {symbol}: {len(fundamental_text)} caractères d'analyses fondamentales ({nb_rapports} rapport(s))")
            logging.info(f"    📋 Extrait: {preview}")

            # Construire l'instruction selon la disponibilité des données structurées
            fin_instruction = ""
            if has_fin_data:
                secteur = "BANQUE" if any([
                    fin_data.get('caisse_banque_centrale') and float(fin_data.get('caisse_banque_centrale') or 0) != 0,
                    fin_data.get('produit_net_bancaire')   and float(fin_data.get('produit_net_bancaire')   or 0) != 0,
                    fin_data.get('dettes_clientele')       and float(fin_data.get('dettes_clientele')       or 0) != 0,
                ]) else "ENTREPRISE"
                fin_instruction = f"""
💰 DONNÉES STRUCTURÉES DISPONIBLES ({fin_data.get('annee', 'N/A')}) — SECTEUR: {secteur}
Ces données proviennent directement de la base brvm_donnees_financieres et sont fiables.
RÈGLES D'INTERPRÉTATION OBLIGATOIRES:
- Toute valeur à 0 ou absente = donnée manquante ou non pertinente pour ce secteur → NE PAS MENTIONNER
- Si BANQUE: interpréter le PNB comme équivalent du chiffre d'affaires, analyser le coût du risque
  (coût_risque = charges de risque / charges financières → plus c'est bas, meilleur est le portefeuille)
  Interpréter le coefficient d'exploitation (< 60% = banque efficace)
  Analyser les dépôts clientèle vs créances clientèle comme indicateur de transformation
- Si ENTREPRISE: analyser délais clients/fournisseurs, BFR, stocks
  Un délai client élevé signifie un risque de trésorerie, un BFR négatif est positif
- Pour TOUS: ROE = rentabilité des fonds propres (>15% = excellent), ROA = rentabilité des actifs
  Taux de croissance CA/PNB = dynamique commerciale
TU DOIS croiser ces données structurées avec les rapports narratifs ci-dessous."""
            
            instruction_fondamentale = f"""
⚠️ INSTRUCTION IMPÉRATIVE — ANALYSES FONDAMENTALES DISPONIBLES:
{len(fundamental_text)} caractères de données financières officielles sont fournis ci-dessous ({nb_rapports} rapport(s)).
{fin_instruction}

TU DOIS OBLIGATOIREMENT:
1. UTILISER CES DONNÉES dans la Partie 3 — c'est une instruction ABSOLUE, non optionnelle
2. Commencer par les données structurées (chiffres précis) puis enrichir avec les rapports narratifs
3. Mentionner explicitement la date de CHAQUE rapport utilisé
4. Citer les chiffres clés: chiffre d'affaires/PNB, résultat net, ratios ROE/ROA
5. Si plusieurs rapports, montrer l'évolution temporelle des indicateurs
6. NE JAMAIS mentionner les variables à 0 ou NULL — elles n'existent pas pour cette société
7. NE JAMAIS écrire que les données sont absentes si elles sont fournies ci-dessus
8. Si les rapports datent d'avant 2025, précise-le mais analyse-les quand même"""
        else:
            logging.warning(f"    ⚠️ {symbol}: Aucune analyse fondamentale en base")
            fin_instruction_fallback = ""
            if has_fin_data:
                fin_instruction_fallback = f"""
💰 DONNÉES STRUCTURÉES DISPONIBLES ({fin_data.get('annee','N/A')}):
Utilise exclusivement les données structurées fournies ci-dessus pour l'analyse fondamentale.
Ignore toute variable à 0 ou absente."""
            instruction_fondamentale = f"""
ℹ️ ABSENCE DE RAPPORTS NARRATIFS:{fin_instruction_fallback}
Aucun rapport narratif n'a été trouvé en base pour cette société.
{'Base ton analyse fondamentale uniquement sur les données structurées fournies.' if has_fin_data else 'Indique clairement cette absence dans la Partie 3 et base ta conclusion uniquement sur les indicateurs techniques et les prédictions.'}"""
        
        prompt = f"""Tu es un analyste financier professionnel spécialisé sur le marché de la BRVM (Bourse Régionale des Valeurs Mobilières, Afrique de l'Ouest). Analyse l'action {symbol} et génère un rapport structuré en 4 parties.

📊 DONNÉES DISPONIBLES:

**Évolution du cours (100 derniers jours) + Statistiques descriptives:**
{data_dict.get('historical_summary', 'Données non disponibles')}

**Ratios de valorisation boursière ({data_dict.get('val_ratios', {}).get('annee_fin', 'N/A')}):**
{self._format_val_ratios_for_prompt(data_dict.get('val_ratios', {}))}

**Indicateurs techniques:**
- Moyennes Mobiles: MM20={data_dict.get('mm_20', 'N/A')}, MM50={data_dict.get('mm_50', 'N/A')}, Décision={data_dict.get('mm_decision', 'N/A')}
- Bandes de Bollinger: Borne supérieure={data_dict.get('bollinger_upper', 'N/A')}, Borne inférieure={data_dict.get('bollinger_lower', 'N/A')}, Prix actuel={data_dict.get('price', 'N/A')}, Décision={data_dict.get('bollinger_decision', 'N/A')}
- MACD: Valeur={data_dict.get('macd_value', 'N/A')}, Signal={data_dict.get('macd_signal', 'N/A')}, Décision={data_dict.get('macd_decision', 'N/A')}
- RSI: Valeur={data_dict.get('rsi_value', 'N/A')}, Décision={data_dict.get('rsi_decision', 'N/A')}
- Stochastique: %K={data_dict.get('stochastic_k', 'N/A')}, %D={data_dict.get('stochastic_d', 'N/A')}, Décision={data_dict.get('stochastic_decision', 'N/A')}

**DONNÉES FINANCIÈRES STRUCTURÉES (brvm_donnees_financieres — chiffres officiels):**
{fin_text if has_fin_data else "Non disponibles dans la base de données structurées."}

**ANALYSES FONDAMENTALES DISPONIBLES (RAPPORTS FINANCIERS OFFICIELS):**
{fundamental_text if has_fundamental else "Aucun rapport financier enregistré en base pour cette société."}

{instruction_fondamentale}

**Prédictions IA (10 prochains jours ouvrables):**
{data_dict.get('predictions_text', 'Aucune prédiction disponible')}

═══════════════════════════════════════════════════════════════

GÉNÈRE UN RAPPORT STRUCTURÉ EN FRANÇAIS AVEC CES 4 PARTIES:

**PARTIE 0 : INDICATEURS DE VALORISATION BOURSIÈRE**

Rédige un paragraphe de 4-5 lignes commentant les ratios de valorisation fournis:
- Capitalisation boursière : son niveau et ce qu'il représente dans le contexte BRVM
- BPA (Bénéfice Par Action) : ce que chaque action rapporte en bénéfice
- PER : si les investisseurs paient cher ou pas par rapport aux bénéfices (référence: PER BRVM moyen ~8-12x)
- P/B : si le marché valorise au-dessus ou en-dessous de la valeur comptable
- EV/EBITDA : comparer la valeur totale à la capacité opérationnelle
Si certains ratios sont absents, indique pourquoi (données manquantes) sans insister.

**PARTIE 1 : ANALYSE DU COURS — STATISTIQUES ET ÉVOLUTION (100 derniers jours)**

Rédige un paragraphe de 6-8 lignes analysant:
- Variation totale sur la période ET variation J-1 (dernière séance)
- Le cours le plus haut et le plus bas atteints (range de trading)
- La tendance générale (haussière, baissière, stable) avec contexte
- **Statistiques descriptives** : commente la moyenne vs médiane (si écart → distribution asymétrique),
  l'écart-type et le CV% (dispersion du cours), le kurtosis (risque de pics) et le skewness (asymétrie)
  en utilisant les interprétations fournies dans les données
- Volatilité annualisée : positionner le titre (faible <15%, modérée 15-30%, élevée >30%)

**PARTIE 2 : ANALYSE TECHNIQUE DÉTAILLÉE**

Pour CHAQUE indicateur, rédige un paragraphe de 2-3 lignes:
- **Moyennes Mobiles**: Interprète MM20 et MM50, leur position relative au cours actuel, justifie la décision
- **Bandes de Bollinger**: Explique la position du cours par rapport aux bornes, la volatilité, justifie la décision
- **MACD**: Analyse la divergence MACD-Signal, le momentum, justifie la décision
- **RSI**: Interprète la valeur (suracheté >70, survente <30, neutre 30-70), justifie la décision
- **Stochastique**: Analyse %K et %D, leur croisement éventuel, justifie la décision

Puis rédige une **conclusion technique** de 3-4 lignes synthétisant tous les indicateurs.

**PARTIE 3 : ANALYSE FONDAMENTALE (SECTION CRITIQUE)**

Rédige un paragraphe détaillé de 8-10 lignes en suivant impérativement cette structure:
1. **Données structurées (états financiers annuels)** : si disponibles, présente les chiffres clés
   (CA/PNB, résultat net, ROE, ROA) en précisant OBLIGATOIREMENT l'année concernée (ex: "En 2025...")
   NE MENTIONNE JAMAIS les variables à 0 ou NULL
2. **Adapte au secteur détecté** :
   - BANQUE: PNB, coefficient d'exploitation (< 60% = efficace), coût du risque (faible = bon portefeuille),
     ratio dépôts/crédits, créances interbancaires
   - ENTREPRISE: CA, marges, BFR, délais clients/fournisseurs (délai client élevé = risque trésorerie),
     endettement, rotation stocks
3. **Estimations basées sur les rapports trimestriels** : si des rapports T1/T2/T3/S1 etc. sont disponibles
   dans les RAPPORTS NARRATIFS, utilise-les pour ESTIMER les tendances et projections annuelles.
   Exemple : "Sur la base du rapport T1 2026 (PNB +12%), on peut estimer que l'exercice 2026 devrait..."
   Distingue clairement les chiffres réels (états financiers annuels) des estimations (rapports trimestriels)
4. Si **plusieurs années** disponibles: montre l'évolution (croissance, amélioration/dégradation des ratios)
5. Conclus avec une recommandation fondamentale (solidité, risques, perspectives)
- NE DIS PAS que les données sont absentes si elles sont fournies ci-dessus

**PARTIE 4 : CONCLUSION D'INVESTISSEMENT**

Rédige un paragraphe de 7-9 lignes synthétisant OBLIGATOIREMENT les 4 parties précédentes:
- **Valorisation (Partie 0)** : les ratios PER/P/B/EV-EBITDA indiquent-ils une sous-évaluation ou surévaluation ?
- **Comportement du cours (Partie 1)** : la tendance récente, la volatilité, les statistiques (kurtosis, skewness) sont-elles favorables ou préoccupantes ?
- **Signaux techniques (Partie 2)** : convergence ou divergence des indicateurs (MM, Bollinger, MACD, RSI, Stochastique)
- **Fondamentaux (Partie 3)** : solidité financière, croissance, estimations issues des rapports trimestriels
- **Prédictions IA (J+1 à J+10)** : la trajectoire prédite confirme-t-elle ou contredit-elle les autres signaux ?
Sur la base de cette synthèse globale:
- Donne une recommandation finale: **ACHAT FORT**, **ACHAT**, **CONSERVER**, **VENTE**, ou **VENTE FORTE**
- Justifie la convergence (ou divergence) entre valorisation, technique, fondamental et prédiction
- Indique le niveau de confiance: Élevé, Moyen, ou Faible
- Mentionne le niveau de risque global: Faible, Moyen, ou Élevé
- Suggère un horizon d'investissement optimal (court terme <3 mois, moyen terme 3-12 mois, long terme >1 an)

═══════════════════════════════════════════════════════════════

RAPPELS IMPÉRATIFS:
- Rédige en français professionnel avec des paragraphes fluides (pas de bullet points)
- Sois précis avec les chiffres — cite les valeurs exactes des données fournies
- Si des analyses fondamentales sont fournies, TU DOIS LES UTILISER — instruction OBLIGATOIRE
- PARTIE 0 : commente tous les ratios de valorisation disponibles avec leur signification pour l'investisseur
- PARTIE 1 : commente OBLIGATOIREMENT les statistiques descriptives (moyenne, médiane, écart-type, kurtosis, skewness)
  en utilisant les interprétations fournies — ces statistiques révèlent le comportement du cours
- PARTIE 3 : précise TOUJOURS l'année des données structurées utilisées
  Si des rapports trimestriels existent, fais des ESTIMATIONS explicites basées dessus
  (distingue données réelles vs estimations — ex: "On estime que..." vs "En 2025, le PNB s'établit à...")
- Mentionne TOUJOURS la date des rapports fondamentaux utilisés
- Reste factuel et objectif
- PARTIE 4 : synthétise OBLIGATOIREMENT les 4 parties + les prédictions IA — c'est la conclusion finale
- LONGUEUR OBLIGATOIRE :
  Partie 0: 4-5 lignes | Partie 1: 6-8 lignes | Partie 2: 2-3 lignes/indicateur + 3-4 lignes conclusion
  Partie 3: 10-12 lignes | Partie 4: 7-9 lignes. Un rapport trop court est un rapport incomplet."""
        
        # ── Rotation Multi-AI: DeepSeek → Claude → Gemini → Mistral ────────────────
        analysis = None
        provider = None

        logging.info(f"    🤖 {symbol}: Tentative DeepSeek...")
        analysis, provider = self._generate_analysis_with_deepseek(symbol, data_dict, prompt)

        if not analysis:
            logging.info(f"    🤖 {symbol}: Tentative Claude (Anthropic)...")
            analysis, provider = self._generate_analysis_with_claude(symbol, data_dict, prompt)

        if not analysis:
            logging.info(f"    🤖 {symbol}: Tentative Gemini...")
            analysis, provider = self._generate_analysis_with_gemini(symbol, data_dict, prompt)

        if not analysis:
            logging.info(f"    🤖 {symbol}: Tentative Mistral...")
            analysis, provider = self._generate_analysis_with_mistral(symbol, data_dict, prompt)

        if not analysis:
            if attempt < max_attempts:
                logging.warning(f"    ⚠️  {symbol}: Toutes API échouées, retry {attempt+1}/{max_attempts}")
                time.sleep(10)
                return self._generate_professional_analysis(symbol, data_dict, attempt + 1, max_attempts)
            else:
                logging.error(f"    ❌ {symbol}: Échec définitif après {max_attempts} tentatives")
                return self._generate_fallback_analysis(symbol, data_dict)
        
        # ── Vérification post-génération: l'IA a-t-elle bien utilisé les fondamentaux ──
        if has_fundamental and "aucune donnée" in analysis.lower() and "fondamental" in analysis.lower():
            logging.warning(f"    ⚠️ {symbol}: L'IA a potentiellement ignoré les analyses fondamentales malgré les instructions")
        
        logging.info(f"    ✅ {symbol}: Analyse générée via {provider.upper()}")
        return analysis

    def _generate_fallback_analysis(self, symbol, data_dict):
        """Analyse de secours structurée"""
        analysis = f"**ANALYSE DE {symbol}**\n\n"
        
        analysis += "**PARTIE 1 : ANALYSE DE L'ÉVOLUTION DU COURS (100 derniers jours)**\n\n"
        if data_dict.get('historical_summary'):
            analysis += f"{data_dict['historical_summary']}\n\n"
        else:
            analysis += "Les données historiques sur 100 jours ne sont pas disponibles pour cette action.\n\n"
        
        analysis += "**PARTIE 2 : ANALYSE TECHNIQUE DÉTAILLÉE**\n\n"
        
        signals = []
        if data_dict.get('mm_decision'):
            signals.append(data_dict['mm_decision'])
            analysis += f"**Moyennes Mobiles**: Les moyennes mobiles (MM20: {data_dict.get('mm_20', 'N/A')}, MM50: {data_dict.get('mm_50', 'N/A')}) suggèrent une tendance {data_dict['mm_decision'].lower()}.\n\n"
        
        if data_dict.get('bollinger_decision'):
            signals.append(data_dict['bollinger_decision'])
            analysis += f"**Bandes de Bollinger**: Les bandes de Bollinger indiquent un signal {data_dict['bollinger_decision'].lower()}.\n\n"
        
        if data_dict.get('macd_decision'):
            signals.append(data_dict['macd_decision'])
            analysis += f"**MACD**: Le MACD confirme une position {data_dict['macd_decision'].lower()}.\n\n"
        
        if data_dict.get('rsi_decision'):
            signals.append(data_dict['rsi_decision'])
            analysis += f"**RSI**: Le RSI ({data_dict.get('rsi_value', 'N/A')}) signale {data_dict['rsi_decision'].lower()}.\n\n"
        
        if data_dict.get('stochastic_decision'):
            signals.append(data_dict['stochastic_decision'])
            analysis += f"**Stochastique**: Le stochastique recommande {data_dict['stochastic_decision'].lower()}.\n\n"
        
        buy_count = signals.count('Achat')
        sell_count = signals.count('Vente')
        
        if buy_count > sell_count:
            analysis += "**Conclusion technique**: Les indicateurs convergent majoritairement vers un signal d'achat.\n\n"
        elif sell_count > buy_count:
            analysis += "**Conclusion technique**: Les indicateurs convergent majoritairement vers un signal de vente.\n\n"
        else:
            analysis += "**Conclusion technique**: Les indicateurs sont mixtes, suggérant une position de conservation.\n\n"
        
        analysis += "**PARTIE 3 : ANALYSE FONDAMENTALE**\n\n"
        fundamental_text = data_dict.get('fundamental_analyses', '')
        if fundamental_text and fundamental_text.strip() and "Aucun rapport" not in fundamental_text:
            analysis += f"{fundamental_text}\n\n"
            analysis += "Ces analyses fondamentales, bien que pouvant dater de périodes antérieures, fournissent des indications précieuses sur la santé financière historique de l'entreprise.\n\n"
        else:
            analysis += "Aucune analyse fondamentale n'est disponible dans la base de données pour cette société.\n\n"
        
        analysis += "**PARTIE 4 : CONCLUSION D'INVESTISSEMENT**\n\n"
        
        if buy_count > sell_count:
            analysis += f"**Recommandation: ACHAT**\n\nEn combinant l'analyse technique ({buy_count} signaux d'achat sur {len(signals)}) et les éléments fondamentaux disponibles, cette action présente des perspectives favorables. Niveau de confiance: Moyen. Niveau de risque: Moyen. Horizon: Moyen terme.\n"
        elif sell_count > buy_count:
            analysis += f"**Recommandation: VENTE**\n\nL'analyse technique ({sell_count} signaux de vente sur {len(signals)}) suggère une prudence. Il est recommandé d'envisager une sortie de position. Niveau de confiance: Moyen. Niveau de risque: Élevé. Horizon: Court terme.\n"
        else:
            analysis += "**Recommandation: CONSERVER**\n\nLes signaux techniques mixtes et l'absence d'éléments fondamentaux déterminants suggèrent de maintenir la position actuelle. Niveau de confiance: Faible. Niveau de risque: Moyen. Horizon: Moyen terme.\n"
        
        return analysis

    # ============================================================================
    # NOUVELLES ANALYSES AVANCÉES (V29.0)
    # ============================================================================
    
    def _calculate_sector_analysis(self, all_company_data):
        """1. ANALYSE PAR SECTEUR: Performance moyenne, sentiment, risque"""
        logging.info("📊 Calcul de l'analyse sectorielle...")
        
        sector_stats = defaultdict(lambda: {
            'companies': [],
            'prices': [],
            'performances': [],
            'recommendations': [],
            'risk_levels': [],
            'volumes': []
        })
        
        for symbol, data in all_company_data.items():
            sector = data.get('sector', 'Non classifié')
            if not sector or str(sector).strip() == '':
                sector = 'Non classifié'
            
            company_name = data.get('company_name', 'N/A')
            sector_stats[sector]['companies'].append(f"{symbol} ({company_name})")
            
            if data.get('price_evolution_100d') and pd.notna(data.get('price_evolution_100d')):
                sector_stats[sector]['performances'].append(data['price_evolution_100d'])
            
            if data.get('recommendation'):
                sector_stats[sector]['recommendations'].append(data['recommendation'])
            
            if data.get('risk_level'):
                sector_stats[sector]['risk_levels'].append(data['risk_level'])
            
            if data.get('current_price') and pd.notna(data.get('current_price')):
                sector_stats[sector]['prices'].append(data['current_price'])
        
        # Calculer les statistiques par secteur
        sector_analysis = {}
        for sector, stats in sector_stats.items():
            # Performance moyenne (avec gestion du cas vide)
            avg_perf = sum(stats['performances']) / len(stats['performances']) if stats['performances'] else 0
            
            # Sentiment général (recommandation la plus fréquente)
            rec_counter = Counter(stats['recommendations'])
            sentiment = rec_counter.most_common(1)[0][0] if rec_counter else 'NEUTRE'
            
            # Niveau de risque moyen
            risk_mapping = {'Faible': 1, 'Moyen': 2, 'Élevé': 3}
            risk_scores = [risk_mapping.get(r, 2) for r in stats['risk_levels']]
            avg_risk_score = sum(risk_scores) / len(risk_scores) if risk_scores else 2
            avg_risk = 'Faible' if avg_risk_score < 1.5 else 'Moyen' if avg_risk_score < 2.5 else 'Élevé'
            
            # Prix moyen (avec gestion du cas vide)
            prix_moyen = sum(stats['prices']) / len(stats['prices']) if stats['prices'] else 0
            
            sector_analysis[sector] = {
                'nb_societes': len(stats['companies']),
                'societes': stats['companies'],
                'performance_moyenne': avg_perf,
                'sentiment_general': sentiment,
                'risque_moyen': avg_risk,
                'prix_moyen': prix_moyen,
                'distribution_recommandations': dict(rec_counter)
            }
        
        return sector_analysis

    def _calculate_signal_convergence_matrix(self, all_company_data):
        """2. MATRICE DE CONVERGENCE: Signaux techniques vs fondamentaux"""
        logging.info("🔄 Calcul de la matrice de convergence...")
        
        matrix = {
            'convergence_forte': [],  # Tech=Achat Fort + Fond=Achat Fort
            'convergence_achat': [],  # Tech=Achat + Fond=Achat
            'convergence_vente': [],  # Tech=Vente + Fond=Vente
            'divergence_forte': [],   # Tech=Achat mais Fond=Vente ou inverse
            'divergence_moderee': [], # Tech=Achat mais Fond=Neutre
            'neutre': []              # Signaux mixtes
        }
        
        for symbol, data in all_company_data.items():
            company_name = data.get('company_name', 'N/A')
            full_name = f"{symbol} ({company_name})"
            
            # Signal technique (basé sur la moyenne des indicateurs)
            tech_signals = []
            if data.get('mm_decision'): tech_signals.append(data['mm_decision'])
            if data.get('bollinger_decision'): tech_signals.append(data['bollinger_decision'])
            if data.get('macd_decision'): tech_signals.append(data['macd_decision'])
            if data.get('rsi_decision'): tech_signals.append(data['rsi_decision'])
            if data.get('stochastic_decision'): tech_signals.append(data['stochastic_decision'])
            
            buy_tech = sum(1 for s in tech_signals if 'Achat' in str(s))
            sell_tech = sum(1 for s in tech_signals if 'Vente' in str(s))
            
            if buy_tech > sell_tech:
                tech_signal = 'ACHAT'
            elif sell_tech > buy_tech:
                tech_signal = 'VENTE'
            else:
                tech_signal = 'NEUTRE'
            
            # Signal fondamental (extrait de la recommandation finale)
            final_rec = data.get('recommendation', 'CONSERVER')
            if 'ACHAT' in final_rec:
                fund_signal = 'ACHAT'
            elif 'VENTE' in final_rec:
                fund_signal = 'VENTE'
            else:
                fund_signal = 'NEUTRE'
            
            # Classement dans la matrice
            if tech_signal == 'ACHAT' and fund_signal == 'ACHAT':
                if 'FORT' in final_rec:
                    matrix['convergence_forte'].append(full_name)
                else:
                    matrix['convergence_achat'].append(full_name)
            elif tech_signal == 'VENTE' and fund_signal == 'VENTE':
                matrix['convergence_vente'].append(full_name)
            elif (tech_signal == 'ACHAT' and fund_signal == 'VENTE') or (tech_signal == 'VENTE' and fund_signal == 'ACHAT'):
                matrix['divergence_forte'].append(full_name)
            elif tech_signal != fund_signal and (tech_signal == 'NEUTRE' or fund_signal == 'NEUTRE'):
                matrix['divergence_moderee'].append(full_name)
            else:
                matrix['neutre'].append(full_name)
        
        return matrix

    def _calculate_liquidity_analysis(self, all_company_data):
        """3. ANALYSE DE LIQUIDITÉ: Volumes moyens et relation avec performance"""
        logging.info("💧 Calcul de l'analyse de liquidité...")
        
        # Récupérer les volumes moyens sur 30 jours pour chaque société
        liquidity_data = []
        
        for symbol, data in all_company_data.items():
            company_id = data.get('company_id')
            company_name = data.get('company_name', 'N/A')
            
            if not company_id:
                continue
            
            # Calculer le volume moyen sur 30 jours
            query = f"""
            SELECT AVG(volume) as avg_volume, AVG(value) as avg_value
            FROM historical_data
            WHERE company_id = {company_id}
              AND trade_date >= CURRENT_DATE - INTERVAL '30 days'
            """
            
            try:
                df = pd.read_sql(query, self.db_conn)
                if not df.empty:
                    avg_volume = df.iloc[0]['avg_volume'] if pd.notna(df.iloc[0]['avg_volume']) else 0
                    avg_value = df.iloc[0]['avg_value'] if pd.notna(df.iloc[0]['avg_value']) else 0
                    
                    # Ne garder que les sociétés avec des données de volume
                    if avg_volume > 0:
                        liquidity_data.append({
                            'symbol': symbol,
                            'company_name': company_name,
                            'avg_volume': avg_volume,
                            'avg_value': avg_value,
                            'performance': data.get('price_evolution_100d', 0) if pd.notna(data.get('price_evolution_100d')) else 0,
                            'recommendation': data.get('recommendation', 'N/A'),
                            'risk_level': data.get('risk_level', 'N/A')
                        })
            except Exception as e:
                logging.error(f"❌ Erreur volume pour {symbol}: {e}")
        
        # Trier par volume décroissant
        liquidity_data.sort(key=lambda x: x['avg_volume'], reverse=True)
        
        # Classifier en catégories (avec protection contre les listes vides)
        total = len(liquidity_data)
        if total == 0:
            return {
                'high_liquidity': [],
                'medium_liquidity': [],
                'low_liquidity': [],
                'all_data': []
            }
        
        # Assurer au moins 1 élément par catégorie si possible
        high_count = max(1, int(total * 0.2)) if total >= 5 else total
        medium_start = high_count
        medium_end = max(medium_start + 1, int(total * 0.6)) if total >= 5 else total
        
        high_liquidity = liquidity_data[:high_count]
        medium_liquidity = liquidity_data[medium_start:medium_end] if total > high_count else []
        low_liquidity = liquidity_data[medium_end:] if total > medium_end else []
        
        return {
            'high_liquidity': high_liquidity,
            'medium_liquidity': medium_liquidity,
            'low_liquidity': low_liquidity,
            'all_data': liquidity_data
        }

    def _calculate_top_divergences(self, all_company_data):
        """4. TOP 10 DES DIVERGENCES: Sociétés avec signaux contradictoires"""
        logging.info("⚠️  Calcul des divergences majeures...")
        
        divergences = []
        
        for symbol, data in all_company_data.items():
            company_name = data.get('company_name', 'N/A')
            
            # Compter signaux d'achat vs vente
            signals = {
                'MM': data.get('mm_decision'),
                'Bollinger': data.get('bollinger_decision'),
                'MACD': data.get('macd_decision'),
                'RSI': data.get('rsi_decision'),
                'Stochastique': data.get('stochastic_decision')
            }
            
            buy_signals = [k for k, v in signals.items() if v and 'Achat' in str(v)]
            sell_signals = [k for k, v in signals.items() if v and 'Vente' in str(v)]
            
            # Score de divergence (écart entre signaux)
            divergence_score = abs(len(buy_signals) - len(sell_signals))
            
            # Divergence avec fondamental
            final_rec = data.get('recommendation', 'CONSERVER')
            fund_is_buy = 'ACHAT' in final_rec
            fund_is_sell = 'VENTE' in final_rec
            
            tech_majority_buy = len(buy_signals) > len(sell_signals)
            tech_majority_sell = len(sell_signals) > len(buy_signals)
            
            if (tech_majority_buy and fund_is_sell) or (tech_majority_sell and fund_is_buy):
                divergence_score += 3  # Bonus si divergence tech/fondamental
            
            if divergence_score > 0:
                divergences.append({
                    'symbol': symbol,
                    'company_name': company_name,
                    'divergence_score': divergence_score,
                    'buy_signals': buy_signals,
                    'sell_signals': sell_signals,
                    'final_recommendation': final_rec,
                    'description': self._describe_divergence(buy_signals, sell_signals, final_rec)
                })
        
        # Trier par score décroissant
        divergences.sort(key=lambda x: x['divergence_score'], reverse=True)
        
        return divergences[:10]  # Top 10

    def _describe_divergence(self, buy_signals, sell_signals, final_rec):
        """Génère une description textuelle de la divergence"""
        desc = []
        
        if buy_signals:
            desc.append(f"Signaux d'achat: {', '.join(buy_signals)}")
        
        if sell_signals:
            desc.append(f"Signaux de vente: {', '.join(sell_signals)}")
        
        desc.append(f"Recommandation finale: {final_rec}")
        
        return " | ".join(desc)

    def _calculate_risk_horizon_matrix(self, all_company_data):
        """5. MATRICE RISQUE vs HORIZON"""
        logging.info("📈 Calcul de la matrice Risque/Horizon...")
        
        matrix = {
            'faible_court': [],
            'faible_moyen': [],
            'faible_long': [],
            'moyen_court': [],
            'moyen_moyen': [],
            'moyen_long': [],
            'eleve_court': [],
            'eleve_moyen': [],
            'eleve_long': []
        }
        
        for symbol, data in all_company_data.items():
            company_name = data.get('company_name', 'N/A')
            full_name = f"{symbol} ({company_name})"
            
            risk = data.get('risk_level', 'Moyen').lower()
            horizon = data.get('investment_horizon', 'Moyen terme').lower()
            
            # Normaliser horizon
            if 'court' in horizon:
                horizon_cat = 'court'
            elif 'long' in horizon:
                horizon_cat = 'long'
            else:
                horizon_cat = 'moyen'
            
            # Normaliser risque
            if 'faible' in risk:
                risk_cat = 'faible'
            elif 'élevé' in risk or 'eleve' in risk:
                risk_cat = 'eleve'
            else:
                risk_cat = 'moyen'
            
            key = f"{risk_cat}_{horizon_cat}"
            if key in matrix:
                matrix[key].append(full_name)
        
        return matrix

    def _calculate_risk_score(self, data, market_indicators=None):
        """
        Score de risque chiffré (0-100) — 5 critères pondérés
        ─────────────────────────────────────────────────────
        1. Volatilité des prix        30 %  (CV = σ_prix / μ_prix)
        2. Bêta réel vs BRVM Composite 20 %  (Cov / Var — régression OLS)
        3. Liquidité (volume moyen)   20 %
        4. Divergence des signaux     15 %  (CORRIGÉ : divergence = risque)
        5. Stabilité des rendements   15 %  (σ des rendements journaliers)
        """
        import numpy as np

        risk_score = 0
        details    = {}

        # ── Données historiques de la société ────────────────────────────────
        hist_df   = self._get_historical_data_100days(data.get('company_id'))
        vol_coeff = 0.0   # coefficient de variation (utilisé aussi pour bêta de secours)

        # ═══════════════════════════════════════════════════════════════════
        # CRITÈRE 1 — VOLATILITÉ DES PRIX (30 %)
        # CV = écart-type(prix) / moyenne(prix)
        # Mesure la dispersion relative des cours autour de leur moyenne.
        # Un CV élevé = fortes oscillations = titre risqué.
        # ═══════════════════════════════════════════════════════════════════
        if not hist_df.empty and len(hist_df) > 1:
            prices     = hist_df['price'].dropna().astype(float).values
            mean_price = float(np.mean(prices))
            if mean_price > 0:
                vol_coeff = float(np.std(prices)) / mean_price

            if vol_coeff < 0.05:
                vol_score = 10
                vol_label = "Faible"
            elif vol_coeff < 0.15:
                vol_score = 30
                vol_label = "Moyenne"
            else:
                vol_score = 60
                vol_label = "Élevée"

            details['volatilite'] = (
                f"{vol_coeff*100:.2f}% ({vol_label}) — "
                f"μ={mean_price:,.0f} FCFA, σ={np.std(prices):,.0f} FCFA"
            )
            risk_score += vol_score * 0.30
        else:
            details['volatilite'] = "Données insuffisantes"
            risk_score += 30 * 0.30   # score moyen par défaut

        # ═══════════════════════════════════════════════════════════════════
        # CRITÈRE 2 — BÊTA RÉEL vs BRVM COMPOSITE (20 %)
        #
        # β = Cov(r_société, r_indice) / Var(r_indice)
        #
        # Méthode :
        #   • Charger les rendements journaliers du titre (hist_df)
        #   • Joindre avec les rendements journaliers du BRVM Composite
        #     (new_market_indicators.brvm_composite) sur la date commune
        #   • Calculer β par régression des moindres carrés (np.polyfit ou
        #     formule directe Cov/Var)
        #
        # Interprétation :
        #   β < 0.8  → titre défensif, moins volatile que le marché
        #   β ≈ 1.0  → titre évolue comme le marché
        #   β > 1.2  → titre amplifie les mouvements du marché (agressif)
        #   β < 0    → titre contre-cyclique (rare sur la BRVM)
        # ═══════════════════════════════════════════════════════════════════
        beta       = None
        beta_label = "N/D"
        beta_score = 20   # neutre par défaut

        try:
            # Charger l'historique BRVM Composite sur 100 jours
            query_idx = """
                SELECT extraction_date::date AS trade_date,
                       brvm_composite
                FROM   new_market_indicators
                WHERE  brvm_composite IS NOT NULL
                  AND  brvm_composite > 0
                  AND  extraction_date >= CURRENT_DATE - INTERVAL '150 days'
                ORDER  BY extraction_date ASC
                LIMIT  100;
            """
            import pandas as pd
            df_idx = pd.read_sql(query_idx, self.db_conn)

            if not hist_df.empty and not df_idx.empty:
                # Préparer les rendements du titre
                df_titre = hist_df[['trade_date', 'price']].copy()
                df_titre['trade_date'] = pd.to_datetime(df_titre['trade_date']).dt.date
                df_titre = df_titre.sort_values('trade_date').drop_duplicates('trade_date')
                df_titre['r_titre'] = df_titre['price'].astype(float).pct_change()

                # Préparer les rendements de l'indice
                df_idx['trade_date'] = pd.to_datetime(df_idx['trade_date']).dt.date
                df_idx = df_idx.sort_values('trade_date').drop_duplicates('trade_date')
                df_idx['r_idx'] = df_idx['brvm_composite'].astype(float).pct_change()

                # Jointure sur date commune
                merged = pd.merge(
                    df_titre[['trade_date','r_titre']],
                    df_idx[['trade_date','r_idx']],
                    on='trade_date', how='inner'
                ).dropna()

                if len(merged) >= 10:
                    r_titre = merged['r_titre'].values
                    r_idx   = merged['r_idx'].values

                    # β = Cov(r_titre, r_idx) / Var(r_idx)
                    cov_matrix = np.cov(r_titre, r_idx, ddof=1)
                    cov_ti     = float(cov_matrix[0, 1])
                    var_idx    = float(cov_matrix[1, 1])

                    if var_idx > 1e-12:   # éviter division par zéro
                        beta = cov_ti / var_idx

                        # Corrélation R pour contexte
                        corr = np.corrcoef(r_titre, r_idx)[0, 1]

                        # Scoring du bêta
                        if beta < 0:
                            beta_score = 25
                            beta_label = "Contre-cyclique"
                        elif beta < 0.8:
                            beta_score = 10
                            beta_label = "Défensif"
                        elif beta <= 1.2:
                            beta_score = 20
                            beta_label = "Neutre (suit le marché)"
                        elif beta <= 1.8:
                            beta_score = 35
                            beta_label = "Agressif"
                        else:
                            beta_score = 50
                            beta_label = "Très agressif"

                        details['beta'] = (
                            f"β={beta:.4f} ({beta_label}) — "
                            f"Cov={cov_ti:.6f}, Var(idx)={var_idx:.6f}, "
                            f"ρ={corr:.3f}, n={len(merged)} obs"
                        )
                    else:
                        details['beta'] = "Var(indice)≈0 — indice sans mouvement sur la période"
                else:
                    # Pas assez de dates communes — fallback sur proxy volatilité
                    beta_fallback = vol_coeff / 0.10 if vol_coeff > 0 else 1.0
                    beta_score = 10 if beta_fallback < 0.8 else (20 if beta_fallback < 1.2 else 40)
                    details['beta'] = (
                        f"β≈{beta_fallback:.2f} (proxy, <10 obs communes) — "
                        f"Données BRVM insuffisantes pour le bêta réel"
                    )
            else:
                details['beta'] = "Données indisponibles (titre ou indice)"
        except Exception as _e_beta:
            # Fallback sans plantage
            beta_fallback = vol_coeff / 0.10 if vol_coeff > 0 else 1.0
            beta_score    = 10 if beta_fallback < 0.8 else (20 if beta_fallback < 1.2 else 40)
            details['beta'] = f"β≈{beta_fallback:.2f} (proxy — erreur calcul : {str(_e_beta)[:60]})"

        risk_score += beta_score * 0.20

        # ═══════════════════════════════════════════════════════════════════
        # CRITÈRE 3 — LIQUIDITÉ (20 %)
        # Volume moyen journalier sur 100 jours.
        # Un titre peu liquide est risqué : difficile à sortir en cas de
        # mauvaise nouvelle. Sur la BRVM, beaucoup de titres ont <500 titres/j.
        # ═══════════════════════════════════════════════════════════════════
        avg_volume = 0.0
        if not hist_df.empty and 'volume' in hist_df.columns:
            avg_volume = float(hist_df['volume'].dropna().mean())

        if avg_volume > 10_000:
            liq_score  = 5
            liq_label  = "Excellente"
        elif avg_volume > 1_000:
            liq_score  = 15
            liq_label  = "Bonne"
        elif avg_volume > 100:
            liq_score  = 30
            liq_label  = "Faible"
        elif avg_volume > 0:
            liq_score  = 45
            liq_label  = "Très faible"
        else:
            liq_score  = 30
            liq_label  = "Non calculable"

        details['liquidite'] = f"{avg_volume:,.0f} titres/j ({liq_label})"
        risk_score += liq_score * 0.20

        # ═══════════════════════════════════════════════════════════════════
        # CRITÈRE 4 — DIVERGENCE DES SIGNAUX TECHNIQUES (15 %)  [CORRIGÉ]
        #
        # LOGIQUE CORRIGÉE :
        #   • Divergence forte (signaux contradictoires) = incertitude = RISQUE ÉLEVÉ
        #   • Convergence forte (tous dans le même sens) = clarté = RISQUE FAIBLE
        #
        # Score de divergence = min(buy, sell) / max(1, total_signaux)
        #   → 0.0 = convergence totale (pas de risque sur ce critère)
        #   → 0.5 = parfaite divergence (50% Achat / 50% Vente) = risque max
        # ═══════════════════════════════════════════════════════════════════
        tech_signals = []
        for key in ['mm_decision', 'bollinger_decision', 'macd_decision',
                    'rsi_decision', 'stochastic_decision']:
            val = data.get(key)
            if val and str(val).strip():
                tech_signals.append(str(val))

        n_sig  = len(tech_signals)
        n_buy  = sum(1 for s in tech_signals if 'Achat' in s)
        n_sell = sum(1 for s in tech_signals if 'Vente' in s)

        # Taux de divergence entre 0 (convergence totale) et 1 (max opposition)
        if n_sig > 0:
            div_rate  = min(n_buy, n_sell) / n_sig   # 0→convergent, 0.5→parfaitement partagé
            div_score = int(div_rate * 60)            # 0–30 points (max 60 × 0.5)
        else:
            div_rate  = 0.0
            div_score = 15   # score moyen si pas de signaux

        details['divergence'] = (
            f"Achat={n_buy} / Vente={n_sell} / Total={n_sig} "
            f"— taux divergence={div_rate*100:.0f}% "
            f"({'Convergent ✅' if div_rate < 0.2 else 'Mixte ⚠️' if div_rate < 0.4 else 'Très divergent ❌'})"
        )
        risk_score += div_score * 0.15

        # ═══════════════════════════════════════════════════════════════════
        # CRITÈRE 5 — STABILITÉ DES RENDEMENTS (15 %)
        # σ des rendements journaliers × 100 (= volatilité journalière %)
        # Différent de la volatilité des prix (critère 1) :
        #   • Critère 1 = CV des niveaux de prix (dispersion long terme)
        #   • Critère 5 = σ des variations quotidiennes (chocs journaliers)
        # ═══════════════════════════════════════════════════════════════════
        if not hist_df.empty and len(hist_df) > 2:
            returns          = hist_df['price'].astype(float).pct_change().dropna()
            daily_vol        = float(returns.std() * 100) if len(returns) > 0 else 10.0
            annualized_vol   = daily_vol * (252 ** 0.5)

            if daily_vol < 1.0:
                perf_score  = 5
                stab_label  = "Très stable"
            elif daily_vol < 3.0:
                perf_score  = 10
                stab_label  = "Stable"
            elif daily_vol < 5.0:
                perf_score  = 20
                stab_label  = "Modérée"
            else:
                perf_score  = 35
                stab_label  = "Instable"

            details['stabilite'] = (
                f"σ_j={daily_vol:.2f}%/j ({stab_label}) — "
                f"Vol. annualisée≈{annualized_vol:.1f}%"
            )
            risk_score += perf_score * 0.15
        else:
            details['stabilite'] = "Données insuffisantes"
            risk_score += 15 * 0.15

        # ── Classification finale du risque ─────────────────────────────────
        risk_score = max(0.0, min(100.0, risk_score))

        if   risk_score < 20: risk_level = "Faible"
        elif risk_score < 45: risk_level = "Moyen"
        elif risk_score < 70: risk_level = "Élevé"
        else:                 risk_level = "Très élevé"

        return {
            'score':   round(risk_score, 2),
            'level':   risk_level,
            'details': details
        }

    # ============================================================================
    # SAUVEGARDE ET GÉNÉRATION DOCUMENT
    # ============================================================================
    
    def _save_to_database(self, report_date, synthesis_text, top_10, flop_10, market_events, all_company_data, filename):
        """Sauvegarde structurée dans la base de données"""
        logging.info("💾 Sauvegarde dans la base de données...")
        
        try:
            with self.db_conn.cursor() as cur:
                cur.execute("""
                    INSERT INTO report_summary (
                        report_date, synthesis_text, top_10_buy, flop_10_sell, 
                        market_events, total_companies_analyzed, report_file_name
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (report_date) DO UPDATE SET
                        synthesis_text = EXCLUDED.synthesis_text,
                        top_10_buy = EXCLUDED.top_10_buy,
                        flop_10_sell = EXCLUDED.flop_10_sell,
                        market_events = EXCLUDED.market_events,
                        total_companies_analyzed = EXCLUDED.total_companies_analyzed,
                        report_file_name = EXCLUDED.report_file_name
                    RETURNING id;
                """, (
                    report_date,
                    synthesis_text,
                    json.dumps(top_10, ensure_ascii=False),
                    json.dumps(flop_10, ensure_ascii=False),
                    market_events,
                    len(all_company_data),
                    filename
                ))
                
                report_summary_id = cur.fetchone()[0]
                
                for symbol, company_data in all_company_data.items():
                    # ✅ V30: HISTORISATION - Plus de ON CONFLICT, toujours INSERT
                    cur.execute("""
                        INSERT INTO report_company_analysis (
                            report_summary_id, company_id, symbol, company_name,
                            current_price, price_evolution_100d, highest_price_100d, lowest_price_100d,
                            mm20, mm50, mm_decision,
                            bollinger_upper, bollinger_lower, bollinger_decision,
                            macd_value, macd_signal, macd_decision,
                            rsi_value, rsi_decision,
                            stochastic_k, stochastic_d, stochastic_decision,
                            price_evolution_text, technical_conclusion, fundamental_analysis, investment_conclusion,
                            recommendation, confidence_level, risk_level, investment_horizon,
                            technical_decision, fundamental_decision, risk_score, risk_details
                        ) VALUES (
                            %s, %s, %s, %s,
                            %s, %s, %s, %s,
                            %s, %s, %s,
                            %s, %s, %s,
                            %s, %s, %s,
                            %s, %s,
                            %s, %s, %s,
                            %s, %s, %s, %s,
                            %s, %s, %s, %s,
                            %s, %s, %s, %s
                        );
                    """, (
                        report_summary_id,
                        company_data.get('company_id'),
                        symbol,
                        company_data.get('company_name'),
                        company_data.get('current_price'),
                        company_data.get('price_evolution_100d'),
                        company_data.get('highest_price_100d'),
                        company_data.get('lowest_price_100d'),
                        company_data.get('mm20'),
                        company_data.get('mm50'),
                        company_data.get('mm_decision'),
                        company_data.get('bollinger_upper'),
                        company_data.get('bollinger_lower'),
                        company_data.get('bollinger_decision'),
                        company_data.get('macd_value'),
                        company_data.get('macd_signal'),
                        company_data.get('macd_decision'),
                        company_data.get('rsi_value'),
                        company_data.get('rsi_decision'),
                        company_data.get('stochastic_k'),
                        company_data.get('stochastic_d'),
                        company_data.get('stochastic_decision'),
                        company_data.get('full_analysis', ''),
                        company_data.get('technical_conclusion', ''),
                        company_data.get('fundamental_analysis', ''),
                        company_data.get('investment_conclusion', ''),
                        company_data.get('recommendation'),
                        company_data.get('confidence_level'),
                        company_data.get('risk_level'),
                        company_data.get('investment_horizon'),
                        company_data.get('technical_decision', 'NEUTRE'),
                        company_data.get('fundamental_decision', 'NEUTRE'),
                        company_data.get('risk_score', 0.0),
                        company_data.get('risk_details', '{}')
                    ))
                
                self.db_conn.commit()
                logging.info(f"   ✅ Rapport sauvegardé (ID: {report_summary_id})")
                
        except Exception as e:
            logging.error(f"❌ Erreur sauvegarde DB: {e}")
            self.db_conn.rollback()

    def _add_table_with_shading(self, doc, data, headers, column_widths=None):
        """Ajoute un tableau avec mise en forme"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # En-têtes
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            # Couleur de fond grise pour l'en-tête
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9D9D9')
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # Données
        for row_data in data:
            row = table.add_row()
            for i, value in enumerate(row_data):
                row.cells[i].text = str(value)
        
        return table

    def _create_word_document(self, all_analyses, all_company_data):
        """Création du document Word professionnel ULTRA-COMPLET"""
        logging.info("📄 Création du document Word ULTIMATE...")
        
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # ========== PAGE DE TITRE ==========
        title = doc.add_heading('RAPPORT D\'ANALYSE BRVM', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        subtitle = doc.add_paragraph(f"Rapport d'investissement professionnel - Édition Ultimate")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
        
        date_p = doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_p.runs[0]
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        
        version_p = doc.add_paragraph(f"Version 30.2 - Analyses Multi-AI (DeepSeek + Gemini + Mistral)")
        version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_run = version_p.runs[0]
        version_run.font.size = Pt(9)
        version_run.font.italic = True
        version_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()

        # ========== RÉSUMÉ EXÉCUTIF (Page 1) ==========
        market_indicators_pre = self._get_market_indicators()
        exec_data = self._build_executive_summary(all_company_data, market_indicators_pre)

        doc.add_paragraph()
        exec_box = doc.add_paragraph()
        exec_box.paragraph_format.space_before = Pt(6)
        exec_box.paragraph_format.space_after  = Pt(6)

        def _add_exec_line(para, label, value, val_color=None):
            r_lbl = para.add_run(f"  {label}  ")
            r_lbl.bold = True
            r_lbl.font.size = Pt(10)
            r_val = para.add_run(str(value) + "\n")
            r_val.font.size = Pt(10)
            if val_color:
                r_val.font.color.rgb = val_color

        # Encadré bleu foncé
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement

        def _shade_para(para, hex_color):
            pPr = para._p.get_or_add_pPr()
            shd = _OxmlElement('w:shd')
            shd.set(_qn('w:val'),  'clear')
            shd.set(_qn('w:color'),'auto')
            shd.set(_qn('w:fill'), hex_color)
            pPr.append(shd)

        lines = [
            ("📊 RÉSUMÉ EXÉCUTIF", "", None),
            ("─" * 55, "", None),
            ("Marché :", exec_data['marche'],
             RGBColor(0,128,0) if 'HAUSSE' in exec_data['marche'].upper()
             else RGBColor(192,0,0) if 'BAIS' in exec_data['marche'].upper()
             else RGBColor(80,80,80)),
            ("Signaux :",
             f"{exec_data['achats']} ACHAT  |  {exec_data['neutres']} NEUTRE  |  {exec_data['ventes']} VENTE  (sur {exec_data['total']} sociétés)",
             RGBColor(0,80,160)),
            ("Top opportunité :",
             f"{exec_data['top_sym']}  (score {exec_data['top_score']}/100)",
             RGBColor(0,128,0)),
            ("Secteur surperformant :",
             f"{exec_data['best_sector']}  ({exec_data['best_sec_pct']:+.1f}%)",
             RGBColor(0,128,0)),
            ("Secteur sous-performant :", exec_data['worst_sector'], RGBColor(160,60,0)),
            ("Divergences tech/fond :",
             f"{exec_data['divergences']} sociétés — vérifier avant d'investir",
             RGBColor(160,100,0)),
            ("⚠️ Liquidité réduite :",
             f"{exec_data['low_liq']} titres à faible liquidité — risque de sortie",
             RGBColor(192,0,0)),
        ]

        for label, value, color in lines:
            ep = doc.add_paragraph()
            ep.paragraph_format.left_indent  = Pt(18)
            ep.paragraph_format.right_indent = Pt(18)
            ep.paragraph_format.space_before = Pt(1)
            ep.paragraph_format.space_after  = Pt(1)
            _shade_para(ep, 'EBF5FB')
            if label.startswith('─') or label.startswith('📊'):
                r = ep.add_run(label if label.startswith('📊') else '')
                r.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0,51,102)
            else:
                r_l = ep.add_run(f"{label}  ")
                r_l.bold = True
                r_l.font.size = Pt(9.5)
                r_v = ep.add_run(value)
                r_v.font.size = Pt(9.5)
                if color:
                    r_v.font.color.rgb = color

        doc.add_paragraph()
        doc.add_page_break()

        # ========== SYNTHÈSE GÉNÉRALE ==========
        doc.add_heading('SYNTHÈSE GÉNÉRALE', level=1)
        
        market_indicators = self._get_market_indicators()
        intro = doc.add_paragraph(
            f"Ce rapport présente une analyse détaillée de {len(all_analyses)} sociétés cotées "
            f"à la Bourse Régionale des Valeurs Mobilières (BRVM). "
        )
        
        if market_indicators and market_indicators.get('composite'):
            intro.add_run(f"L'indice BRVM Composite s'établit à {market_indicators['composite']:.2f} points ")
            
            if market_indicators.get('composite_var_day') is not None:
                var_day = market_indicators['composite_var_day']
                if var_day > 0:
                    run = intro.add_run(f"(+{var_day:.2f}%)")
                    run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    run = intro.add_run(f"({var_day:.2f}%)")
                    run.font.color.rgb = RGBColor(192, 0, 0)
                intro.add_run(f" sur la séance. ")
            
            # ✅ Capitalisation corrigée: valeur brute en FCFA → diviser par 1e9 pour obtenir des milliards
            # Exemple: 15 660 629 773 994 FCFA / 1e9 = 15 660,630 milliards FCFA
            if market_indicators.get('capitalisation'):
                cap_raw = market_indicators['capitalisation']
                cap_milliards = cap_raw / 1e9  # toujours en milliards FCFA
                # Formatage français: séparateur de milliers = espace, décimale = virgule
                cap_entier = int(cap_milliards)
                cap_decimale = round((cap_milliards - cap_entier) * 1000)
                cap_display = f"{cap_entier:,}".replace(",", " ") + f",{cap_decimale:03d}"
                intro.add_run(f"La capitalisation globale du marché atteint {cap_display} milliards FCFA.")
        else:
            intro.add_run("Les indicateurs de marché seront mis à jour prochainement.")
        
        intro.paragraph_format.space_after = Pt(12)
        doc.add_paragraph()
        
        # ✅ Commentaire évolution BRVM Composite sur 100 derniers jours
        if market_indicators and market_indicators.get('history_100d') is not None:
            df_hist = market_indicators['history_100d']
            # Les données sont déjà filtrées (brvm_composite NOT NULL > 0) et triées par id croissant
            
            if len(df_hist) >= 2:
                doc.add_heading('📊 Évolution de l\'indice BRVM Composite (100 derniers jours)', level=3)
                
                composite_first = float(df_hist.iloc[0]['brvm_composite'])
                composite_last  = float(df_hist.iloc[-1]['brvm_composite'])
                composite_max   = float(df_hist['brvm_composite'].max())
                composite_min   = float(df_hist['brvm_composite'].min())
                composite_evol  = ((composite_last - composite_first) / composite_first * 100)
                nb_jours        = len(df_hist)
                date_debut      = df_hist.iloc[0]['extraction_date']
                date_fin        = df_hist.iloc[-1]['extraction_date']
                
                # ── Stats descriptives BRVM Composite ──────────────────
                import numpy as np
                from scipy import stats as _sc
                comp_vals = df_hist['brvm_composite'].dropna().astype(float)
                c_mean   = float(comp_vals.mean())   if len(comp_vals) > 0 else 0
                c_median = float(comp_vals.median()) if len(comp_vals) > 0 else 0
                try:
                    c_mode = float(comp_vals.mode().iloc[0])
                except Exception:
                    c_mode = c_mean
                c_std    = float(comp_vals.std())    if len(comp_vals) > 1 else 0
                c_cv     = (c_std / c_mean * 100)    if c_mean > 0 else 0
                try:
                    c_kurt = float(_sc.kurtosis(comp_vals, fisher=True))
                    c_skew = float(_sc.skew(comp_vals))
                except Exception:
                    c_kurt = 0.0; c_skew = 0.0
                c_kurt_lbl = ("leptokurtique — risque de mouvements brusques" if c_kurt > 1
                              else "platykurtique — fluctuations régulières" if c_kurt < -1
                              else "distribution normale")
                c_skew_lbl = ("asymétrie positive — plus de jours sous la moyenne" if c_skew > 0.3
                              else "asymétrie négative — plus de jours au-dessus de la moyenne" if c_skew < -0.3
                              else "distribution symétrique")
                # ── Génération du commentaire composite ────────────────────────
                p_composite = doc.add_paragraph()
                p_composite.add_run("BRVM Composite — ").bold = True
                p_composite.add_run(
                    f"Sur les {nb_jours} derniers jours de cotation "
                    f"({date_debut} au {date_fin}), l'indice BRVM Composite a évolué de "
                )
                perf_run = p_composite.add_run(f"{composite_evol:+.2f}%")
                perf_run.bold = True
                perf_run.font.color.rgb = RGBColor(0, 128, 0) if composite_evol >= 0 else RGBColor(192, 0, 0)
                p_composite.add_run(
                    f", passant de {composite_first:.2f} pts à {composite_last:.2f} pts. "
                    f"Le plus haut atteint est {composite_max:.2f} pts et le plus bas {composite_min:.2f} pts. "
                    f"Sur l'ensemble de la période, la moyenne de l'indice s'établit à {c_mean:.2f} pts, "
                    f"la médiane à {c_median:.2f} pts et le mode à {c_mode:.2f} pts"
                )
                # Interpréter écart moyenne/médiane
                ecart_mm = abs(c_mean - c_median)
                if ecart_mm > c_std * 0.3:
                    p_composite.add_run(
                        f" — l'écart entre moyenne et médiane ({ecart_mm:.2f} pts) révèle une distribution asymétrique "
                        f"({c_skew_lbl}). "
                    )
                else:
                    p_composite.add_run(f" ({c_skew_lbl}). ")
                p_composite.add_run(
                    f"L'écart-type de {c_std:.2f} pts (CV={c_cv:.1f}%) traduit une "
                    f"{'forte' if c_cv > 15 else 'modérée' if c_cv > 7 else 'faible'} dispersion des cotations. "
                    f"Le kurtosis de {c_kurt:.2f} indique une distribution {c_kurt_lbl}. "
                )
                if composite_evol > 2:
                    p_composite.add_run("L'indice affiche une tendance haussière sur la période, témoignant d'un regain de confiance des investisseurs.")
                elif composite_evol < -2:
                    p_composite.add_run("L'indice accuse une tendance baissière sur la période, reflétant une pression vendeuse sur le marché.")
                else:
                    p_composite.add_run("L'indice évolue dans une phase de consolidation, sans tendance directrice nette sur la période.")

                # ── Graphique BRVM Composite (courbe seule) ──────────────
                buf_comp, buf_cap = self._generate_composite_chart(df_hist)
                if buf_comp:
                    try:
                        doc.add_picture(buf_comp, width=Inches(6.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as _ce:
                        logging.warning(f"⚠️  Insertion graphique composite: {_ce}")
                doc.add_paragraph()

                # Commentaire capitalisation 100j
                cap_vals = df_hist['capitalisation_globale'].dropna()
                cap_vals = cap_vals[cap_vals > 0]
                if len(cap_vals) >= 2:
                    doc.add_paragraph()
                    p_cap = doc.add_paragraph()
                    p_cap.add_run("Capitalisation globale — ").bold = True
                    cap_first = float(cap_vals.iloc[0])
                    cap_last  = float(cap_vals.iloc[-1])
                    cap_max   = float(cap_vals.max())
                    cap_min   = float(cap_vals.min())
                    cap_evol  = ((cap_last - cap_first) / cap_first * 100)
                    div = 1e9

                    def fmt_mds(val):
                        """Formate une valeur en milliards avec séparateur français (3 décimales)"""
                        mds = val / div
                        entier = int(mds)
                        dec = round((mds - entier) * 1000)
                        return f"{entier:,}".replace(",", " ") + f",{dec:03d}"

                    # ── Stats descriptives capitalisation ───────────────
                    cap_mean   = float(cap_vals.mean())   if len(cap_vals) > 0 else 0
                    cap_median = float(cap_vals.median()) if len(cap_vals) > 0 else 0
                    try:
                        cap_mode = float(cap_vals.mode().iloc[0])
                    except Exception:
                        cap_mode = cap_mean
                    cap_std  = float(cap_vals.std()) if len(cap_vals) > 1 else 0
                    cap_cv   = (cap_std / cap_mean * 100) if cap_mean > 0 else 0
                    try:
                        from scipy import stats as _sc2
                        cap_kurt = float(_sc2.kurtosis(cap_vals, fisher=True))
                        cap_skew = float(_sc2.skew(cap_vals))
                    except Exception:
                        cap_kurt = 0.0; cap_skew = 0.0
                    cap_kurt_lbl = ("leptokurtique — risque de pics de valorisation" if cap_kurt > 1
                                    else "platykurtique — valorisation régulière" if cap_kurt < -1
                                    else "distribution normale")
                    cap_skew_lbl = ("asymétrie positive" if cap_skew > 0.3
                                    else "asymétrie négative" if cap_skew < -0.3
                                    else "distribution symétrique")
                    # ── Commentaire ─────────────────────────────────────────────
                    p_cap.add_run("La capitalisation boursière totale a évolué de ")
                    cap_run = p_cap.add_run(f"{cap_evol:+.2f}%")
                    cap_run.bold = True
                    cap_run.font.color.rgb = RGBColor(0, 128, 0) if cap_evol >= 0 else RGBColor(192, 0, 0)
                    p_cap.add_run(
                        f" sur la période, passant de {fmt_mds(cap_first)} Mds FCFA à {fmt_mds(cap_last)} Mds FCFA. "
                        f"Le pic de capitalisation observé est de {fmt_mds(cap_max)} Mds FCFA "
                        f"et le plancher de {fmt_mds(cap_min)} Mds FCFA. "
                        f"La moyenne de la capitalisation sur la période s'établit à {fmt_mds(cap_mean)} Mds FCFA, "
                        f"la médiane à {fmt_mds(cap_median)} Mds FCFA et le mode à {fmt_mds(cap_mode)} Mds FCFA"
                    )
                    ecart_cap = abs(cap_mean - cap_median)
                    if ecart_cap > cap_std * 0.3:
                        p_cap.add_run(
                            f" — l'écart moyenne/médiane ({fmt_mds(ecart_cap)} Mds) signale une distribution {cap_skew_lbl}. "
                        )
                    else:
                        p_cap.add_run(f" ({cap_skew_lbl}). ")
                    p_cap.add_run(
                        f"L'écart-type de {fmt_mds(cap_std)} Mds FCFA (CV={cap_cv:.1f}%) reflète une "
                        f"{'forte' if cap_cv > 15 else 'modérée' if cap_cv > 7 else 'faible'} variabilité "
                        f"de la valorisation du marché. "
                        f"Le kurtosis de {cap_kurt:.2f} indique une distribution {cap_kurt_lbl}."
                    )
                    if cap_evol > 3:
                        p_cap.add_run(" Le marché a globalement capté davantage de richesse sur la période, signe d'une confiance accrue des investisseurs.")
                    elif cap_evol < -3:
                        p_cap.add_run(" La contraction de la capitalisation traduit une sortie de capitaux et une pression baissière structurelle.")
                    else:
                        p_cap.add_run(" La stabilité relative de la capitalisation reflète un marché en phase de consolidation.")

                # ── Graphique capitalisation (séparé) ───────────────────
                if buf_cap:
                    try:
                        doc.add_picture(buf_cap, width=Inches(6.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as _ce2:
                        logging.warning(f"⚠️  Insertion graphique capitalisation: {_ce2}")
                doc.add_paragraph()
        
        # ========== TOP 10 ACHATS ==========
        doc.add_heading('📈 TOP 10 DES OPPORTUNITÉS D\'ACHAT', level=2)
        
        # ✅ TOP 10 achat = uniquement ACHAT FORT (score 5) et ACHAT (score 4), triés par score décroissant
        sorted_buy = sorted(
            [(symbol, data) for symbol, data in all_company_data.items()
             if data.get('recommendation_score', 0) >= 4],  # ACHAT ou ACHAT FORT uniquement
            key=lambda x: x[1].get('recommendation_score', 0),
            reverse=True
        )[:10]
        
        # Si moins de 10 ACHAT/ACHAT FORT, compléter avec CONSERVER trié par score
        if len(sorted_buy) < 10:
            conserver_list = sorted(
                [(symbol, data) for symbol, data in all_company_data.items()
                 if data.get('recommendation_score', 0) == 3
                 and symbol not in [s for s, _ in sorted_buy]],
                key=lambda x: x[1].get('recommendation_score', 0),
                reverse=True
            )
            sorted_buy = sorted_buy + conserver_list[:10 - len(sorted_buy)]
        
        top_10_list = []
        for idx, (symbol, data) in enumerate(sorted_buy, 1):
            p = doc.add_paragraph(style='List Number')
            company_name = data.get('company_name', 'N/A')
            p.add_run(f"{symbol} - {company_name}").bold = True
            p.add_run(f" | Prix: {data.get('current_price', 0):.0f} FCFA | ")
            
            rec_run = p.add_run(f"{data.get('recommendation', 'N/A')}")
            rec_run.font.color.rgb = RGBColor(0, 128, 0)
            rec_run.bold = True
            
            p.add_run(f" | Confiance: {data.get('confidence_level', 'N/A')} | Risque: {data.get('risk_level', 'N/A')}")
            
            top_10_list.append({
                'symbol': symbol,
                'name': company_name,
                'price': float(data.get('current_price', 0)),
                'recommendation': data.get('recommendation'),
                'confidence': data.get('confidence_level'),
                'risk': data.get('risk_level')
            })
        
        doc.add_paragraph()
        
        # ========== FLOP 10 VENTES ==========
        doc.add_heading('📉 TOP 10 DES ACTIONS À ÉVITER', level=2)
        doc.add_paragraph(
            "Ces actions présentent des signaux de vente ou de dégradation fondamentale. "
            "Il est conseillé de les céder immédiatement ou d'éviter d'y investir."
        )
        
        # ✅ TOP 10 à éviter = uniquement VENTE FORTE (score 1) et VENTE (score 2), pires en premier
        sorted_sell = sorted(
            [(symbol, data) for symbol, data in all_company_data.items()
             if data.get('recommendation_score', 3) <= 2],  # VENTE ou VENTE FORTE uniquement
            key=lambda x: x[1].get('recommendation_score', 3)
        )[:10]
        
        # Si moins de 10 VENTE/VENTE FORTE, compléter avec les CONSERVER les plus faibles (risk_score le plus élevé)
        if len(sorted_sell) < 10:
            conserver_risky = sorted(
                [(symbol, data) for symbol, data in all_company_data.items()
                 if data.get('recommendation_score', 3) == 3
                 and symbol not in [s for s, _ in sorted_sell]],
                key=lambda x: x[1].get('risk_score', 0),
                reverse=True
            )
            sorted_sell = sorted_sell + conserver_risky[:10 - len(sorted_sell)]
        
        flop_10_list = []
        if sorted_sell:
            for idx, (symbol, data) in enumerate(sorted_sell, 1):
                p = doc.add_paragraph(style='List Number')
                company_name = data.get('company_name', 'N/A')
                p.add_run(f"{symbol} - {company_name}").bold = True
                p.add_run(f" | Prix: {data.get('current_price', 0):.0f} FCFA | ")
                
                rec_run = p.add_run(f"{data.get('recommendation', 'N/A')}")
                rec_run.font.color.rgb = RGBColor(192, 0, 0)
                rec_run.bold = True
                
                p.add_run(f" | Confiance: {data.get('confidence_level', 'N/A')} | Risque: {data.get('risk_level', 'N/A')}")
                
                flop_10_list.append({
                    'symbol': symbol,
                    'name': company_name,
                    'price': float(data.get('current_price', 0)),
                    'recommendation': data.get('recommendation'),
                    'confidence': data.get('confidence_level'),
                    'risk': data.get('risk_level')
                })
        else:
            p_empty = doc.add_paragraph()
            p_empty.add_run("✅ Aucune action à éviter ce jour.").italic = True
            p_empty.add_run(
                " L'ensemble des 47 sociétés analysées présentent des signaux neutres ou positifs. "
                "Aucun signal de vente fort n'a été détecté par l'analyse Multi-AI."
            )
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # ========== 1. ANALYSE PAR SECTEUR ==========
        doc.add_heading('📊 ANALYSE PAR SECTEUR', level=1)
        
        sector_analysis = self._calculate_sector_analysis(all_company_data)
        
        doc.add_paragraph(
            "Cette section présente une analyse comparative de tous les secteurs représentés à la BRVM, "
            "incluant la performance moyenne, le sentiment général du marché et le niveau de risque moyen."
        )
        doc.add_paragraph()
        
        # Récupérer la perf BRVM Composite pour comparaison relative
        _brvm_perf = 0.0
        try:
            _mi = market_indicators_pre if 'market_indicators_pre' in dir() else self._get_market_indicators()
            if _mi and _mi.get('history_100d') is not None:
                _h = _mi['history_100d']
                if len(_h) >= 2:
                    _first = float(_h.iloc[0]['brvm_composite'])
                    _last  = float(_h.iloc[-1]['brvm_composite'])
                    _brvm_perf = ((_last - _first) / _first * 100) if _first else 0
        except Exception:
            _brvm_perf = 0.0

        for sector, stats in sorted(sector_analysis.items(), key=lambda x: x[1]['performance_moyenne'], reverse=True):
            doc.add_heading(f"Secteur: {sector}", level=3)

            p = doc.add_paragraph()
            p.add_run("Nombre de sociétés: ").bold = True
            p.add_run(f"{stats['nb_societes']}  |  ")

            # Performance vs BRVM Composite
            perf = stats['performance_moyenne']
            vs   = perf - _brvm_perf
            p.add_run("Performance moyenne (100j): ").bold = True
            perf_run = p.add_run(f"{perf:.2f}%")
            perf_run.font.color.rgb = RGBColor(0,128,0) if perf >= 0 else RGBColor(192,0,0)
            p.add_run("  |  ")
            p.add_run("vs BRVM Composite: ").bold = True
            vs_run = p.add_run(f"{'+'if vs>=0 else ''}{vs:.2f}%  "
                               + ("🟢 Surperformance" if vs > 0 else "🔴 Sous-performance"))
            vs_run.font.color.rgb = RGBColor(0,128,0) if vs >= 0 else RGBColor(192,0,0)
            vs_run.bold = True
            p.add_run("\n")

            p.add_run("Sentiment général: ").bold = True
            p.add_run(f"{stats['sentiment_general']}  |  ")
            p.add_run("Risque moyen: ").bold = True
            p.add_run(f"{stats['risque_moyen']}  |  ")
            p.add_run("Prix moyen: ").bold = True
            p.add_run(f"{stats['prix_moyen']:.0f} FCFA\n")

            p.add_run("\nSociétés: ").bold = True
            p.add_run(", ".join(stats['societes']))

            doc.add_paragraph()
        
        doc.add_page_break()
        
        # ========== 2. MATRICE DE CONVERGENCE ==========
        doc.add_heading('🔄 MATRICE DE CONVERGENCE DES SIGNAUX', level=1)
        
        doc.add_paragraph(
            "Cette matrice segmente les sociétés cotées selon la combinaison de leur signal technique "
            "(majorité des 5 indicateurs) et de leur signal fondamental (recommandation finale de l'IA). "
            "Les cellules surlignées en vert indiquent une convergence favorable, en rouge une convergence défavorable."
        )
        doc.add_paragraph()
        
        # ✅ Construction de la matrice 3x3 à partir de technical_decision × fundamental_decision
        matrix_3x3 = {
            ('ACHAT',  'ACHAT'):  [],
            ('ACHAT',  'NEUTRE'): [],
            ('ACHAT',  'VENTE'):  [],
            ('NEUTRE', 'ACHAT'):  [],
            ('NEUTRE', 'NEUTRE'): [],
            ('NEUTRE', 'VENTE'):  [],
            ('VENTE',  'ACHAT'):  [],
            ('VENTE',  'NEUTRE'): [],
            ('VENTE',  'VENTE'):  [],
        }
        
        for symbol, data in all_company_data.items():
            td = (data.get('technical_decision') or 'NEUTRE').upper()
            fd = (data.get('fundamental_decision') or 'NEUTRE').upper()
            # Normaliser
            if 'ACHAT' in td:
                td = 'ACHAT'
            elif 'VENTE' in td:
                td = 'VENTE'
            else:
                td = 'NEUTRE'
            if 'ACHAT' in fd:
                fd = 'ACHAT'
            elif 'VENTE' in fd:
                fd = 'VENTE'
            else:
                fd = 'NEUTRE'
            key = (td, fd)
            if key in matrix_3x3:
                matrix_3x3[key].append(symbol)
        
        # Construire le tableau Word 4×4 (en-têtes inclus)
        tbl = doc.add_table(rows=4, cols=4)
        tbl.style = 'Table Grid'
        
        # Couleurs
        GREEN_STRONG = '00B050'   # convergence achat
        GREEN_LIGHT  = 'C6EFCE'
        RED_STRONG   = 'FF0000'   # convergence vente
        RED_LIGHT    = 'FFC7CE'
        ORANGE       = 'FFEB9C'   # divergence tech/fond
        GREY_HDR     = 'D9D9D9'
        WHITE        = 'FFFFFF'
        
        def _set_cell_bg(cell, hex_color):
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), hex_color)
            shd.set(qn('w:val'), 'clear')
            cell._element.get_or_add_tcPr().append(shd)
        
        def _cell_text(cell, text, bold=False, font_size=9, color=None):
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        
        headers_col = ['', 'Fond. ACHAT', 'Fond. NEUTRE', 'Fond. VENTE']
        headers_row = ['Tech. ACHAT', 'Tech. NEUTRE', 'Tech. VENTE']
        
        # Couleurs de fond par cellule (td, fd)
        cell_colors = {
            ('ACHAT',  'ACHAT'):  GREEN_STRONG,
            ('ACHAT',  'NEUTRE'): GREEN_LIGHT,
            ('ACHAT',  'VENTE'):  ORANGE,
            ('NEUTRE', 'ACHAT'):  GREEN_LIGHT,
            ('NEUTRE', 'NEUTRE'): WHITE,
            ('NEUTRE', 'VENTE'):  RED_LIGHT,
            ('VENTE',  'ACHAT'):  ORANGE,
            ('VENTE',  'NEUTRE'): RED_LIGHT,
            ('VENTE',  'VENTE'):  RED_STRONG,
        }
        
        # Ligne d'en-tête (ligne 0)
        for col_idx, hdr in enumerate(headers_col):
            cell = tbl.rows[0].cells[col_idx]
            _set_cell_bg(cell, GREY_HDR)
            _cell_text(cell, hdr, bold=True, font_size=9)
        
        # Lignes 1-3
        tech_order = ['ACHAT', 'NEUTRE', 'VENTE']
        fund_order = ['ACHAT', 'NEUTRE', 'VENTE']
        
        for row_idx, td in enumerate(tech_order):
            row = tbl.rows[row_idx + 1]
            # Cellule d'en-tête de ligne
            _set_cell_bg(row.cells[0], GREY_HDR)
            _cell_text(row.cells[0], headers_row[row_idx], bold=True, font_size=9)
            
            for col_idx, fd in enumerate(fund_order):
                cell = row.cells[col_idx + 1]
                color = cell_colors.get((td, fd), WHITE)
                _set_cell_bg(cell, color)
                symbols_list = matrix_3x3.get((td, fd), [])
                content = ', '.join(symbols_list) if symbols_list else '—'
                _cell_text(cell, content, font_size=8)
        
        doc.add_paragraph()
        
        # Légende
        legend_p = doc.add_paragraph()
        legend_p.add_run("Légende : ").bold = True
        legend_p.add_run("🟩 Vert foncé = Double achat (signal fort) | 🟩 Vert clair = Signal partiellement haussier | "
                         "🟧 Orange = Divergence tech/fondamental | 🟥 Rouge clair = Signal partiellement baissier | "
                         "🟥 Rouge foncé = Double vente (éviter) | ⬜ Blanc = Neutre")
        legend_p.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()
        
        doc.add_page_break()
        
        # ========== 3. ANALYSE DE LIQUIDITÉ ==========
        doc.add_heading('💧 ANALYSE DE LIQUIDITÉ', level=1)
        
        liquidity_analysis = self._calculate_liquidity_analysis(all_company_data)
        
        doc.add_paragraph(
            "Cette analyse compare les volumes moyens échangés pour identifier les titres les plus liquides "
            "(faciles à acheter/vendre) et ceux présentant un risque de liquidité."
        )
        doc.add_paragraph()
        
        # Haute liquidité
        doc.add_heading('🟢 Titres à Haute Liquidité (Top 20%)', level=3)
        high_liq_data = []
        for item in liquidity_analysis['high_liquidity']:
            high_liq_data.append([
                f"{item['symbol']} ({item['company_name']})",
                f"{item['avg_volume']:.0f}",
                f"{item['avg_value']:.0f} FCFA",
                f"{item['performance']:.2f}%",
                item['recommendation']
            ])
        
        if high_liq_data:
            self._add_table_with_shading(doc, high_liq_data, 
                                        ['Société', 'Vol. Moy.', 'Valeur Moy.', 'Perf. 100j', 'Recom.'])
        else:
            doc.add_paragraph("ℹ️ Données insuffisantes pour calculer la liquidité (besoin de 30 jours minimum).")
        doc.add_paragraph()
        
        # Faible liquidité
        doc.add_heading('🔴 Titres à Faible Liquidité (Bottom 40%) - RISQUE ÉLEVÉ', level=3)
        low_liq_data = []
        for item in liquidity_analysis['low_liquidity'][:10]:  # Top 10 des moins liquides
            low_liq_data.append([
                f"{item['symbol']} ({item['company_name']})",
                f"{item['avg_volume']:.0f}",
                f"{item['avg_value']:.0f} FCFA",
                f"{item['performance']:.2f}%",
                item['recommendation']
            ])
        
        if low_liq_data:
            self._add_table_with_shading(doc, low_liq_data, 
                                        ['Société', 'Vol. Moy.', 'Valeur Moy.', 'Perf. 100j', 'Recom.'])
        else:
            doc.add_paragraph("ℹ️ Données insuffisantes pour identifier les titres à faible liquidité.")
        
        doc.add_paragraph()
        p_warning = doc.add_paragraph()
        p_warning.add_run("⚠️ ATTENTION: ").bold = True
        p_warning.add_run(
            "Les titres à faible liquidité présentent un risque de ne pas pouvoir sortir facilement "
            "de sa position. À privilégier uniquement pour un investissement de long terme."
        )
        
        doc.add_page_break()
        
        # ========== 4. TOP 10 DIVERGENCES ==========
        doc.add_heading('⚠️ TOP 10 DES DIVERGENCES MAJEURES', level=1)
        
        top_divergences = self._calculate_top_divergences(all_company_data)
        
        doc.add_paragraph(
            "Cette section liste les sociétés présentant les écarts les plus importants "
            "entre les différents indicateurs techniques et fondamentaux."
        )
        doc.add_paragraph()
        
        for idx, div in enumerate(top_divergences, 1):
            doc.add_heading(f"{idx}. {div['symbol']} - {div['company_name']}", level=3)
            
            p = doc.add_paragraph()
            p.add_run(f"Score de divergence: ").bold = True
            p.add_run(f"{div['divergence_score']}/8\n")
            
            p.add_run(f"Description: ").bold = True
            p.add_run(f"{div['description']}\n")
            
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # ========== 5. MATRICE RISQUE vs HORIZON ==========
        doc.add_heading('📈 MATRICE RISQUE vs HORIZON DE PLACEMENT', level=1)
        
        risk_horizon_matrix = self._calculate_risk_horizon_matrix(all_company_data)
        
        doc.add_paragraph(
            "Cette matrice croise le niveau de risque avec l'horizon de placement recommandé "
            "pour faciliter la sélection de titres selon votre profil d'investisseur."
        )
        doc.add_paragraph()
        
        # Tableau récapitulatif
        matrix_data = [
            ['RISQUE / HORIZON', 'Court Terme', 'Moyen Terme', 'Long Terme'],
            [
                'Risque Faible',
                str(len(risk_horizon_matrix['faible_court'])),
                str(len(risk_horizon_matrix['faible_moyen'])),
                str(len(risk_horizon_matrix['faible_long']))
            ],
            [
                'Risque Moyen',
                str(len(risk_horizon_matrix['moyen_court'])),
                str(len(risk_horizon_matrix['moyen_moyen'])),
                str(len(risk_horizon_matrix['moyen_long']))
            ],
            [
                'Risque Élevé',
                str(len(risk_horizon_matrix['eleve_court'])),
                str(len(risk_horizon_matrix['eleve_moyen'])),
                str(len(risk_horizon_matrix['eleve_long']))
            ]
        ]
        
        self._add_table_with_shading(doc, matrix_data[1:], matrix_data[0])
        doc.add_paragraph()
        
        # Détail par catégorie
        categories = [
            ('Faible Risque - Court Terme', 'faible_court', '🟢'),
            ('Faible Risque - Moyen Terme', 'faible_moyen', '🟢'),
            ('Faible Risque - Long Terme', 'faible_long', '🟢'),
            ('Risque Moyen - Court Terme', 'moyen_court', '🟡'),
            ('Risque Moyen - Moyen Terme', 'moyen_moyen', '🟡'),
            ('Risque Moyen - Long Terme', 'moyen_long', '🟡'),
            ('Risque Élevé - Court Terme', 'eleve_court', '🔴'),
            ('Risque Élevé - Moyen Terme', 'eleve_moyen', '🔴'),
            ('Risque Élevé - Long Terme', 'eleve_long', '🔴'),
        ]
        
        for cat_name, cat_key, emoji in categories:
            if risk_horizon_matrix[cat_key]:
                doc.add_heading(f"{emoji} {cat_name}", level=3)
                for company in risk_horizon_matrix[cat_key]:
                    doc.add_paragraph(f"• {company}", style='List Bullet')
                doc.add_paragraph()
        
        doc.add_page_break()

        # ================================================================
        # ========== ANALYSE MACRO INTERNATIONALE (page dédiée) ==========
        # ================================================================
        doc.add_heading('🌍 ANALYSE MACRO — CONTEXTE INTERNATIONAL, AFRICAIN & UEMOA', level=1)
        doc.add_paragraph(
            "Analyse de l'environnement macro-économique mondial et régional susceptible "
            "d'impacter la BRVM et les sociétés cotées. Générée par IA à partir des "
            "actualités collectées via flux RSS Google News (BRVM, BCEAO, UEMOA, matières premières, marchés mondiaux)."
        ).runs[0].font.size = Pt(9)
        doc.add_paragraph()

        # ── Chargement des données macro et génération IA ────────────────────
        macro_news_data    = self._get_macro_news()
        macro_indicators   = market_indicators_pre if 'market_indicators_pre' in dir() else self._get_market_indicators()
        macro_result       = self._generate_macro_analysis(
            macro_news_data, all_company_data, macro_indicators
        )
        macro_text    = macro_result.get('analysis_text', '')
        macro_ai_prov = macro_result.get('ai_provider', '—')

        # ── Badge IA + compteurs ─────────────────────────────────────────────
        flat_df    = macro_news_data.get('_flat', pd.DataFrame())
        total_arts = len(flat_df) if isinstance(flat_df, pd.DataFrame) and not flat_df.empty else 0
        n_intl  = sum(len(macro_news_data.get(t,{}).get('international', pd.DataFrame())) for t in ['macroeconomique','politique','financiere'])
        n_afr   = sum(len(macro_news_data.get(t,{}).get('afrique', pd.DataFrame()))       for t in ['macroeconomique','politique','financiere'])
        n_aow   = sum(len(macro_news_data.get(t,{}).get('afrique_ouest', pd.DataFrame())) for t in ['macroeconomique','politique','financiere'])
        n_uemoa = sum(len(macro_news_data.get(t,{}).get('uemoa', pd.DataFrame()))         for t in ['macroeconomique','politique','financiere'])
        n_brvm2 = sum(len(macro_news_data.get(t,{}).get('brvm', pd.DataFrame()))          for t in ['macroeconomique','politique','financiere'])

        badge_p = doc.add_paragraph()
        badge_p.paragraph_format.space_before = Pt(2)
        badge_p.paragraph_format.space_after  = Pt(4)
        br1 = badge_p.add_run(f"🤖 Analyse générée par : {macro_ai_prov.upper()}  |  ")
        br1.font.size = Pt(8); br1.font.color.rgb = RGBColor(100,100,100)
        br2 = badge_p.add_run(
            f"Sources : {total_arts} actualités  |  "
            f"Mondial: {n_intl}  Afrique: {n_afr}  Afr.Ouest: {n_aow}  UEMOA: {n_uemoa}  BRVM: {n_brvm2}"
        )
        br2.font.size = Pt(8); br2.font.color.rgb = RGBColor(80,80,80)

        # ── Marqueurs de section ──────────────────────────────────────────────
        SECTION_STYLES = {
            '## TYPE_1':         ('📊 1. ACTUALITÉS MACRO-ÉCONOMIQUES',          RGBColor(0,70,127)),
            '## TYPE_2':         ('🏛️ 2. ACTUALITÉS POLITIQUES & GÉOPOLITIQUES', RGBColor(140,0,0)),
            '## TYPE_3':         ('💹 3. ACTUALITÉS FINANCIÈRES & MARCHÉS',       RGBColor(0,100,40)),
            '## SYNTHESE_FINALE':('🎯 SYNTHÈSE & RECOMMANDATION FINALE',          RGBColor(0,51,102)),
        }
        ZONE_STYLES = {
            '### PLAN MONDIAL':               ('🌐 Plan Mondial',                RGBColor(0,70,127),  True),
            '### PLAN AFRICAIN':              ('🌍 Plan Africain',               RGBColor(0,110,50),  True),
            '### PLAN AFRIQUE DE L':          ("🌍 Plan Afrique de l'Ouest",     RGBColor(20,90,40),  True),
            '### PLAN UEMOA':                 ('🏦 Plan UEMOA / Zone Franc',     RGBColor(140,70,0),  True),
            '### PLAN BRVM':                  ('📈 Plan BRVM',                   RGBColor(0,51,102),  True),
            '#### Impact estimé sur la BRVM': ('⚡ Impact estimé sur la BRVM',   RGBColor(80,0,80),   False),
        }

        if macro_text:
            current_impact = False
            buffer_lines   = []

            def _flush_buf(doc, buf, is_impact=False):
                for line in buf:
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith('- ') or line.startswith('• '):
                        txt = line.lstrip('-• ').strip()
                        bp  = doc.add_paragraph(style='List Bullet')
                        bp.paragraph_format.left_indent  = Pt(28 if is_impact else 18)
                        bp.paragraph_format.space_before = Pt(1)
                        bp.paragraph_format.space_after  = Pt(1)
                        rr  = bp.add_run(txt)
                        rr.font.size = Pt(9)
                        up  = txt.upper()
                        if any(k in up for k in ['POSITIF','HAUSSE','FAVORABLE','OPPORTUNITE']):
                            rr.font.color.rgb = RGBColor(0,128,0)
                        elif any(k in up for k in ['NEGATIF','BAISSE','RISQUE','PRESSION','DEFAVORABLE']):
                            rr.font.color.rgb = RGBColor(192,0,0)
                        elif is_impact:
                            rr.font.color.rgb = RGBColor(80,0,80)
                    elif any(k in line.upper() for k in ["NIVEAU D'ALERTE", "ALERTE :"]):
                        al_p = doc.add_paragraph()
                        al_p.paragraph_format.space_before = Pt(6)
                        al_r = al_p.add_run(line)
                        al_r.bold = True; al_r.font.size = Pt(11)
                        if 'VERT'    in line.upper(): al_r.font.color.rgb = RGBColor(0,128,0)
                        elif 'ROUGE' in line.upper(): al_r.font.color.rgb = RGBColor(192,0,0)
                        elif 'ORANGE'in line.upper(): al_r.font.color.rgb = RGBColor(200,80,0)
                    else:
                        pp = doc.add_paragraph()
                        pp.paragraph_format.space_before = Pt(1)
                        pp.paragraph_format.space_after  = Pt(2)
                        pp.paragraph_format.left_indent  = Pt(14 if is_impact else 0)
                        rr = pp.add_run(line)
                        rr.font.size = Pt(9)
                        if is_impact: rr.font.italic = True

            def _add_hdg(doc, label, color, level):
                sh = doc.add_heading(label, level=level)
                if sh.runs: sh.runs[0].font.color.rgb = color

            for raw_line in macro_text.split('\n'):
                s_line = raw_line.strip()
                matched = False

                for sec_key, (sec_title, sec_color) in SECTION_STYLES.items():
                    if s_line.startswith(sec_key):
                        _flush_buf(doc, buffer_lines, current_impact)
                        buffer_lines = []; current_impact = False
                        _add_hdg(doc, sec_title, sec_color, level=1)
                        matched = True; break
                if matched: continue

                for zone_key, (zone_title, zone_color, is_zone) in ZONE_STYLES.items():
                    if s_line.startswith(zone_key):
                        _flush_buf(doc, buffer_lines, current_impact)
                        buffer_lines = []; current_impact = not is_zone
                        _add_hdg(doc, zone_title, zone_color, level=2 if is_zone else 3)
                        matched = True; break
                if matched: continue

                buffer_lines.append(raw_line)

            _flush_buf(doc, buffer_lines, current_impact)

        else:
            doc.add_paragraph("ℹ️ Analyse macro non disponible.")

        # ── Tableaux sources par type et zone ────────────────────────────────
        doc.add_paragraph()
        doc.add_heading('📋 Sources utilisées pour cette analyse', level=3)

        TYPE_HDR = {
            'macroeconomique': '📊 Macro-économique',
            'politique':       '🏛️ Politique & Géopolitique',
            'financiere':      '💹 Financière & Marchés',
        }
        ZONE_HDR = {
            'international': '🌐 Mondial',
            'afrique':       '🌍 Afrique',
            'afrique_ouest': "🌍 Afrique de l'Ouest",
            'uemoa':         '🏦 UEMOA',
            'brvm':          '📈 BRVM',
        }

        for t_key, t_label in TYPE_HDR.items():
            t_data   = macro_news_data.get(t_key, {})
            has_data = any(
                isinstance(df, pd.DataFrame) and not df.empty
                for df in t_data.values()
            ) if isinstance(t_data, dict) else False
            if not has_data:
                continue
            doc.add_heading(t_label, level=4)
            for z_key, z_label in ZONE_HDR.items():
                df = t_data.get(z_key, pd.DataFrame()) if isinstance(t_data, dict) else pd.DataFrame()
                if not isinstance(df, pd.DataFrame) or df.empty:
                    continue
                zp = doc.add_paragraph()
                zr = zp.add_run(f"  {z_label}")
                zr.bold = True; zr.font.size = Pt(8.5); zr.font.color.rgb = RGBColor(60,60,60)
                tbl = doc.add_table(rows=1, cols=5)
                tbl.style = 'Light Grid Accent 1'
                for ci, hdr in enumerate(['Date', 'Titre', 'Sentiment', 'Impact BRVM', 'Source']):
                    cell = tbl.rows[0].cells[ci]
                    cell.text = hdr
                    run = cell.paragraphs[0].runs[0]
                    run.bold = True; run.font.size = Pt(7.5)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), '2E4057')
                    shd.set(qn('w:val'), 'clear')
                    cell._element.get_or_add_tcPr().append(shd)
                    run.font.color.rgb = RGBColor(255,255,255)
                for _, row in df.head(6).iterrows():
                    tr = tbl.add_row().cells
                    vals = [
                        str(row.get('mail_date',''))[:10],
                        str(row.get('titre') or row.get('mail_subject',''))[:70],
                        str(row.get('sentiment','') or '').capitalize(),
                        str(row.get('impact_brvm','') or '').capitalize(),
                        str(row.get('source_rss',''))[:30],
                    ]
                    for ci, v in enumerate(vals):
                        tr[ci].text = v
                        r2 = (tr[ci].paragraphs[0].runs[0]
                              if tr[ci].paragraphs[0].runs
                              else tr[ci].paragraphs[0].add_run(v))
                        r2.font.size = Pt(7)
                        if ci in (2, 3):
                            if 'Positif' in v:
                                r2.font.color.rgb = RGBColor(0,128,0)
                            elif 'Negatif' in v or 'Négatif' in v:
                                r2.font.color.rgb = RGBColor(192,0,0)
                doc.add_paragraph()

        doc.add_page_break()

        # ================================================================
        # ========== ACTUALITÉS DU MARCHÉ (page dédiée) ==================
        # ================================================================
        doc.add_heading('📰 ACTUALITÉS DU MARCHÉ BRVM', level=1)
        doc.add_paragraph(
            "Cette section regroupe les derniers documents officiels publiés par les "
            "sociétés cotées (AG, dividendes, convocations, résultats) ainsi que les "
            "alertes d'actualité financière collectées via Google Alerts."
        ).runs[0].font.size = Pt(9)
        doc.add_paragraph()

        # ── A) Documents officiels BRVM : AG / dividendes / convocations ─────
        doc.add_heading('📄 A — Documents officiels des sociétés cotées', level=2)
        brvm_actu_df = self._get_brvm_actualites()

        if not brvm_actu_df.empty:
            # Grouper par catégorie pour structurer la section
            CAT_ICONS = {
                'ag':           ('🏛️',  'Assemblées Générales'),
                'dividende':    ('💰',  'Dividendes'),
                'convocation':  ('📩',  'Convocations'),
                'résultats':    ('📊',  'Résultats financiers'),
                'résultat':     ('📊',  'Résultats financiers'),
                'rapport':      ('📋',  'Rapports annuels'),
                'communiqué':   ('📢',  'Communiqués de presse'),
                'avis':         ('ℹ️',  'Avis & informations'),
            }

            # Grouper les documents par catégorie
            cat_groups = {}
            for _, row in brvm_actu_df.iterrows():
                cat_raw = str(row.get('categorie') or row.get('type_document') or row.get('rapport_type') or 'Autre').lower().strip()
                # Normaliser la catégorie
                cat_key = 'autre'
                for key in CAT_ICONS:
                    if key in cat_raw:
                        cat_key = key
                        break
                cat_groups.setdefault(cat_key, []).append(row)

            # Afficher par catégorie
            for cat_key, rows in sorted(cat_groups.items()):
                icon, cat_label = CAT_ICONS.get(cat_key, ('📌', cat_key.capitalize()))
                doc.add_heading(f"{icon} {cat_label} ({len(rows)} document(s))", level=3)

                # Tableau compact pour cette catégorie
                tbl = doc.add_table(rows=1, cols=5)
                tbl.style = 'Light Grid Accent 1'
                hdrs = ['Date', 'Société', 'Titre', 'Impact', 'Résumé']
                hcells = tbl.rows[0].cells
                for ci, h in enumerate(hdrs):
                    hcells[ci].text = h
                    r = hcells[ci].paragraphs[0].runs[0]
                    r.bold = True
                    r.font.size = Pt(8)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), '1F4E79')
                    shd.set(qn('w:val'), 'clear')
                    hcells[ci]._element.get_or_add_tcPr().append(shd)
                    r.font.color.rgb = RGBColor(255,255,255)

                for row in rows[:10]:
                    tr = tbl.add_row().cells
                    # Date
                    date_val = str(row.get('date_doc') or row.get('date_publication') or '')[:10]
                    # Société
                    soc = str(row.get('societe_confirmee') or '')[:12]
                    # Titre
                    titre_doc = str(row.get('titre') or '')[:50]
                    # Impact
                    impact_raw = str(row.get('impact') or 'neutre').lower()
                    impact_map = {'positif':'🟢 +','negatif':'🔴 -','neutre':'⚪ ='}
                    impact_txt = impact_map.get(impact_raw, '⚪ =')
                    # Résumé court
                    res_court = str(row.get('resume') or '')[:120]

                    vals = [date_val, soc, titre_doc, impact_txt, res_court]
                    for ci, v in enumerate(vals):
                        tr[ci].text = v
                        run = tr[ci].paragraphs[0].runs[0] if tr[ci].paragraphs[0].runs else tr[ci].paragraphs[0].add_run(v)
                        run.font.size = Pt(7.5)
                        # Colorer impact
                        if ci == 3:
                            if '🟢' in v: run.font.color.rgb = RGBColor(0,128,0)
                            elif '🔴' in v: run.font.color.rgb = RGBColor(192,0,0)
                            run.bold = True

                    # Points clés en dessous si disponibles
                    pk = row.get('points_cles')
                    pts = []
                    if isinstance(pk, list): pts = [str(p) for p in pk[:3]]
                    elif isinstance(pk, str) and pk.strip():
                        try:
                            import json as _j
                            parsed = _j.loads(pk)
                            pts = [str(p) for p in parsed[:3]] if isinstance(parsed,list) else []
                        except Exception: pass

                    if pts:
                        pk_row = tbl.add_row().cells
                        # Fusionner les 5 colonnes pour les points clés
                        pk_para = pk_row[0].paragraphs[0]
                        pk_para.add_run("  → " + " | ".join(pts)).font.size = Pt(7)
                        pk_row[0].paragraphs[0].runs[-1].font.color.rgb = RGBColor(80,80,80)
                        # Étendre sur 5 cols (merge)
                        for merge_ci in range(1, 5):
                            pk_row[0].merge(pk_row[merge_ci])

                doc.add_paragraph()
        else:
            doc.add_paragraph("ℹ️ Aucun document officiel récent disponible dans la base de données.")

        doc.add_paragraph()

        # ── B) Alertes Google — actualités marché ────────────────────────────
        doc.add_heading('🔔 B — Alertes Google & actualités financières', level=2)
        alerts_df = self._get_google_alerts_events()

        if not alerts_df.empty:
            # Grouper par mot-clé / catégorie pour éviter la liste plate
            alert_groups = {}
            for _, alert_row in alerts_df.iterrows():
                kw = str(alert_row.get('alert_keyword') or alert_row.get('mot_cle') or
                         alert_row.get('categorie') or 'Marché BRVM').strip()
                alert_groups.setdefault(kw, []).append(alert_row)

            for kw, rows in alert_groups.items():
                grp_hdr = doc.add_paragraph()
                grp_hdr.paragraph_format.space_before = Pt(6)
                grp_hdr.paragraph_format.space_after  = Pt(2)
                rk = grp_hdr.add_run(f"🔑 {kw}  ({len(rows)} alerte(s))")
                rk.bold = True
                rk.font.size = Pt(10)
                rk.font.color.rgb = RGBColor(0,70,127)

                for alert_row in rows[:5]:
                    ap = doc.add_paragraph()
                    ap.paragraph_format.left_indent  = Pt(18)
                    ap.paragraph_format.space_before = Pt(2)
                    ap.paragraph_format.space_after  = Pt(2)

                    # Date
                    mail_date = alert_row.get('mail_date')
                    date_str  = mail_date.strftime('%d/%m/%Y') if hasattr(mail_date,'strftime') else str(mail_date)[:10]

                    # Sentiment
                    sent = str(alert_row.get('sentiment','') or '').lower()
                    s_color   = RGBColor(0,128,0) if 'positif' in sent else (RGBColor(192,0,0) if 'negatif' in sent else RGBColor(80,80,80))
                    sent_icon = '🟢' if 'positif' in sent else ('🔴' if 'negatif' in sent else '⚪')

                    # Titre : mail_subject > titre > resume[:80]
                    titre_a = str(
                        alert_row.get('mail_subject') or
                        alert_row.get('titre') or
                        str(alert_row.get('resume',''))[:80]
                    )[:120]

                    resume_a = str(alert_row.get('resume') or '')[:350]
                    cat_a    = str(alert_row.get('categorie') or alert_row.get('rapport_type') or '')
                    pert_a   = alert_row.get('pertinence')
                    url_a    = str(alert_row.get('source_url') or '')[:80]

                    # Ligne titre
                    rd = ap.add_run(f"[{date_str}] {sent_icon}  ")
                    rd.font.size = Pt(8.5)
                    rd.font.color.rgb = RGBColor(100,100,100)
                    rt = ap.add_run(titre_a)
                    rt.bold = True
                    rt.font.size = Pt(9)
                    rt.font.color.rgb = s_color

                    if cat_a:
                        rc = ap.add_run(f"  [{cat_a}]")
                        rc.font.size = Pt(7.5)
                        rc.font.color.rgb = RGBColor(120,120,120)

                    if pert_a and pd.notna(pert_a):
                        rp = ap.add_run(f"  ★ {int(pert_a)}/10")
                        rp.font.size = Pt(8)
                        rp.font.color.rgb = RGBColor(180,120,0)

                    # Résumé sur la ligne suivante
                    if resume_a and resume_a.strip() and resume_a[:80] not in titre_a:
                        ap_res = doc.add_paragraph()
                        ap_res.paragraph_format.left_indent  = Pt(30)
                        ap_res.paragraph_format.space_before = Pt(0)
                        ap_res.paragraph_format.space_after  = Pt(1)
                        rr = ap_res.add_run(resume_a)
                        rr.font.size = Pt(8.5)
                        rr.font.color.rgb = RGBColor(50,50,50)

                    # Points clés de l'alerte
                    pk_a = alert_row.get('points_cles')
                    pts_a = []
                    if isinstance(pk_a, list): pts_a = [str(p) for p in pk_a[:4]]
                    elif isinstance(pk_a, str) and pk_a.strip():
                        try:
                            import json as _jj
                            parsed = _jj.loads(pk_a)
                            pts_a = [str(p) for p in parsed[:4]] if isinstance(parsed,list) else []
                        except Exception: pass

                    if pts_a:
                        pk_para = doc.add_paragraph()
                        pk_para.paragraph_format.left_indent  = Pt(30)
                        pk_para.paragraph_format.space_before = Pt(0)
                        pk_para.paragraph_format.space_after  = Pt(2)
                        pk_run = pk_para.add_run("Points clés : " + "  •  ".join(pts_a))
                        pk_run.font.size = Pt(8)
                        pk_run.font.color.rgb = RGBColor(0,100,60)

                    # URL source
                    if url_a and url_a.startswith('http'):
                        url_para = doc.add_paragraph()
                        url_para.paragraph_format.left_indent  = Pt(30)
                        url_para.paragraph_format.space_before = Pt(0)
                        url_para.paragraph_format.space_after  = Pt(3)
                        ur = url_para.add_run(f"Source : {url_a}")
                        ur.font.size = Pt(7.5)
                        ur.font.color.rgb = RGBColor(0,100,180)

        else:
            doc.add_paragraph("ℹ️ Aucune alerte Google disponible pour cette période.")

        doc.add_page_break()

        # ========== CLASSEMENT 47 SOCIÉTÉS PAR SCORE COMPOSITE ==========
        doc.add_heading('🏆 CLASSEMENT DES SOCIÉTÉS — SCORE COMPOSITE /100', level=1)
        doc.add_paragraph(
            "Les 47 sociétés sont classées par score composite (Technique 30% + Fondamental 40% "
            "+ Risque 20% + Liquidité 10%). Ce tableau est l'outil de référence pour une décision rapide."
        )
        doc.add_paragraph()

        # Trier par score décroissant
        ranked = sorted(
            [(sym, d) for sym, d in all_company_data.items()],
            key=lambda x: x[1].get('investment_score', 0),
            reverse=True
        )

        score_tbl = doc.add_table(rows=1, cols=8)
        score_tbl.style = 'Light Grid Accent 1'
        hdrs = ['#', 'Symbole', 'Secteur', 'Prix (FCFA)', 'Score /100', 'Signal Tech', 'Signal Fond', 'Recommandation']
        hcells = score_tbl.rows[0].cells
        for ci, h in enumerate(hdrs):
            hcells[ci].text = h
            hcells[ci].paragraphs[0].runs[0].bold = True
            hcells[ci].paragraphs[0].runs[0].font.size = Pt(8.5)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1F4E79')
            shd.set(qn('w:val'), 'clear')
            hcells[ci]._element.get_or_add_tcPr().append(shd)
            hcells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

        for rank, (sym, d) in enumerate(ranked, 1):
            row_cells = score_tbl.add_row().cells
            sc = d.get('investment_score', 0)
            lbl = d.get('investment_label','—')
            rec = d.get('recommendation','—')
            td  = d.get('technical_decision','—')
            fd  = d.get('fundamental_decision','—')

            # Couleur ligne selon score
            if   sc >= 70: row_bg = 'C6EFCE'
            elif sc >= 55: row_bg = 'EBF5FB'
            elif sc >= 40: row_bg = 'FEFEFE'
            else:          row_bg = 'FCE4D6'

            vals = [
                str(rank),
                sym,
                str(d.get('sector','—') or '—')[:18],
                f"{d.get('current_price',0):,.0f}" if d.get('current_price') else '—',
                f"{sc}/100  ({lbl})",
                td,
                fd,
                rec,
            ]
            for ci, val in enumerate(vals):
                row_cells[ci].text = val
                run = row_cells[ci].paragraphs[0].runs[0] if row_cells[ci].paragraphs[0].runs else row_cells[ci].paragraphs[0].add_run(val)
                run.font.size = Pt(8)
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:fill'), row_bg)
                shd2.set(qn('w:val'), 'clear')
                row_cells[ci]._element.get_or_add_tcPr().append(shd2)
                # Coloration recommandation
                if ci == 7:
                    if   'ACHAT' in rec.upper(): run.font.color.rgb = RGBColor(0,128,0)
                    elif 'VENTE' in rec.upper(): run.font.color.rgb = RGBColor(192,0,0)
                # Score en gras
                if ci == 4:
                    run.bold = True

        doc.add_paragraph()
        doc.add_page_break()

        # ========== PORTEFEUILLES MODÈLES ==========
        doc.add_heading('💼 PORTEFEUILLES MODÈLES', level=1)
        doc.add_paragraph(
            "Trois portefeuilles construits automatiquement selon le score composite, "
            "le niveau de risque et la liquidité. Les pondérations sont indicatives et "
            "équipondérées. Position cash recommandée incluse."
        )
        doc.add_paragraph()

        portfolios = self._build_model_portfolios(all_company_data)
        port_configs = [
            ('defensif',  '🔵 Portefeuille DÉFENSIF',
             'Faible risque, haute liquidité, score ≥ 45. Adapté aux investisseurs prudents.'),
            ('equilibre', '🟢 Portefeuille ÉQUILIBRÉ',
             'Risque faible à moyen, score ≥ 50. Le meilleur rapport rendement/risque.'),
            ('offensif',  '🔴 Portefeuille OFFENSIF',
             'Meilleurs scores (≥ 60), tous niveaux de risque. Pour investisseurs avertis.'),
        ]
        PORT_COLORS = {'defensif':'DEEAF1','equilibre':'E2EFDA','offensif':'FCE4D6'}

        for port_key, port_title, port_desc in port_configs:
            entries_w, cash_pct = portfolios[port_key]
            doc.add_heading(port_title, level=3)
            dp = doc.add_paragraph(port_desc)
            dp.runs[0].font.size = Pt(9)
            dp.runs[0].font.italic = True

            if not entries_w:
                doc.add_paragraph("ℹ️ Aucun titre ne satisfait les critères ce jour.")
                doc.add_paragraph()
                continue

            # Tableau du portefeuille
            ptbl = doc.add_table(rows=1, cols=5)
            ptbl.style = 'Light Grid Accent 1'
            phdrs = ['Symbole', 'Société', 'Prix (FCFA)', 'Score /100', 'Poids %']
            phcells = ptbl.rows[0].cells
            for ci, ph in enumerate(phdrs):
                phcells[ci].text = ph
                phcells[ci].paragraphs[0].runs[0].bold = True
                phcells[ci].paragraphs[0].runs[0].font.size = Pt(8.5)
                pshd = OxmlElement('w:shd')
                pshd.set(qn('w:fill'), '2E74B5')
                pshd.set(qn('w:val'), 'clear')
                phcells[ci]._element.get_or_add_tcPr().append(pshd)
                phcells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

            for entry, weight in entries_w:
                pr = ptbl.add_row().cells
                vals = [
                    entry['sym'],
                    (entry['name'] or '')[:28],
                    f"{entry['price']:,.0f}" if entry['price'] else '—',
                    f"{entry['score']}/100",
                    f"{weight}%",
                ]
                for ci, vp in enumerate(vals):
                    pr[ci].text = vp
                    run = pr[ci].paragraphs[0].runs[0] if pr[ci].paragraphs[0].runs else pr[ci].paragraphs[0].add_run(vp)
                    run.font.size = Pt(8.5)
                    pshd2 = OxmlElement('w:shd')
                    pshd2.set(qn('w:fill'), PORT_COLORS.get(port_key,'FFFFFF'))
                    pshd2.set(qn('w:val'), 'clear')
                    pr[ci]._element.get_or_add_tcPr().append(pshd2)

            # Ligne cash
            cash_row = ptbl.add_row().cells
            cash_vals = ['CASH', 'Liquidités', '—', '—', f"{cash_pct}%"]
            for ci, vp in enumerate(cash_vals):
                cash_row[ci].text = vp
                run = cash_row[ci].paragraphs[0].runs[0] if cash_row[ci].paragraphs[0].runs else cash_row[ci].paragraphs[0].add_run(vp)
                run.font.size = Pt(8.5)
                run.font.italic = True
                run.font.color.rgb = RGBColor(80,80,80)

            doc.add_paragraph()

        doc.add_page_break()

        # ========== ALERTES DU JOUR ==========
        doc.add_heading('⚡ ALERTES DU JOUR', level=1)
        doc.add_paragraph(
            "Sociétés présentant des signaux extrêmes ou des incohérences majeures "
            "nécessitant une attention particulière avant toute décision."
        )
        doc.add_paragraph()

        alerts_list = []
        for sym, d in all_company_data.items():
            rsi_val  = d.get('rsi_value')
            price    = d.get('current_price') or 0
            high100  = d.get('highest_price_100d') or 0
            low100   = d.get('lowest_price_100d') or 0
            td       = str(d.get('technical_decision','')).upper()
            fd       = str(d.get('fundamental_decision','')).upper()
            rec      = str(d.get('recommendation','')).upper()

            alert_msgs = []

            # RSI extrêmes
            if rsi_val and pd.notna(rsi_val):
                rsi_f = float(rsi_val)
                if rsi_f > 70:
                    alert_msgs.append(f"🔴 RSI surachat ({rsi_f:.1f} > 70) — risque de correction")
                elif rsi_f < 30:
                    alert_msgs.append(f"🟢 RSI survente ({rsi_f:.1f} < 30) — opportunité potentielle")

            # Cours proche des bornes 100j
            if price and high100 and high100 > 0:
                dist_high = (high100 - price) / high100 * 100
                if dist_high < 3:
                    alert_msgs.append(f"⚠️ Cours à {dist_high:.1f}% du plus haut 100j ({high100:,.0f} FCFA)")
            if price and low100 and low100 > 0:
                dist_low = (price - low100) / low100 * 100
                if dist_low < 3:
                    alert_msgs.append(f"⚠️ Cours à {dist_low:.1f}% du plus bas 100j ({low100:,.0f} FCFA)")

            # Incohérence tech ↔ fond + recommandation ⚠️
            if '⚠️' in rec or (
                ('ACHAT' in td and 'VENTE' in fd) or
                ('VENTE' in td and 'ACHAT' in fd)
            ):
                alert_msgs.append(
                    f"🟠 Divergence tech({td}) / fond({fd}) — recommandation ajustée à {rec}"
                )

            if alert_msgs:
                alerts_list.append((sym, d.get('company_name',''), alert_msgs))

        if alerts_list:
            for sym, cname, msgs in sorted(alerts_list, key=lambda x: len(x[2]), reverse=True):
                ahdr = doc.add_paragraph()
                ahdr.paragraph_format.space_before = Pt(6)
                ahdr.add_run(f"{sym} — {cname}").bold = True
                ahdr.runs[-1].font.size = Pt(10)
                ahdr.runs[-1].font.color.rgb = RGBColor(0,80,160)
                for msg in msgs:
                    am = doc.add_paragraph(style='List Bullet')
                    am.paragraph_format.left_indent = Pt(20)
                    am.paragraph_format.space_before = Pt(1)
                    am.add_run(msg).font.size = Pt(9)
        else:
            doc.add_paragraph("✅ Aucune alerte majeure ce jour.")

        doc.add_page_break()

        # ══════════════════════════════════════════════════════════════════════
        # SYNTHÈSE RÉCAPITULATIVE DES PRÉDICTIONS IA
        # ══════════════════════════════════════════════════════════════════════
        doc.add_heading('🔮 SYNTHÈSE RÉCAPITULATIVE DES PRÉDICTIONS IA (J+1 → J+10)', level=1)
        pred_intro = doc.add_paragraph()
        pred_intro.add_run(
            "Cette section synthétise les prédictions des modèles GRU/LSTM/BiGRU pour les 47 sociétés. "
            "La variation J+10 est calculée entre le cours actuel et le cours prédit à 10 jours ouvrables. "
            "Le niveau de confiance agrégé reflète la proportion de jours avec confiance Élevée ou Moyenne "
            "dans la séquence de prédiction. Les prédictions sont données à titre indicatif — "
            "les marchés peu liquides comme la BRVM peuvent diverger des modèles statistiques."
        ).font.size = Pt(9)
        pred_intro.runs[0].font.italic = True
        pred_intro.runs[0].font.color.rgb = RGBColor(80, 80, 80)
        doc.add_paragraph()

        # ── Calcul des variations J+10 pour toutes les sociétés ──────────────
        pred_summary = []
        for sym, cdata in all_company_data.items():
            preds   = cdata.get('predictions_full', [])
            cur_p   = cdata.get('current_price') or 0
            cname   = cdata.get('company_name', sym)
            sector  = cdata.get('sector', 'Autre') or 'Autre'
            conf_lv = cdata.get('confidence_level', 'Moyen')

            if not preds or cur_p == 0:
                continue

            # Cours J+10 (dernier de la liste)
            valid_preds = [p for p in preds if p.get('price') is not None]
            if len(valid_preds) < 2:
                continue

            p_j1  = valid_preds[0]['price']   # J+1
            p_j10 = valid_preds[-1]['price']  # J+10

            var_j10 = ((p_j10 - cur_p) / cur_p * 100) if cur_p > 0 else 0
            var_j1  = ((p_j1  - cur_p) / cur_p * 100) if cur_p > 0 else 0

            # Confiance agrégée sur les 10 jours
            conf_scores = {'Élevée': 3, 'Moyenne': 2, 'Moyen': 2, 'Faible': 1, '': 1}
            conf_vals   = [conf_scores.get(p.get('confidence', ''), 1) for p in valid_preds]
            conf_avg    = sum(conf_vals) / len(conf_vals) if conf_vals else 1
            if conf_avg >= 2.5:   conf_label = 'Élevée'
            elif conf_avg >= 1.8: conf_label = 'Moyenne'
            else:                 conf_label = 'Faible'

            # Tendance de la trajectoire (J+1 → J+10)
            prix_seq = [p['price'] for p in valid_preds if p.get('price') is not None]
            hausses  = sum(1 for i in range(1, len(prix_seq)) if prix_seq[i] > prix_seq[i-1])
            baisses  = sum(1 for i in range(1, len(prix_seq)) if prix_seq[i] < prix_seq[i-1])
            if hausses > baisses * 2:   traj = "📈 Hausse continue"
            elif baisses > hausses * 2: traj = "📉 Baisse continue"
            elif hausses > baisses:     traj = "↗️ Tendance haussière"
            elif baisses > hausses:     traj = "↘️ Tendance baissière"
            else:                       traj = "➡️ Consolidation"

            pred_summary.append({
                'symbol':    sym,
                'name':      cname,
                'sector':    sector,
                'cur_price': cur_p,
                'p_j1':      p_j1,
                'p_j10':     p_j10,
                'var_j1':    var_j1,
                'var_j10':   var_j10,
                'conf':      conf_label,
                'traj':      traj,
                'n_preds':   len(valid_preds),
                'conf_avg':  conf_avg,
            })

        if not pred_summary:
            doc.add_paragraph("⚠️ Aucune prédiction disponible pour cette analyse.")
        else:
            # ── TENDANCE GÉNÉRALE ────────────────────────────────────────────
            doc.add_heading("📊 1. Tendance générale du marché", level=2)

            nb_hausse  = sum(1 for p in pred_summary if p['var_j10'] > 1)
            nb_baisse  = sum(1 for p in pred_summary if p['var_j10'] < -1)
            nb_neutre  = len(pred_summary) - nb_hausse - nb_baisse
            var_moyenne = sum(p['var_j10'] for p in pred_summary) / len(pred_summary)
            var_mediane = sorted(p['var_j10'] for p in pred_summary)[len(pred_summary)//2]

            # Confiance globale
            conf_count  = {'Élevée': 0, 'Moyenne': 0, 'Faible': 0}
            for p in pred_summary:
                conf_count[p['conf']] = conf_count.get(p['conf'], 0) + 1
            conf_dominante = max(conf_count, key=conf_count.get)
            pct_elevee = (conf_count.get('Élevée', 0) / len(pred_summary)) * 100
            pct_faible = (conf_count.get('Faible', 0) / len(pred_summary)) * 100

            # Signal marché
            if var_moyenne > 3 and nb_hausse > nb_baisse * 2:
                signal_marche  = "🟢 HAUSSIER"
                signal_color   = RGBColor(0, 128, 0)
                signal_bg      = 'C6EFCE'
                signal_detail  = (f"Les modèles anticipent une progression généralisée des cours. "
                                  f"La majorité des sociétés ({nb_hausse}/{len(pred_summary)}) sont "
                                  f"orientées à la hausse avec une variation moyenne de +{var_moyenne:.2f}%.")
            elif var_moyenne < -3 and nb_baisse > nb_hausse * 2:
                signal_marche  = "🔴 BAISSIER"
                signal_color   = RGBColor(180, 0, 0)
                signal_bg      = 'FFC7CE'
                signal_detail  = (f"Les modèles anticipent une correction généralisée. "
                                  f"{nb_baisse} sociétés sur {len(pred_summary)} sont orientées à la baisse "
                                  f"avec une variation moyenne de {var_moyenne:.2f}%.")
            elif var_moyenne > 1:
                signal_marche  = "🟡 LÉGÈREMENT HAUSSIER"
                signal_color   = RGBColor(150, 100, 0)
                signal_bg      = 'FFEB9C'
                signal_detail  = (f"Légère tendance haussière mais signal mixte : {nb_hausse} sociétés en hausse "
                                  f"vs {nb_baisse} en baisse. Variation moyenne : +{var_moyenne:.2f}%.")
            elif var_moyenne < -1:
                signal_marche  = "🟡 LÉGÈREMENT BAISSIER"
                signal_color   = RGBColor(150, 50, 0)
                signal_bg      = 'FFEB9C'
                signal_detail  = (f"Légère tendance baissière avec signal mixte : {nb_baisse} sociétés en baisse "
                                  f"vs {nb_hausse} en hausse. Variation moyenne : {var_moyenne:.2f}%.")
            else:
                signal_marche  = "⚪ NEUTRE / CONSOLIDATION"
                signal_color   = RGBColor(80, 80, 80)
                signal_bg      = 'F2F2F2'
                signal_detail  = (f"Les modèles n'anticipent pas de mouvement directionnel clair. "
                                  f"La variation moyenne est de {var_moyenne:+.2f}% avec {nb_neutre} sociétés "
                                  f"en zone neutre (±1%).")

            # Tableau synthèse générale
            tbl_tend = doc.add_table(rows=4, cols=2)
            tbl_tend.style = 'Table Grid'
            rows_tend = [
                ("Signal marché (J+10)",  signal_marche),
                ("Variation moyenne J+10", f"{var_moyenne:+.2f}%  |  Médiane : {var_mediane:+.2f}%"),
                ("Répartition",
                 f"📈 {nb_hausse} sociétés en hausse (>{'+1%'})   "
                 f"📉 {nb_baisse} en baisse (<{'-1%'})   "
                 f"➡️ {nb_neutre} neutres"),
                ("Confiance globale",
                 f"Élevée : {conf_count.get('Élevée',0)} sociétés ({pct_elevee:.0f}%)   "
                 f"Moyenne : {conf_count.get('Moyenne',0)}   "
                 f"Faible : {conf_count.get('Faible',0)} ({pct_faible:.0f}%)   "
                 f"→ Confiance dominante : {conf_dominante}"),
            ]
            for ri, (lbl, val) in enumerate(rows_tend):
                tr = tbl_tend.rows[ri].cells
                # Col label
                tr[0].text = lbl
                rl = tr[0].paragraphs[0].runs[0] if tr[0].paragraphs[0].runs else tr[0].paragraphs[0].add_run(lbl)
                rl.bold = True; rl.font.size = Pt(9)
                shd_l = OxmlElement('w:shd'); shd_l.set(qn('w:fill'), 'DEEAF1'); shd_l.set(qn('w:val'), 'clear')
                tr[0]._element.get_or_add_tcPr().append(shd_l)
                # Col valeur
                tr[1].text = val
                rv = tr[1].paragraphs[0].runs[0] if tr[1].paragraphs[0].runs else tr[1].paragraphs[0].add_run(val)
                rv.bold = True; rv.font.size = Pt(9)
                if ri == 0:
                    rv.font.color.rgb = signal_color
                    shd_v = OxmlElement('w:shd'); shd_v.set(qn('w:fill'), signal_bg); shd_v.set(qn('w:val'), 'clear')
                    tr[1]._element.get_or_add_tcPr().append(shd_v)

            # Largeurs
            try:
                for ci_t, col_t in enumerate(tbl_tend.columns):
                    for cell_t in col_t.cells:
                        cell_t.width = Cm(5.0) if ci_t == 0 else Cm(13.0)
            except: pass

            # Commentaire narratif du signal
            doc.add_paragraph()
            sig_p = doc.add_paragraph()
            sig_r = sig_p.add_run(f"📝 Analyse : {signal_detail}")
            sig_r.font.size = Pt(9); sig_r.font.italic = True
            sig_r.font.color.rgb = RGBColor(40, 40, 40)

            # Confiance globale commentaire
            if pct_elevee >= 50:
                conf_comment = (f"✅ La confiance est globalement ÉLEVÉE ({pct_elevee:.0f}% des sociétés) — "
                                f"les modèles ont été entraînés sur des données suffisamment régulières pour "
                                f"produire des prédictions fiables. Les résultats sont exploitables.")
            elif pct_faible >= 50:
                conf_comment = (f"⚠️ La confiance est globalement FAIBLE ({pct_faible:.0f}% des sociétés) — "
                                f"forte volatilité ou données insuffisantes dans les historiques. "
                                f"Les prédictions sont indicatives et doivent être croisées avec l'analyse technique.")
            else:
                conf_comment = (f"🔵 Confiance MIXTE — certains modèles sont robustes, d'autres sont limités "
                                f"par la faible liquidité ou la volatilité des cours sur la BRVM. "
                                f"Privilégier les sociétés avec confiance Élevée.")
            conf_p = doc.add_paragraph()
            conf_r = conf_p.add_run(f"🎯 Fiabilité : {conf_comment}")
            conf_r.font.size = Pt(9); conf_r.font.italic = True
            conf_r.font.color.rgb = RGBColor(40, 40, 40)
            doc.add_paragraph()

            # ── TOP 5 HAUSSES ────────────────────────────────────────────────
            doc.add_heading("📈 2. TOP 5 — Cours attendus en HAUSSE à J+10", level=2)
            top5_hausse = sorted(pred_summary, key=lambda x: x['var_j10'], reverse=True)[:5]

            note_h = doc.add_paragraph()
            note_h.add_run(
                "Classement par variation prédite J+10 (cours actuel → cours prédit dans 10 jours ouvrables). "
                "Sont indiqués : la variation attendue, la trajectoire sur 10 jours et le niveau de confiance du modèle."
            ).font.size = Pt(8.5)
            note_h.runs[0].font.italic = True
            doc.add_paragraph()

            # Tableau Top 5 Hausse
            tbl_h5 = doc.add_table(rows=1, cols=7)
            tbl_h5.style = 'Table Grid'
            hdrs_h5 = ['Rang', 'Symbole / Société', 'Secteur', 'Cours actuel',
                       'Cours J+10 prédit', 'Variation J+10', 'Confiance']
            hdr_h5 = tbl_h5.rows[0].cells
            for ci, ht in enumerate(hdrs_h5):
                hdr_h5[ci].text = ht
                rh = hdr_h5[ci].paragraphs[0].runs[0] if hdr_h5[ci].paragraphs[0].runs else hdr_h5[ci].paragraphs[0].add_run(ht)
                rh.bold = True; rh.font.size = Pt(8); rh.font.color.rgb = RGBColor(255,255,255)
                shd_hh = OxmlElement('w:shd'); shd_hh.set(qn('w:fill'), '375623'); shd_hh.set(qn('w:val'), 'clear')
                hdr_h5[ci]._element.get_or_add_tcPr().append(shd_hh)

            for rang, p in enumerate(top5_hausse, 1):
                tr5 = tbl_h5.add_row().cells
                vals5 = [
                    f"#{rang}",
                    f"{p['symbol']} — {p['name'][:35]}{'...' if len(p['name'])>35 else ''}",
                    p['sector'][:20],
                    f"{p['cur_price']:,.0f} FCFA",
                    f"{p['p_j10']:,.0f} FCFA",
                    f"+{p['var_j10']:.2f}%",
                    p['conf'],
                ]
                bg_conf = 'C6EFCE' if p['conf'] == 'Élevée' else 'FFEB9C' if p['conf'] == 'Moyenne' else 'FFC7CE'
                bgs = ['FFFFFF', 'C6EFCE', 'FFFFFF', 'FFFFFF', 'FFFFFF', 'C6EFCE', bg_conf]
                for ci, (val_s, bg_s) in enumerate(zip(vals5, bgs)):
                    tr5[ci].text = val_s
                    rc5 = tr5[ci].paragraphs[0].runs[0] if tr5[ci].paragraphs[0].runs else tr5[ci].paragraphs[0].add_run(val_s)
                    rc5.font.size = Pt(8)
                    if ci == 5: rc5.bold = True; rc5.font.color.rgb = RGBColor(0, 120, 0)
                    if ci == 0: rc5.bold = True
                    shd5 = OxmlElement('w:shd'); shd5.set(qn('w:fill'), bg_s); shd5.set(qn('w:val'), 'clear')
                    tr5[ci]._element.get_or_add_tcPr().append(shd5)
            try:
                widths_h5 = [Cm(1.2), Cm(5.0), Cm(3.0), Cm(2.5), Cm(2.5), Cm(2.0), Cm(1.8)]
                for ci5, col5 in enumerate(tbl_h5.columns):
                    for cell5 in col5.cells:
                        cell5.width = widths_h5[ci5]
            except: pass
            doc.add_paragraph()

            # Commentaire analytique Top 5 hausse
            for rang, p in enumerate(top5_hausse, 1):
                var_abs = abs(p['var_j10'])
                if var_abs > 20:
                    amp_comment = f"hausse très forte (+{p['var_j10']:.1f}%) — signal potentiellement excessif, à croiser avec l'analyse technique"
                elif var_abs > 10:
                    amp_comment = f"hausse significative (+{p['var_j10']:.1f}%) — potentiel de revalorisation notable"
                elif var_abs > 5:
                    amp_comment = f"hausse modérée (+{p['var_j10']:.1f}%) — signal raisonnable et crédible"
                else:
                    amp_comment = f"légère hausse (+{p['var_j10']:.1f}%) — progression graduelle attendue"

                conf_detail = {
                    'Élevée':  "Le modèle est bien calibré sur cet actif (faible MAPE, R² élevé). Résultat exploitable.",
                    'Moyenne': "Confiance partielle — les premiers jours sont fiables, les jours lointains moins précis.",
                    'Faible':  "⚠️ Confiance limitée — volatilité élevée ou données insuffisantes. Résultat indicatif uniquement.",
                }.get(p['conf'], "Confiance non déterminée.")

                comm_p = doc.add_paragraph(style='List Bullet')
                comm_r1 = comm_p.add_run(f"#{rang} {p['symbol']} ")
                comm_r1.bold = True; comm_r1.font.size = Pt(9); comm_r1.font.color.rgb = RGBColor(0, 100, 0)
                comm_r2 = comm_p.add_run(
                    f"({p['name'][:40]}) — {amp_comment}. "
                    f"Trajectoire : {p['traj']}. "
                    f"Cours actuel : {p['cur_price']:,.0f} FCFA → prédit J+10 : {p['p_j10']:,.0f} FCFA. "
                    f"{conf_detail}"
                )
                comm_r2.font.size = Pt(8.5)
            doc.add_paragraph()

            # ── TOP 5 BAISSES ────────────────────────────────────────────────
            doc.add_heading("📉 3. TOP 5 — Cours attendus en BAISSE à J+10", level=2)
            top5_baisse = sorted(pred_summary, key=lambda x: x['var_j10'])[:5]

            note_b = doc.add_paragraph()
            note_b.add_run(
                "Les sociétés présentant les prédictions les plus négatives à J+10. "
                "Une baisse prédite ne signifie pas nécessairement une recommandation de vente — "
                "elle doit être croisée avec les fondamentaux et la liquidité du titre."
            ).font.size = Pt(8.5)
            note_b.runs[0].font.italic = True
            doc.add_paragraph()

            # Tableau Top 5 Baisse
            tbl_b5 = doc.add_table(rows=1, cols=7)
            tbl_b5.style = 'Table Grid'
            hdr_b5 = tbl_b5.rows[0].cells
            hdrs_b5 = ['Rang', 'Symbole / Société', 'Secteur', 'Cours actuel',
                       'Cours J+10 prédit', 'Variation J+10', 'Confiance']
            for ci, ht in enumerate(hdrs_b5):
                hdr_b5[ci].text = ht
                rh = hdr_b5[ci].paragraphs[0].runs[0] if hdr_b5[ci].paragraphs[0].runs else hdr_b5[ci].paragraphs[0].add_run(ht)
                rh.bold = True; rh.font.size = Pt(8); rh.font.color.rgb = RGBColor(255,255,255)
                shd_hb = OxmlElement('w:shd'); shd_hb.set(qn('w:fill'), '843C0C'); shd_hb.set(qn('w:val'), 'clear')
                hdr_b5[ci]._element.get_or_add_tcPr().append(shd_hb)

            for rang, p in enumerate(top5_baisse, 1):
                tr5b = tbl_b5.add_row().cells
                vals5b = [
                    f"#{rang}",
                    f"{p['symbol']} — {p['name'][:35]}{'...' if len(p['name'])>35 else ''}",
                    p['sector'][:20],
                    f"{p['cur_price']:,.0f} FCFA",
                    f"{p['p_j10']:,.0f} FCFA",
                    f"{p['var_j10']:.2f}%",
                    p['conf'],
                ]
                bg_conf_b = 'C6EFCE' if p['conf'] == 'Élevée' else 'FFEB9C' if p['conf'] == 'Moyenne' else 'FFC7CE'
                bgs_b = ['FFFFFF', 'FFC7CE', 'FFFFFF', 'FFFFFF', 'FFFFFF', 'FFC7CE', bg_conf_b]
                for ci, (val_s, bg_s) in enumerate(zip(vals5b, bgs_b)):
                    tr5b[ci].text = val_s
                    rc5b = tr5b[ci].paragraphs[0].runs[0] if tr5b[ci].paragraphs[0].runs else tr5b[ci].paragraphs[0].add_run(val_s)
                    rc5b.font.size = Pt(8)
                    if ci == 5: rc5b.bold = True; rc5b.font.color.rgb = RGBColor(180, 0, 0)
                    if ci == 0: rc5b.bold = True
                    shd5b = OxmlElement('w:shd'); shd5b.set(qn('w:fill'), bg_s); shd5b.set(qn('w:val'), 'clear')
                    tr5b[ci]._element.get_or_add_tcPr().append(shd5b)
            try:
                for ci5b, col5b in enumerate(tbl_b5.columns):
                    for cell5b in col5b.cells:
                        cell5b.width = widths_h5[ci5b]
            except: pass
            doc.add_paragraph()

            # Commentaire analytique Top 5 baisse
            for rang, p in enumerate(top5_baisse, 1):
                var_abs = abs(p['var_j10'])
                if var_abs > 30:
                    amp_b = f"baisse très forte ({p['var_j10']:.1f}%) — amplitude potentiellement surestimée par le modèle"
                elif var_abs > 15:
                    amp_b = f"correction importante ({p['var_j10']:.1f}%) — signal de dégradation significative"
                elif var_abs > 5:
                    amp_b = f"baisse modérée ({p['var_j10']:.1f}%) — repli technique attendu"
                else:
                    amp_b = f"légère baisse ({p['var_j10']:.1f}%) — mouvement de faible amplitude"

                conf_detail_b = {
                    'Élevée':  "Modèle bien calibré — ce signal de baisse est à prendre au sérieux.",
                    'Moyenne': "Confiance partielle — surveiller l'évolution réelle des premiers jours.",
                    'Faible':  "⚠️ Confiance limitée — volatilité élevée. Ne pas agir sur ce seul signal.",
                }.get(p['conf'], "")

                comm_pb = doc.add_paragraph(style='List Bullet')
                comm_rb1 = comm_pb.add_run(f"#{rang} {p['symbol']} ")
                comm_rb1.bold = True; comm_rb1.font.size = Pt(9); comm_rb1.font.color.rgb = RGBColor(180, 0, 0)
                comm_rb2 = comm_pb.add_run(
                    f"({p['name'][:40]}) — {amp_b}. "
                    f"Trajectoire : {p['traj']}. "
                    f"Cours actuel : {p['cur_price']:,.0f} FCFA → prédit J+10 : {p['p_j10']:,.0f} FCFA. "
                    f"{conf_detail_b}"
                )
                comm_rb2.font.size = Pt(8.5)
            doc.add_paragraph()

            # ── TABLEAU COMPLET TOUTES SOCIÉTÉS ─────────────────────────────
            doc.add_heading("📋 4. Tableau complet — Toutes les sociétés", level=2)
            note_all = doc.add_paragraph()
            note_all.add_run(
                "Vue d'ensemble de toutes les sociétés avec prédictions disponibles, "
                "triées par variation J+10 décroissante."
            ).font.size = Pt(8.5)
            note_all.runs[0].font.italic = True
            doc.add_paragraph()

            tbl_all = doc.add_table(rows=1, cols=6)
            tbl_all.style = 'Table Grid'
            hdr_all = tbl_all.rows[0].cells
            hdrs_all = ['Symbole', 'Société', 'Cours actuel', 'Cours J+10', 'Var. J+10', 'Confiance']
            for ci, ht in enumerate(hdrs_all):
                hdr_all[ci].text = ht
                rh = hdr_all[ci].paragraphs[0].runs[0] if hdr_all[ci].paragraphs[0].runs else hdr_all[ci].paragraphs[0].add_run(ht)
                rh.bold = True; rh.font.size = Pt(8); rh.font.color.rgb = RGBColor(255,255,255)
                shd_ha = OxmlElement('w:shd'); shd_ha.set(qn('w:fill'), '1F4E79'); shd_ha.set(qn('w:val'), 'clear')
                hdr_all[ci]._element.get_or_add_tcPr().append(shd_ha)

            sorted_all = sorted(pred_summary, key=lambda x: x['var_j10'], reverse=True)
            for p in sorted_all:
                tra = tbl_all.add_row().cells
                var_str = f"{p['var_j10']:+.2f}%"
                vals_a = [
                    p['symbol'],
                    p['name'][:35] + ('...' if len(p['name'])>35 else ''),
                    f"{p['cur_price']:,.0f}",
                    f"{p['p_j10']:,.0f}",
                    var_str,
                    p['conf'],
                ]
                is_up   = p['var_j10'] > 1
                is_down = p['var_j10'] < -1
                bg_var  = 'C6EFCE' if is_up else 'FFC7CE' if is_down else 'FFEB9C'
                bg_conf_a = 'C6EFCE' if p['conf'] == 'Élevée' else 'FFEB9C' if p['conf'] == 'Moyenne' else 'FFC7CE'
                bgs_a = ['FFFFFF', 'FFFFFF', 'FFFFFF', 'FFFFFF', bg_var, bg_conf_a]
                for ci, (val_s, bg_s) in enumerate(zip(vals_a, bgs_a)):
                    tra[ci].text = val_s
                    rca = tra[ci].paragraphs[0].runs[0] if tra[ci].paragraphs[0].runs else tra[ci].paragraphs[0].add_run(val_s)
                    rca.font.size = Pt(8)
                    if ci == 4:
                        rca.bold = True
                        rca.font.color.rgb = RGBColor(0,120,0) if is_up else RGBColor(180,0,0) if is_down else RGBColor(100,100,0)
                    shd_a = OxmlElement('w:shd'); shd_a.set(qn('w:fill'), bg_s); shd_a.set(qn('w:val'), 'clear')
                    tra[ci]._element.get_or_add_tcPr().append(shd_a)

            try:
                widths_a = [Cm(1.8), Cm(6.0), Cm(2.5), Cm(2.5), Cm(2.0), Cm(2.0)]
                for cia, cola in enumerate(tbl_all.columns):
                    for cella in cola.cells:
                        cella.width = widths_a[cia]
            except: pass
            doc.add_paragraph()

        doc.add_page_break()

        # ══════════════════════════════════════════════════════════════════════
        # ANALYSE FINANCIÈRE COMPARATIVE PAR SECTEUR (brvm_donnees_financieres)
        # ══════════════════════════════════════════════════════════════════════
        doc.add_heading('💰 ANALYSE FINANCIÈRE COMPARATIVE PAR SECTEUR', level=1)
        intro_afcs = doc.add_paragraph()
        intro_afcs.add_run(
            "Cette section compare les performances financières des sociétés cotées sur la base des données "
            "structurées annuelles issues de la table brvm_donnees_financieres. Seules les sociétés disposant "
            "de données renseignées sont analysées. Pour chaque secteur, les meilleures et moins bonnes "
            "performances sont identifiées sur les indicateurs clés de rentabilité, structure et liquidité."
        ).font.size = Pt(9)
        intro_afcs.runs[0].font.italic = True
        intro_afcs.runs[0].font.color.rgb = RGBColor(80, 80, 80)
        doc.add_paragraph()

        # ── Charger toutes les données financières disponibles ─────────────────
        fin_all = {}
        for sym in all_company_data:
            fd = self._get_donnees_financieres(sym)
            if fd:
                fin_all[sym] = fd

        # ── Mapping secteur → symboles (depuis all_company_data) ──────────────
        sector_map = {}
        for sym, cdata in all_company_data.items():
            sec = cdata.get('sector', 'Autre')
            if not sec or sec.strip() == '': sec = 'Autre'
            sector_map.setdefault(sec, []).append(sym)

        # ── Helper : valeur numérique sûre ────────────────────────────────────
        def _safe(d, key):
            v = d.get(key)
            if v is None: return None
            try:
                f = float(v)
                return f if f != 0.0 else None
            except: return None

        def _fmt_mds(v):
            if v is None: return 'N/D'
            if abs(v) >= 1e9: return f"{v/1e9:.3f} Mds FCFA"
            if abs(v) >= 1e6: return f"{v/1e6:.2f} M FCFA"
            return f"{v:,.0f} FCFA"

        def _fmt_pct(v):
            if v is None: return 'N/D'
            return f"{v*100:.2f}%"

        # ── Indicateurs à comparer (label, clé, format, sens : +1=plus grand=mieux, -1=plus petit=mieux) ──
        KPI_DEFS = [
            # Rentabilité
            ("ROE",                    "roe",                      "pct",  +1),
            ("ROA",                    "roa",                      "pct",  +1),
            ("Marge nette",            "marge_nette",              "pct",  +1),
            ("Résultat net",           "resultat_net",             "mds",  +1),
            ("Taux croissance CA/PNB", "taux_croissance_ca",       "pct",  +1),
            # Structure
            ("Autonomie financière",   "autonomie_financiere",     "pct",  +1),
            ("Ratio endettement",      "ratio_endettement",        "pct",  -1),
            ("Solvabilité générale",   "solvabilite_generale",     "pct",  +1),
            # Liquidité
            ("Liquidité générale",     "liquidite_generale",       "pct",  +1),
            ("Free Cash Flow",         "free_cash_flow",           "mds",  +1),
            # Efficacité
            ("Coeff. exploitation",    "coefficient_exploitation",  "pct",  -1),  # banques uniquement
            ("Marge opérationnelle",   "marge_operationnelle",     "pct",  +1),
            # Délais
            ("Délai clients",          "delai_clients",            "jours", -1),
            ("Durée stockage",         "duree_stockage",           "jours", -1),
        ]

        # Palettes couleur
        CLR_POS_HDR = '1F4E79'   # bleu foncé
        CLR_POS     = 'C6EFCE'   # vert
        CLR_NEG     = 'FFC7CE'   # rouge
        CLR_MID     = 'FFEB9C'   # orange
        CLR_SEC_HDR = 'D9E1F2'   # bleu clair (entête secteur)
        CLR_KPI_HDR = 'EEF2F7'   # gris clair

        def _kpi_interp(key, val, is_bank):
            """Retourne une interprétation courte et détaillée selon l'indicateur et le secteur."""
            if val is None: return "Donnée absente"
            v = val
            if key == "roe":
                if v > 0.20: return f"✅ ROE de {v*100:.1f}% — excellente rentabilité pour les actionnaires, bien au-dessus du coût du capital UEMOA (≈10-12%). Création de valeur forte."
                if v > 0.12: return f"✅ ROE de {v*100:.1f}% — rentabilité satisfaisante, rémunère correctement les capitaux propres."
                if v > 0.05: return f"⚠️ ROE de {v*100:.1f}% — rentabilité modeste, inférieure au coût du capital. Les actionnaires ne sont pas pleinement rémunérés."
                if v > 0:    return f"⚠️ ROE de {v*100:.1f}% — très faible rentabilité. Risque de dépréciation de la valeur actionnariale."
                return f"❌ ROE négatif ({v*100:.1f}%) — la société détruit de la valeur. Perte nette sur les fonds propres."
            if key == "roa":
                seuil = 0.01 if is_bank else 0.05
                bench = "seuil banque UEMOA ≈ 1%" if is_bank else "seuil entreprise ≈ 5%"
                if v > seuil * 2: return f"✅ ROA de {v*100:.2f}% — très bonne utilisation des actifs ({bench}). Efficacité opérationnelle supérieure à la moyenne sectorielle."
                if v > seuil:    return f"✅ ROA de {v*100:.2f}% — bonne utilisation des actifs ({bench})."
                if v > 0:        return f"⚠️ ROA de {v*100:.2f}% — en dessous du seuil sectoriel ({bench}). Les actifs ne génèrent pas assez de résultat."
                return f"❌ ROA négatif ({v*100:.2f}%) — actifs non rentables."
            if key == "marge_nette":
                if v > 0.15: return f"✅ Marge nette de {v*100:.1f}% — excellente profitabilité finale. La société conserve plus de 15 centimes par FCFA de revenu après toutes les charges."
                if v > 0.05: return f"✅ Marge nette de {v*100:.1f}% — profitabilité satisfaisante."
                if v > 0.01: return f"⚠️ Marge nette de {v*100:.1f}% — marge très fine, vulnérable aux chocs de coûts ou de revenus."
                if v > 0:    return f"⚠️ Marge nette quasi-nulle ({v*100:.2f}%) — activité à peine bénéficiaire."
                return f"❌ Marge nette négative ({v*100:.1f}%) — la société est déficitaire."
            if key == "resultat_net":
                if v > 0: return f"✅ Résultat net positif de {_fmt_mds(v)} — société bénéficiaire, capacité à distribuer des dividendes et à renforcer ses fonds propres."
                return f"❌ Perte nette de {_fmt_mds(abs(v))} — impact négatif sur les fonds propres et impossibilité de distribuer des dividendes."
            if key == "taux_croissance_ca":
                if v > 0.15: return f"✅ Croissance de {v*100:.1f}% — forte expansion, bien au-dessus du PIB UEMOA (≈6%). Gain de parts de marché ou hausse de volume significative."
                if v > 0.05: return f"✅ Croissance de {v*100:.1f}% — dynamisme commercial solide, supérieur à l'inflation régionale."
                if v > 0:    return f"⚠️ Croissance faible ({v*100:.1f}%) — activité quasi-stagnante. Risque de perte de compétitivité en termes réels."
                if v > -0.05: return f"⚠️ Légère baisse ({v*100:.1f}%) — contraction modérée du revenu principal."
                return f"❌ Recul marqué de {abs(v)*100:.1f}% — perte de revenus significative, signal d'alerte fort."
            if key == "autonomie_financiere":
                if v > 0.40: return f"✅ Autonomie financière de {v*100:.1f}% — société très peu dépendante des créanciers. Capacité élevée à absorber les chocs."
                if v > 0.20: return f"✅ Autonomie financière correcte ({v*100:.1f}%) — équilibre sain entre fonds propres et dettes."
                if v > 0.10: return f"⚠️ Autonomie limitée ({v*100:.1f}%) — dépendance significative aux financeurs externes. Risque accru en cas de resserrement du crédit."
                return f"❌ Autonomie très faible ({v*100:.1f}%) — quasi-totalité des actifs financée par des dettes. Fragilité structurelle."
            if key == "ratio_endettement":
                if v < 0.50: return f"✅ Gearing de {v*100:.0f}% — endettement très modéré par rapport aux fonds propres. Capacité d'emprunt préservée."
                if v < 1.50: return f"✅ Gearing de {v*100:.0f}% — levier raisonnable, standard pour le secteur."
                if v < 3.00: return f"⚠️ Gearing de {v*100:.0f}% — levier élevé, la dette dépasse les fonds propres. Sensibilité aux variations de taux."
                return f"❌ Gearing de {v*100:.0f}% — sur-endettement critique. Risque de défaillance financière si les revenus baissent."
            if key == "solvabilite_generale":
                if v > 2.0: return f"✅ Solvabilité de {v*100:.0f}% — actifs très largement supérieurs aux dettes. Confort absolu pour les créanciers."
                if v > 1.2: return f"✅ Solvabilité satisfaisante ({v*100:.0f}%) — les actifs couvrent bien les dettes."
                if v > 1.0: return f"⚠️ Solvabilité limite ({v*100:.0f}%) — marge de sécurité faible pour les créanciers."
                return f"❌ Insolvable ({v*100:.0f}%) — les dettes dépassent la valeur totale des actifs. Situation critique."
            if key == "liquidite_generale":
                if v > 2.0: return f"✅ Liquidité générale de {v*100:.0f}x — actif circulant deux fois supérieur aux dettes CT. Trésorerie très confortable."
                if v > 1.2: return f"✅ Liquidité générale de {v*100:.2f}x — bonne capacité à honorer les obligations CT."
                if v > 1.0: return f"⚠️ Liquidité générale limite ({v*100:.2f}x) — couverture juste suffisante."
                return f"❌ Liquidité insuffisante ({v*100:.2f}x) — actif circulant ne couvre pas les dettes CT. Risque de défaut de paiement."
            if key == "free_cash_flow":
                if v and v > 0: return f"✅ FCF positif de {_fmt_mds(v)} — la société génère du cash après investissements. Peut rembourser la dette, verser des dividendes ou investir davantage."
                if v and v > -5e8: return f"⚠️ FCF légèrement négatif ({_fmt_mds(v)}) — investissements supérieurs au cash opérationnel. Phase de développement."
                return f"❌ FCF négatif significatif ({_fmt_mds(v) if v else 'N/D'}) — la société consomme de la trésorerie. Risque si la situation perdure."
            if key == "coefficient_exploitation":
                if v < 0.45: return f"✅ Coeff. d'exploitation de {v*100:.1f}% — banque très efficiente. Moins de 45 centimes de charges par FCFA de PNB. Excellence opérationnelle."
                if v < 0.60: return f"✅ Coeff. d'exploitation de {v*100:.1f}% — dans la norme UEMOA. Efficience satisfaisante."
                if v < 0.75: return f"⚠️ Coeff. d'exploitation de {v*100:.1f}% — charges absorbant trop du PNB. Marge de progression nécessaire."
                return f"❌ Coeff. d'exploitation de {v*100:.1f}% — structure de coûts critique. La banque est peu viable à ce niveau."
            if key == "marge_operationnelle":
                if v > 0.20: return f"✅ Marge opérationnelle de {v*100:.1f}% — excellente efficacité de l'exploitation principale."
                if v > 0.08: return f"✅ Marge opérationnelle correcte ({v*100:.1f}%)."
                if v > 0:    return f"⚠️ Marge opérationnelle faible ({v*100:.1f}%) — l'exploitation laisse peu de bénéfice avant charges financières."
                return f"❌ Marge opérationnelle négative ({v*100:.1f}%) — l'activité principale est déficitaire."
            if key == "delai_clients":
                if v < 30: return f"✅ Délai clients de {v:.0f} jours — recouvrement rapide. Trésorerie peu sollicitée par les créances clients."
                if v < 60: return f"✅ Délai standard ({v:.0f} jours) — dans les normes commerciales habituelles."
                if v < 90: return f"⚠️ Délai long ({v:.0f} jours) — risque d'impayés accru, pression sur la trésorerie CT."
                return f"❌ Délai critique de {v:.0f} jours — process de recouvrement à revoir. Immobilisation de capital excessive."
            if key == "duree_stockage":
                if v < 30: return f"✅ Stock écoulé en {v:.0f} jours — rotation excellente, capital peu immobilisé."
                if v < 60: return f"✅ Durée de stockage standard ({v:.0f} jours)."
                if v < 120: return f"⚠️ Stock lent ({v:.0f} jours) — immobilisation de capital importante. Risque d'obsolescence."
                return f"❌ Stock quasi-bloqué ({v:.0f} jours) — capital fortement immobilisé, risque d'obsolescence élevé."
            return f"Valeur : {_fmt_pct(v) if isinstance(v, float) and abs(v) < 10 else _fmt_mds(v)}"

        # ── Générer la section par secteur ─────────────────────────────────────
        sectors_with_data = {}
        for sec, syms in sorted(sector_map.items()):
            syms_with_fin = [s for s in syms if s in fin_all]
            if len(syms_with_fin) >= 1:
                sectors_with_data[sec] = syms_with_fin

        if not sectors_with_data:
            doc.add_paragraph("⚠️ Aucune donnée financière structurée disponible pour cette analyse.")
        else:
            note_p = doc.add_paragraph()
            note_r = note_p.add_run(
                f"Données disponibles pour {len(fin_all)} société(s) sur {len(all_company_data)}. "
                f"Sociétés sans données ignorées. "
                f"🟢 = performance positive  🔴 = performance à surveiller."
            )
            note_r.font.size = Pt(8.5)
            note_r.font.italic = True
            note_r.font.color.rgb = RGBColor(60, 60, 60)
            doc.add_paragraph()

            for sec_name, syms_fin in sectors_with_data.items():
                # ── Titre secteur ──────────────────────────────────────────────
                doc.add_page_break()
                sec_hdg = doc.add_heading(f"📂 Secteur : {sec_name}", level=2)
                sec_hdg.runs[0].font.color.rgb = RGBColor(0, 51, 102)

                n_total = len(sector_map.get(sec_name, []))
                n_data  = len(syms_fin)
                sub_p = doc.add_paragraph()
                sub_p.add_run(
                    f"{n_data} société(s) avec données financières sur {n_total} cotée(s) dans ce secteur. "
                    f"Exercice(s) analysé(s) : {', '.join(str(fin_all[s].get('annee','?')) for s in syms_fin)}."
                ).font.size = Pt(9)
                doc.add_paragraph()

                # ── Tableau récapitulatif KPI par société ─────────────────────
                doc.add_heading("📊 Tableau comparatif des indicateurs clés", level=3)

                # Sélectionner les KPI pertinents selon le secteur
                is_bank_sector = any(
                    any([
                        fin_all[s].get('caisse_banque_centrale') and float(fin_all[s].get('caisse_banque_centrale') or 0) != 0,
                        fin_all[s].get('produit_net_bancaire')   and float(fin_all[s].get('produit_net_bancaire')   or 0) != 0,
                        fin_all[s].get('dettes_clientele')       and float(fin_all[s].get('dettes_clientele')       or 0) != 0,
                    ])
                    for s in syms_fin
                )
                kpis = [(l, k, f, s) for l, k, f, s in KPI_DEFS
                        if not (k == "coefficient_exploitation" and not is_bank_sector)
                        and not (k in ["delai_clients","duree_stockage"] and is_bank_sector)]

                # Tableau : ligne = KPI, colonne = société
                n_soc = len(syms_fin)
                tbl_cmp = doc.add_table(rows=1, cols=1 + n_soc)
                tbl_cmp.style = 'Table Grid'
                # En-tête
                hdr_cmp = tbl_cmp.rows[0].cells
                hdr_cmp[0].text = "Indicateur"
                r0h = hdr_cmp[0].paragraphs[0].runs[0] if hdr_cmp[0].paragraphs[0].runs else hdr_cmp[0].paragraphs[0].add_run("Indicateur")
                r0h.bold = True; r0h.font.size = Pt(8); r0h.font.color.rgb = RGBColor(255,255,255)
                shd0h = OxmlElement('w:shd'); shd0h.set(qn('w:fill'), CLR_POS_HDR); shd0h.set(qn('w:val'), 'clear')
                hdr_cmp[0]._element.get_or_add_tcPr().append(shd0h)
                for si, sym in enumerate(syms_fin):
                    c = hdr_cmp[si+1]
                    cname = all_company_data.get(sym, {}).get('company_name', sym)
                    c.text = f"{sym} ({fin_all[sym].get('annee','?')})"
                    cr = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(c.text)
                    cr.bold = True; cr.font.size = Pt(7.5); cr.font.color.rgb = RGBColor(255,255,255)
                    shdc = OxmlElement('w:shd'); shdc.set(qn('w:fill'), CLR_POS_HDR); shdc.set(qn('w:val'), 'clear')
                    c._element.get_or_add_tcPr().append(shdc)

                # Lignes KPI
                for kpi_lbl, kpi_key, kpi_fmt, kpi_sens in kpis:
                    vals = {s: _safe(fin_all[s], kpi_key) for s in syms_fin}
                    valid_vals = {s: v for s, v in vals.items() if v is not None}
                    if not valid_vals:
                        continue  # Sauter ce KPI si aucune donnée

                    # Trouver meilleur et pire
                    best_sym = max(valid_vals, key=lambda s: valid_vals[s] * kpi_sens)
                    worst_sym = min(valid_vals, key=lambda s: valid_vals[s] * kpi_sens)
                    best_val = valid_vals[best_sym]
                    worst_val = valid_vals[worst_sym]

                    tr = tbl_cmp.add_row().cells
                    # Col 0 : label KPI
                    tr[0].text = kpi_lbl
                    rl = tr[0].paragraphs[0].runs[0] if tr[0].paragraphs[0].runs else tr[0].paragraphs[0].add_run(kpi_lbl)
                    rl.bold = True; rl.font.size = Pt(7.5)
                    shd_lbl = OxmlElement('w:shd'); shd_lbl.set(qn('w:fill'), CLR_KPI_HDR); shd_lbl.set(qn('w:val'), 'clear')
                    tr[0]._element.get_or_add_tcPr().append(shd_lbl)

                    for si, sym in enumerate(syms_fin):
                        v = vals.get(sym)
                        if kpi_fmt == "pct":
                            vstr = _fmt_pct(v) if v is not None else "N/D"
                        elif kpi_fmt == "mds":
                            vstr = _fmt_mds(v) if v is not None else "N/D"
                        else:
                            vstr = f"{v:.1f} j" if v is not None else "N/D"

                        c = tr[si+1]
                        c.text = vstr
                        rc = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(vstr)
                        rc.font.size = Pt(8); rc.bold = True

                        # Colorier fond selon rang
                        if v is None:
                            bg = 'F2F2F2'
                        elif sym == best_sym and len(valid_vals) > 1:
                            bg = CLR_POS
                        elif sym == worst_sym and len(valid_vals) > 1:
                            bg = CLR_NEG
                        else:
                            bg = 'FFFFFF'
                        shd_c = OxmlElement('w:shd'); shd_c.set(qn('w:fill'), bg); shd_c.set(qn('w:val'), 'clear')
                        c._element.get_or_add_tcPr().append(shd_c)

                        # Couleur texte
                        if v is not None and v != 0:
                            rc.font.color.rgb = RGBColor(0,100,0) if sym == best_sym and len(valid_vals) > 1 else (
                                RGBColor(180,0,0) if sym == worst_sym and len(valid_vals) > 1 else RGBColor(0,0,0))

                # Largeurs colonnes
                try:
                    total_w = Cm(17.5)
                    lbl_w   = Cm(4.5)
                    val_w   = Cm((17.5 - 4.5) / max(n_soc, 1))
                    for ci, col in enumerate(tbl_cmp.columns):
                        for cell in col.cells:
                            cell.width = lbl_w if ci == 0 else val_w
                except: pass
                doc.add_paragraph()

                # ── Analyse narrative par société ──────────────────────────────
                doc.add_heading("🔍 Analyse détaillée par société", level=3)

                for sym in syms_fin:
                    fd = fin_all[sym]
                    annee_fd = fd.get('annee', '?')
                    cname = all_company_data.get(sym, {}).get('company_name', sym)
                    is_bank_s = any([
                        fd.get('caisse_banque_centrale') and float(fd.get('caisse_banque_centrale') or 0) != 0,
                        fd.get('produit_net_bancaire')   and float(fd.get('produit_net_bancaire')   or 0) != 0,
                        fd.get('dettes_clientele')       and float(fd.get('dettes_clientele')       or 0) != 0,
                    ])
                    secteur_lbl = "🏦 Banque" if is_bank_s else "🏢 Entreprise"

                    # Titre société
                    soc_p = doc.add_paragraph()
                    soc_p.paragraph_format.space_before = Pt(8)
                    soc_r1 = soc_p.add_run(f"▶ {sym} — {cname} ")
                    soc_r1.bold = True; soc_r1.font.size = Pt(10)
                    soc_r1.font.color.rgb = RGBColor(0, 51, 102)
                    soc_r2 = soc_p.add_run(f"[{secteur_lbl} — Exercice {annee_fd}]")
                    soc_r2.font.size = Pt(8.5); soc_r2.font.italic = True
                    soc_r2.font.color.rgb = RGBColor(80, 80, 80)

                    # Tableau interprétation : Indicateur | Valeur | Interprétation détaillée
                    tbl_soc = doc.add_table(rows=1, cols=3)
                    tbl_soc.style = 'Table Grid'
                    hdr_s = tbl_soc.rows[0].cells
                    for ci_s, ht in enumerate(['Indicateur', 'Valeur', 'Interprétation détaillée']):
                        hdr_s[ci_s].text = ht
                        rh = hdr_s[ci_s].paragraphs[0].runs[0] if hdr_s[ci_s].paragraphs[0].runs else hdr_s[ci_s].paragraphs[0].add_run(ht)
                        rh.bold = True; rh.font.size = Pt(8); rh.font.color.rgb = RGBColor(255,255,255)
                        shd_h = OxmlElement('w:shd'); shd_h.set(qn('w:fill'), '2F5496'); shd_h.set(qn('w:val'), 'clear')
                        hdr_s[ci_s]._element.get_or_add_tcPr().append(shd_h)

                    for kpi_lbl, kpi_key, kpi_fmt, kpi_sens in kpis:
                        v = _safe(fd, kpi_key)
                        if v is None: continue
                        if kpi_fmt == "pct":  vstr = _fmt_pct(v)
                        elif kpi_fmt == "mds": vstr = _fmt_mds(v)
                        else:                  vstr = f"{v:.1f} j"

                        interp_txt = _kpi_interp(kpi_key, v, is_bank_s)
                        # Couleur fond selon signal
                        if "✅" in interp_txt:    bg_s = CLR_POS
                        elif "❌" in interp_txt:  bg_s = CLR_NEG
                        elif "⚠️" in interp_txt: bg_s = CLR_MID
                        else:                     bg_s = 'FFFFFF'

                        tr_s = tbl_soc.add_row().cells
                        tr_s[0].text = kpi_lbl
                        r_lbl = tr_s[0].paragraphs[0].runs[0] if tr_s[0].paragraphs[0].runs else tr_s[0].paragraphs[0].add_run(kpi_lbl)
                        r_lbl.bold = True; r_lbl.font.size = Pt(7.5)
                        shd_l = OxmlElement('w:shd'); shd_l.set(qn('w:fill'), CLR_KPI_HDR); shd_l.set(qn('w:val'), 'clear')
                        tr_s[0]._element.get_or_add_tcPr().append(shd_l)

                        tr_s[1].text = vstr
                        r_val = tr_s[1].paragraphs[0].runs[0] if tr_s[1].paragraphs[0].runs else tr_s[1].paragraphs[0].add_run(vstr)
                        r_val.bold = True; r_val.font.size = Pt(8.5)
                        try:
                            nv = float(vstr.replace('%','').replace(' Mds FCFA','').replace(' M FCFA','').replace(' FCFA','').replace(' j','').replace(',','').strip())
                            r_val.font.color.rgb = RGBColor(0,120,0) if nv > 0 else RGBColor(180,0,0) if nv < 0 else RGBColor(80,80,80)
                        except: r_val.font.color.rgb = RGBColor(0,0,0)
                        shd_v = OxmlElement('w:shd'); shd_v.set(qn('w:fill'), bg_s); shd_v.set(qn('w:val'), 'clear')
                        tr_s[1]._element.get_or_add_tcPr().append(shd_v)

                        tr_s[2].text = interp_txt
                        r_int = tr_s[2].paragraphs[0].runs[0] if tr_s[2].paragraphs[0].runs else tr_s[2].paragraphs[0].add_run(interp_txt)
                        r_int.font.size = Pt(7.5)
                        shd_i = OxmlElement('w:shd'); shd_i.set(qn('w:fill'), 'FFFFFF'); shd_i.set(qn('w:val'), 'clear')
                        tr_s[2]._element.get_or_add_tcPr().append(shd_i)

                    # Largeurs
                    try:
                        for ci_s, col_s in enumerate(tbl_soc.columns):
                            for cell_s in col_s.cells:
                                cell_s.width = [Cm(4.0), Cm(3.5), Cm(10.0)][ci_s]
                    except: pass
                    doc.add_paragraph()

                # ── Synthèse du secteur : TOP / BOTTOM ────────────────────────
                doc.add_heading("🏆 Synthèse sectorielle — Points forts et points faibles", level=3)

                podium_rows = []
                # Pour chaque KPI clé, identifier le meilleur et le moins bon
                kpis_synthese = [
                    ("ROE",             "roe",              "pct", +1),
                    ("Marge nette",     "marge_nette",      "pct", +1),
                    ("Croissance CA/PNB","taux_croissance_ca","pct",+1),
                    ("Autonomie fin.",  "autonomie_financiere","pct",+1),
                    ("Liquidité gén.",  "liquidite_generale","pct", +1),
                    ("FCF",            "free_cash_flow",    "mds", +1),
                ]
                if is_bank_sector:
                    kpis_synthese.append(("Coeff. exploit.", "coefficient_exploitation","pct",-1))

                for k_lbl, k_key, k_fmt, k_sens in kpis_synthese:
                    vals_syn = {s: _safe(fin_all[s], k_key) for s in syms_fin if _safe(fin_all[s], k_key) is not None}
                    if len(vals_syn) < 2: continue
                    best  = max(vals_syn, key=lambda s: vals_syn[s] * k_sens)
                    worst = min(vals_syn, key=lambda s: vals_syn[s] * k_sens)
                    bv = vals_syn[best]; wv = vals_syn[worst]
                    bstr = _fmt_pct(bv) if k_fmt == "pct" else _fmt_mds(bv)
                    wstr = _fmt_pct(wv) if k_fmt == "pct" else _fmt_mds(wv)
                    podium_rows.append((k_lbl, best, bstr, worst, wstr))

                if podium_rows:
                    tbl_pod = doc.add_table(rows=1, cols=5)
                    tbl_pod.style = 'Table Grid'
                    hdr_pod = tbl_pod.rows[0].cells
                    for ci_p, ht_p in enumerate(['Indicateur','🥇 Meilleure','Valeur','🔻 À surveiller','Valeur']):
                        hdr_pod[ci_p].text = ht_p
                        rh_p = hdr_pod[ci_p].paragraphs[0].runs[0] if hdr_pod[ci_p].paragraphs[0].runs else hdr_pod[ci_p].paragraphs[0].add_run(ht_p)
                        rh_p.bold = True; rh_p.font.size = Pt(8); rh_p.font.color.rgb = RGBColor(255,255,255)
                        shd_ph = OxmlElement('w:shd'); shd_ph.set(qn('w:fill'), '375623' if ci_p == 1 else '843C0C' if ci_p == 3 else '1F4E79'); shd_ph.set(qn('w:val'), 'clear')
                        hdr_pod[ci_p]._element.get_or_add_tcPr().append(shd_ph)
                    for k_lbl, best, bstr, worst, wstr in podium_rows:
                        tr_p = tbl_pod.add_row().cells
                        tr_p[0].text = k_lbl
                        r0p = tr_p[0].paragraphs[0].runs[0] if tr_p[0].paragraphs[0].runs else tr_p[0].paragraphs[0].add_run(k_lbl)
                        r0p.bold = True; r0p.font.size = Pt(8)
                        shd_0p = OxmlElement('w:shd'); shd_0p.set(qn('w:fill'), CLR_KPI_HDR); shd_0p.set(qn('w:val'), 'clear')
                        tr_p[0]._element.get_or_add_tcPr().append(shd_0p)
                        # Meilleur
                        tr_p[1].text = best
                        r1p = tr_p[1].paragraphs[0].runs[0] if tr_p[1].paragraphs[0].runs else tr_p[1].paragraphs[0].add_run(best)
                        r1p.bold = True; r1p.font.size = Pt(8); r1p.font.color.rgb = RGBColor(0,100,0)
                        shd_1p = OxmlElement('w:shd'); shd_1p.set(qn('w:fill'), CLR_POS); shd_1p.set(qn('w:val'), 'clear')
                        tr_p[1]._element.get_or_add_tcPr().append(shd_1p)
                        tr_p[2].text = bstr
                        r2p = tr_p[2].paragraphs[0].runs[0] if tr_p[2].paragraphs[0].runs else tr_p[2].paragraphs[0].add_run(bstr)
                        r2p.font.size = Pt(8); r2p.font.color.rgb = RGBColor(0,100,0)
                        shd_2p = OxmlElement('w:shd'); shd_2p.set(qn('w:fill'), CLR_POS); shd_2p.set(qn('w:val'), 'clear')
                        tr_p[2]._element.get_or_add_tcPr().append(shd_2p)
                        # Moins bon
                        tr_p[3].text = worst
                        r3p = tr_p[3].paragraphs[0].runs[0] if tr_p[3].paragraphs[0].runs else tr_p[3].paragraphs[0].add_run(worst)
                        r3p.bold = True; r3p.font.size = Pt(8); r3p.font.color.rgb = RGBColor(180,0,0)
                        shd_3p = OxmlElement('w:shd'); shd_3p.set(qn('w:fill'), CLR_NEG); shd_3p.set(qn('w:val'), 'clear')
                        tr_p[3]._element.get_or_add_tcPr().append(shd_3p)
                        tr_p[4].text = wstr
                        r4p = tr_p[4].paragraphs[0].runs[0] if tr_p[4].paragraphs[0].runs else tr_p[4].paragraphs[0].add_run(wstr)
                        r4p.font.size = Pt(8); r4p.font.color.rgb = RGBColor(180,0,0)
                        shd_4p = OxmlElement('w:shd'); shd_4p.set(qn('w:fill'), CLR_NEG); shd_4p.set(qn('w:val'), 'clear')
                        tr_p[4]._element.get_or_add_tcPr().append(shd_4p)
                    doc.add_paragraph()

        doc.add_page_break()

        # ========== TABLE DES MATIÈRES ==========
        doc.add_heading('TABLE DES MATIÈRES - ANALYSES DÉTAILLÉES', level=1)
        for idx, (symbol, data) in enumerate(sorted(all_company_data.items()), 1):
            company_name = data.get('company_name', 'N/A')
            toc_p = doc.add_paragraph(f"{idx}. {symbol} - {company_name}")
            toc_p.paragraph_format.left_indent = Pt(20)
        
        doc.add_page_break()
        
        # ========== ANALYSES DÉTAILLÉES ==========
        doc.add_heading('ANALYSES DÉTAILLÉES PAR SOCIÉTÉ', level=1)
        
        for idx, (symbol, analysis) in enumerate(sorted(all_analyses.items()), 1):
            company_data = all_company_data.get(symbol, {})
            company_name = company_data.get('company_name', 'N/A')
            
            company_heading = doc.add_heading(f"{idx}. {symbol} - {company_name}", level=2)
            company_heading.paragraph_format.space_before = Pt(18)
            company_heading_run = company_heading.runs[0]
            company_heading_run.font.color.rgb = RGBColor(0, 102, 204)
            
            doc.add_paragraph("─" * 80)

            # ══════════════════════════════════════════════════════════════════
            # CARTE D'IDENTITÉ DE LA SOCIÉTÉ
            # ══════════════════════════════════════════════════════════════════
            sector_raw      = company_data.get('sector') or 'Non classifié'
            cap_txt_raw     = company_data.get('capitalisation_txt', 'N/D')
            cap_val_raw     = company_data.get('capitalisation')
            cur_price_raw   = company_data.get('current_price')
            evol_100d_raw   = company_data.get('price_evolution_100d')
            hi_100d_raw     = company_data.get('highest_price_100d')
            lo_100d_raw     = company_data.get('lowest_price_100d')
            vol_moy_raw     = company_data.get('volume_moyen_jour')
            vol_ann_raw     = company_data.get('vol_annualisee')
            reco_raw        = company_data.get('recommendation', 'N/A')
            risk_lv_raw     = company_data.get('risk_level', 'N/A')
            conf_lv_raw     = company_data.get('confidence_level', 'N/A')
            inv_lbl_raw     = company_data.get('investment_label', 'N/A')

            # ── Couleur recommandation ─────────────────────────────────────────
            reco_up = str(reco_raw).upper()
            if   'ACHAT FORT' in reco_up: reco_col, reco_bg = RGBColor(0,120,0),  '00B050'
            elif 'ACHAT'      in reco_up: reco_col, reco_bg = RGBColor(0,150,0),  'C6EFCE'
            elif 'VENTE FORT' in reco_up: reco_col, reco_bg = RGBColor(180,0,0),  'FF0000'
            elif 'VENTE'      in reco_up: reco_col, reco_bg = RGBColor(200,0,0),  'FFC7CE'
            else:                         reco_col, reco_bg = RGBColor(100,100,0), 'FFEB9C'

            # ── Couleur capitalisation ─────────────────────────────────────────
            if cap_val_raw:
                if cap_val_raw >= 500e9:   cap_label = "Grande cap (≥500 Mds)"
                elif cap_val_raw >= 100e9: cap_label = "Moyenne cap (100-500 Mds)"
                elif cap_val_raw >= 10e9:  cap_label = "Petite cap (10-100 Mds)"
                else:                      cap_label = "Micro cap (<10 Mds)"
            else:
                cap_label = "N/D"

            # ── Variation 100j couleur ─────────────────────────────────────────
            if evol_100d_raw is not None:
                evol_str = f"{evol_100d_raw:+.2f}%"
                evol_col = RGBColor(0,120,0) if evol_100d_raw >= 0 else RGBColor(180,0,0)
            else:
                evol_str = "N/D"; evol_col = RGBColor(80,80,80)

            # ── Construction du tableau carte d'identité ─────────────────────
            id_card = doc.add_table(rows=1, cols=4)
            id_card.style = 'Table Grid'

            # Ligne 1 : en-têtes grandes catégories
            hdr_id = id_card.rows[0].cells
            id_sections = [
                ('🏢 IDENTITÉ',       '1F4E79'),
                ('💰 CAPITALISATION', '375623'),
                ('📈 COURS & MARCHÉ', '833C00'),
                ('🎯 RECOMMANDATION', reco_bg if len(reco_bg) == 6 else '595959'),
            ]
            for ci_id, (lbl_id, bg_id) in enumerate(id_sections):
                c = hdr_id[ci_id]
                c.text = lbl_id
                r = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(lbl_id)
                r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(255,255,255)
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), bg_id); shd.set(qn('w:val'), 'clear')
                c._element.get_or_add_tcPr().append(shd)

            # Ligne 2 : valeurs
            tr_id = id_card.add_row().cells

            # Col 0 : Identité
            id_lines = [
                ("Symbole",  symbol),
                ("Société",  company_name[:40] + ('...' if len(company_name) > 40 else '')),
                ("Secteur",  sector_raw),
            ]
            id_para = tr_id[0].paragraphs[0]
            for li_lbl, li_val in id_lines:
                if id_para.runs:
                    id_para = tr_id[0].add_paragraph()
                r_lbl = id_para.add_run(f"{li_lbl} : ")
                r_lbl.bold = True; r_lbl.font.size = Pt(8)
                r_lbl.font.color.rgb = RGBColor(50, 50, 80)
                r_val = id_para.add_run(li_val)
                r_val.font.size = Pt(8)
                r_val.font.color.rgb = RGBColor(0, 0, 0)
            shd0_id = OxmlElement('w:shd'); shd0_id.set(qn('w:fill'), 'DEEAF1'); shd0_id.set(qn('w:val'), 'clear')
            tr_id[0]._element.get_or_add_tcPr().append(shd0_id)

            # Col 1 : Capitalisation
            cap_lines = [
                ("Capitalisation", cap_txt_raw),
                ("Catégorie",      cap_label),
                ("Vol. moy/jour",  f"{vol_moy_raw:,.0f} titres" if vol_moy_raw else "N/D"),
            ]
            cap_para = tr_id[1].paragraphs[0]
            for li_lbl, li_val in cap_lines:
                if cap_para.runs:
                    cap_para = tr_id[1].add_paragraph()
                r_lbl = cap_para.add_run(f"{li_lbl} : ")
                r_lbl.bold = True; r_lbl.font.size = Pt(8)
                r_lbl.font.color.rgb = RGBColor(30, 70, 30)
                r_val = cap_para.add_run(li_val)
                r_val.font.size = Pt(8.5); r_val.bold = True
                r_val.font.color.rgb = RGBColor(0, 100, 0)
            shd1_id = OxmlElement('w:shd'); shd1_id.set(qn('w:fill'), 'E2EFDA'); shd1_id.set(qn('w:val'), 'clear')
            tr_id[1]._element.get_or_add_tcPr().append(shd1_id)

            # Col 2 : Cours & Marché
            cours_lines = [
                ("Cours actuel",   f"{cur_price_raw:,.0f} FCFA" if cur_price_raw else "N/D"),
                ("Var. 100 jours", evol_str),
                ("Plus haut 100j", f"{hi_100d_raw:,.0f} FCFA" if hi_100d_raw else "N/D"),
                ("Plus bas 100j",  f"{lo_100d_raw:,.0f} FCFA"  if lo_100d_raw  else "N/D"),
                ("Volatilité ann.", f"{vol_ann_raw:.1f}%" if vol_ann_raw else "N/D"),
            ]
            cours_para = tr_id[2].paragraphs[0]
            for li_i, (li_lbl, li_val) in enumerate(cours_lines):
                if cours_para.runs:
                    cours_para = tr_id[2].add_paragraph()
                r_lbl = cours_para.add_run(f"{li_lbl} : ")
                r_lbl.bold = True; r_lbl.font.size = Pt(8)
                r_lbl.font.color.rgb = RGBColor(80, 50, 0)
                r_val = cours_para.add_run(li_val)
                r_val.font.size = Pt(8.5)
                if li_i == 1:  # Variation : colorer
                    r_val.bold = True; r_val.font.color.rgb = evol_col
                else:
                    r_val.font.color.rgb = RGBColor(0, 0, 0)
            shd2_id = OxmlElement('w:shd'); shd2_id.set(qn('w:fill'), 'FFF2CC'); shd2_id.set(qn('w:val'), 'clear')
            tr_id[2]._element.get_or_add_tcPr().append(shd2_id)

            # Col 3 : Recommandation
            reco_lines = [
                ("Recommandation", reco_raw),
                ("Confiance",      conf_lv_raw),
                ("Risque",         risk_lv_raw),
                ("Horizon",        company_data.get('investment_horizon', 'Moyen terme')),
                ("Score invest.",  inv_lbl_raw),
            ]
            reco_para = tr_id[3].paragraphs[0]
            for li_i, (li_lbl, li_val) in enumerate(reco_lines):
                if reco_para.runs:
                    reco_para = tr_id[3].add_paragraph()
                r_lbl = reco_para.add_run(f"{li_lbl} : ")
                r_lbl.bold = True; r_lbl.font.size = Pt(8)
                r_lbl.font.color.rgb = RGBColor(80, 0, 0)
                r_val = reco_para.add_run(str(li_val))
                r_val.font.size = Pt(8.5)
                if li_i == 0:  # Recommandation en gros
                    r_val.bold = True; r_val.font.size = Pt(9.5)
                    r_val.font.color.rgb = reco_col
                else:
                    r_val.font.color.rgb = RGBColor(0, 0, 0)
            shd3_id = OxmlElement('w:shd'); shd3_id.set(qn('w:fill'), 'FCE4D6'); shd3_id.set(qn('w:val'), 'clear')
            tr_id[3]._element.get_or_add_tcPr().append(shd3_id)

            # Largeurs colonnes : identité=4.5cm, capi=4cm, cours=4cm, reco=5cm
            try:
                widths_id = [Cm(4.5), Cm(4.0), Cm(4.0), Cm(5.0)]
                for ci_id2, col_id in enumerate(id_card.columns):
                    for cell_id in col_id.cells:
                        cell_id.width = widths_id[ci_id2]
            except: pass
            doc.add_paragraph()

            # ══════════════════════════════════════════════════════════════
            # TABLEAU DU SCORE DE RISQUE DÉTAILLÉ
            # ══════════════════════════════════════════════════════════════
            risk_details_raw = company_data.get('risk_details', '{}')
            try:
                risk_details = json.loads(risk_details_raw) if isinstance(risk_details_raw, str) else risk_details_raw
            except Exception:
                risk_details = {}

            risk_score_val  = company_data.get('risk_score', 0)
            risk_level_val  = company_data.get('risk_level', 'N/A')

            # Couleur selon niveau de risque
            if   risk_level_val == "Faible":     rs_bg, rs_txt = 'C6EFCE', RGBColor(0, 100, 0)
            elif risk_level_val == "Moyen":       rs_bg, rs_txt = 'FFEB9C', RGBColor(130, 90, 0)
            elif risk_level_val == "Élevé":       rs_bg, rs_txt = 'FFC7CE', RGBColor(180, 0, 0)
            elif risk_level_val == "Très élevé":  rs_bg, rs_txt = 'FF0000', RGBColor(255, 255, 255)
            else:                                 rs_bg, rs_txt = 'F2F2F2', RGBColor(80, 80, 80)

            # ── Titre de la section ──────────────────────────────────────
            risk_hdg_p = doc.add_paragraph()
            risk_hdg_p.paragraph_format.space_before = Pt(6)
            risk_hdg_p.paragraph_format.space_after  = Pt(2)
            rh1 = risk_hdg_p.add_run("⚠️ SCORE DE RISQUE — ")
            rh1.bold = True; rh1.font.size = Pt(10); rh1.font.color.rgb = RGBColor(0, 51, 102)
            rh2 = risk_hdg_p.add_run(f"{risk_score_val:.1f}/100 → {risk_level_val}")
            rh2.bold = True; rh2.font.size = Pt(11); rh2.font.color.rgb = rs_txt

            # ── Barre de progression visuelle ───────────────────────────
            bar_p = doc.add_paragraph()
            bar_p.paragraph_format.space_before = Pt(2)
            bar_p.paragraph_format.space_after  = Pt(4)
            filled = int(risk_score_val / 5)   # 20 blocs max
            empty  = 20 - filled
            bar_r  = bar_p.add_run("█" * filled + "░" * empty + f"  {risk_score_val:.1f}/100")
            bar_r.font.size  = Pt(9)
            bar_r.font.color.rgb = rs_txt
            bar_r.font.name  = 'Courier New'

            # ── Tableau 5 critères ───────────────────────────────────────
            if risk_details:
                CRITERIA_META = {
                    'volatilite': {
                        'label':  '📊 Volatilité des prix (30%)',
                        'formule':'CV = σ(prix) / μ(prix)',
                        'interp': {
                            'Faible':  ('C6EFCE', '✅ Faible dispersion — titre stable'),
                            'Moyenne': ('FFEB9C', '⚠️ Dispersion modérée'),
                            'Élevée':  ('FFC7CE', '❌ Forte dispersion — titre très volatile'),
                        }
                    },
                    'beta': {
                        'label':  '📐 Bêta vs BRVM Composite (20%)',
                        'formule':'β = Cov(r_titre, r_indice) / Var(r_indice)',
                        'interp': {
                            'Défensif':             ('C6EFCE', '✅ β<0.8 — titre amplifie moins que le marché'),
                            'Neutre':               ('DEEAF1', '🔵 β≈1.0 — suit le marché'),
                            'Neutre (suit le marché)': ('DEEAF1', '🔵 β≈1.0 — suit le marché'),
                            'Agressif':             ('FFEB9C', '⚠️ β>1.2 — amplifie les hausses ET les baisses'),
                            'Très agressif':        ('FFC7CE', '❌ β>1.8 — très sensible aux chocs de marché'),
                            'Contre-cyclique':      ('E2EFDA', 'ℹ️ β<0 — évolue en sens inverse du marché'),
                        }
                    },
                    'liquidite': {
                        'label':  '💧 Liquidité du titre (20%)',
                        'formule':'Volume moyen journalier (100 jours)',
                        'interp': {
                            'Excellente': ('C6EFCE', '✅ >10 000 titres/j — sortie rapide possible'),
                            'Bonne':      ('DEEAF1', '🔵 1 000–10 000 titres/j — liquidité correcte'),
                            'Faible':     ('FFEB9C', '⚠️ 100–1 000 titres/j — risque de blocage'),
                            'Très faible':('FFC7CE', '❌ <100 titres/j — titre quasi-illiquide'),
                        }
                    },
                    'divergence': {
                        'label':  '🔀 Divergence signaux techniques (15%)',
                        'formule':'min(Achat,Vente) / Total signaux',
                        'interp': {
                            'Convergent': ('C6EFCE', '✅ Signaux alignés — clarté directionnelle'),
                            'Mixte':      ('FFEB9C', '⚠️ Signaux mixtes — incertitude modérée'),
                            'Très divergent': ('FFC7CE', '❌ Signaux contradictoires — forte incertitude'),
                        }
                    },
                    'stabilite': {
                        'label':  '📉 Stabilité des rendements (15%)',
                        'formule':'σ des rendements journaliers × 100',
                        'interp': {
                            'Très stable': ('C6EFCE', '✅ <1%/j — rendements très réguliers'),
                            'Stable':      ('DEEAF1', '🔵 1-3%/j — quelques variations normales'),
                            'Modérée':     ('FFEB9C', '⚠️ 3-5%/j — pics occasionnels'),
                            'Instable':    ('FFC7CE', '❌ >5%/j — forte agitation journalière'),
                        }
                    },
                }

                # Construction du tableau 4 colonnes
                risk_tbl = doc.add_table(rows=1, cols=4)
                risk_tbl.style = 'Table Grid'
                hdr_r = risk_tbl.rows[0].cells
                for ci_r, ht_r in enumerate(['Critère (poids)', 'Valeur mesurée', 'Formule', 'Interprétation']):
                    hdr_r[ci_r].text = ht_r
                    rhr = hdr_r[ci_r].paragraphs[0].runs[0] if hdr_r[ci_r].paragraphs[0].runs else hdr_r[ci_r].paragraphs[0].add_run(ht_r)
                    rhr.bold = True; rhr.font.size = Pt(8); rhr.font.color.rgb = RGBColor(255,255,255)
                    shd_rh = OxmlElement('w:shd'); shd_rh.set(qn('w:fill'), '243F60'); shd_rh.set(qn('w:val'), 'clear')
                    hdr_r[ci_r]._element.get_or_add_tcPr().append(shd_rh)

                metrics_order = ['volatilite', 'beta', 'liquidite', 'divergence', 'stabilite']
                for key in metrics_order:
                    raw_val = risk_details.get(key, '—')
                    if not raw_val or raw_val == '—':
                        continue
                    meta = CRITERIA_META.get(key, {})
                    lbl_col   = meta.get('label', key.capitalize())
                    formule   = meta.get('formule', '—')
                    raw_str   = str(raw_val)

                    # Déduire le niveau pour choisir la couleur
                    interp_map = meta.get('interp', {})
                    row_bg  = 'FFFFFF'
                    interp_txt = raw_str
                    for kw, (bg, txt) in interp_map.items():
                        if kw.lower() in raw_str.lower():
                            row_bg    = bg
                            interp_txt = txt
                            break

                    tr_r = risk_tbl.add_row().cells

                    # Col 0 : Critère
                    tr_r[0].text = lbl_col
                    r0r = tr_r[0].paragraphs[0].runs[0] if tr_r[0].paragraphs[0].runs else tr_r[0].paragraphs[0].add_run(lbl_col)
                    r0r.bold = True; r0r.font.size = Pt(7.5)
                    shd_r0 = OxmlElement('w:shd'); shd_r0.set(qn('w:fill'), 'EEF2F7'); shd_r0.set(qn('w:val'), 'clear')
                    tr_r[0]._element.get_or_add_tcPr().append(shd_r0)

                    # Col 1 : Valeur mesurée
                    tr_r[1].text = raw_str
                    r1r = tr_r[1].paragraphs[0].runs[0] if tr_r[1].paragraphs[0].runs else tr_r[1].paragraphs[0].add_run(raw_str)
                    r1r.font.size = Pt(7.5); r1r.bold = True
                    shd_r1 = OxmlElement('w:shd'); shd_r1.set(qn('w:fill'), row_bg); shd_r1.set(qn('w:val'), 'clear')
                    tr_r[1]._element.get_or_add_tcPr().append(shd_r1)

                    # Col 2 : Formule
                    tr_r[2].text = formule
                    r2r = tr_r[2].paragraphs[0].runs[0] if tr_r[2].paragraphs[0].runs else tr_r[2].paragraphs[0].add_run(formule)
                    r2r.font.size = Pt(7); r2r.font.italic = True
                    r2r.font.color.rgb = RGBColor(80, 80, 120)
                    shd_r2 = OxmlElement('w:shd'); shd_r2.set(qn('w:fill'), 'F7F7F7'); shd_r2.set(qn('w:val'), 'clear')
                    tr_r[2]._element.get_or_add_tcPr().append(shd_r2)

                    # Col 3 : Interprétation
                    tr_r[3].text = interp_txt
                    r3r = tr_r[3].paragraphs[0].runs[0] if tr_r[3].paragraphs[0].runs else tr_r[3].paragraphs[0].add_run(interp_txt)
                    r3r.font.size = Pt(7.5)
                    shd_r3 = OxmlElement('w:shd'); shd_r3.set(qn('w:fill'), row_bg); shd_r3.set(qn('w:val'), 'clear')
                    tr_r[3]._element.get_or_add_tcPr().append(shd_r3)

                # Largeurs : 5cm | 5.5cm | 3.5cm | 3.5cm
                try:
                    wds_r = [Cm(5.0), Cm(5.5), Cm(3.5), Cm(3.5)]
                    for cir2, colr in enumerate(risk_tbl.columns):
                        for cellr in colr.cells:
                            cellr.width = wds_r[cir2]
                except: pass

            doc.add_paragraph()

            # ── Feux tricolores indicateurs techniques ──────────────────────────
            tech_indicators = [
                ('MM',          company_data.get('mm_decision')),
                ('Bollinger',   company_data.get('bollinger_decision')),
                ('MACD',        company_data.get('macd_decision')),
                ('RSI',         company_data.get('rsi_decision')),
                ('Stochastique',company_data.get('stochastic_decision')),
            ]
            feux_p = doc.add_paragraph()
            feux_p.paragraph_format.space_before = Pt(4)
            feux_p.paragraph_format.space_after  = Pt(2)
            feux_p.add_run("Signaux techniques : ").bold = True
            for ind_name, ind_val in tech_indicators:
                iv = str(ind_val or '').upper()
                if   'ACHAT' in iv: icon, col = '🟢', RGBColor(0,128,0)
                elif 'VENTE' in iv: icon, col = '🔴', RGBColor(192,0,0)
                else:               icon, col = '🟡', RGBColor(160,120,0)
                r_icon = feux_p.add_run(f"{icon} {ind_name}  ")
                r_icon.font.size = Pt(9)
                r_icon.font.color.rgb = col

            # ── Score composite bandeau ──────────────────────────────────────────
            inv_score = company_data.get('investment_score', 0)
            inv_label = company_data.get('investment_label', '—')
            score_p = doc.add_paragraph()
            score_p.paragraph_format.space_before = Pt(4)
            score_p.paragraph_format.space_after  = Pt(4)
            sr1 = score_p.add_run("Score composite : ")
            sr1.bold = True
            sr1.font.size = Pt(11)
            sr2 = score_p.add_run(f"{inv_score}/100  ({inv_label})")
            sr2.bold = True
            sr2.font.size = Pt(13)
            if   inv_score >= 70: sr2.font.color.rgb = RGBColor(0,128,0)
            elif inv_score >= 55: sr2.font.color.rgb = RGBColor(0,100,160)
            elif inv_score >= 40: sr2.font.color.rgb = RGBColor(160,100,0)
            else:                 sr2.font.color.rgb = RGBColor(192,0,0)
            doc.add_paragraph()

            # ── Graphique cours réels + prédictions ML ────────────────────────
            hist_df_chart = self._get_historical_data_100days(company_data.get('company_id'))
            # Récupérer les prédictions déjà chargées dans company_data
            preds_for_chart = company_data.get('predictions_full', [])
            chart_buf = self._generate_price_chart_with_predictions(symbol, hist_df_chart, preds_for_chart)
            if chart_buf:
                try:
                    doc.add_picture(chart_buf, width=Inches(6.2))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except Exception as ce:
                    logging.warning(f"⚠️  Insertion image {symbol}: {ce}")

            # ── Tableau prédictions J+1 → J+10 avec IC ──────────────────────────
            preds_full = company_data.get('predictions_full', [])
            current_price = company_data.get('current_price') or 0
            if preds_full:
                doc.add_heading('🔮 Prédictions J+1 à J+10 (Intervalle de Confiance 90%)', level=3)
                pred_tbl = doc.add_table(rows=1, cols=6)
                pred_tbl.style = 'Light Grid Accent 1'
                pred_hdrs = ['Jour', 'Date', 'Borne basse', 'Prix prédit', 'Borne haute', 'Var. % / actuel']
                ph_cells = pred_tbl.rows[0].cells
                for ci, ph in enumerate(pred_hdrs):
                    ph_cells[ci].text = ph
                    r = ph_cells[ci].paragraphs[0].runs[0]
                    r.bold = True
                    r.font.size = Pt(8.5)
                    shd_ph = OxmlElement('w:shd')
                    shd_ph.set(qn('w:fill'), '1F4E79')
                    shd_ph.set(qn('w:val'), 'clear')
                    ph_cells[ci]._element.get_or_add_tcPr().append(shd_ph)
                    r.font.color.rgb = RGBColor(255,255,255)

                for ji, pred in enumerate(preds_full[:10], 1):
                    pr = pred_tbl.add_row().cells
                    pprice = pred.get('price') or 0
                    plower = pred.get('lower')
                    pupper = pred.get('upper')
                    pconf  = pred.get('confidence','')
                    pdate  = str(pred.get('date',''))[:10]
                    var_pct= ((pprice - current_price) / current_price * 100) if current_price else 0

                    # Couleur ligne
                    if   var_pct >  2: row_cl = 'C6EFCE'
                    elif var_pct < -2: row_cl = 'FFDCE0'
                    else:              row_cl = 'FAFAFA'

                    conf_stars = {'Élevée':'⭐⭐⭐','Moyen':'⭐⭐','Faible':'⭐'}.get(pconf, pconf)

                    vals_p = [
                        f"J+{ji}  {conf_stars}",
                        pdate,
                        f"{plower:,.0f}" if plower else '—',
                        f"{pprice:,.0f}",
                        f"{pupper:,.0f}" if pupper else '—',
                        f"{'+'if var_pct>=0 else ''}{var_pct:.1f}%",
                    ]
                    for ci, vp in enumerate(vals_p):
                        pr[ci].text = vp
                        run_p = pr[ci].paragraphs[0].runs[0] if pr[ci].paragraphs[0].runs else pr[ci].paragraphs[0].add_run(vp)
                        run_p.font.size = Pt(8.5)
                        shd_pr = OxmlElement('w:shd')
                        shd_pr.set(qn('w:fill'), row_cl)
                        shd_pr.set(qn('w:val'), 'clear')
                        pr[ci]._element.get_or_add_tcPr().append(shd_pr)
                        if ci == 3:   # Prix prédit en gras
                            run_p.bold = True
                        if ci == 5:   # Variation colorée
                            run_p.font.color.rgb = RGBColor(0,128,0) if var_pct >= 0 else RGBColor(192,0,0)
                            run_p.bold = True

                doc.add_paragraph()
                doc.add_paragraph(
                    "⚠️ Les prédictions sont des estimations statistiques (modèles GRU/LSTM). "
                    "L'intervalle de confiance à 90% indique que 9 fois sur 10, le cours devrait "
                    "se situer entre la borne basse et la borne haute. ⭐⭐⭐ = confiance élevée."
                ).runs[0].font.size = Pt(8)
                doc.add_paragraph()

            # ══════════════════════════════════════════════════════════════════
            # TABLEAU DONNÉES FINANCIÈRES STRUCTURÉES (brvm_donnees_financieres)
            # ══════════════════════════════════════════════════════════════════
            fin_data_word = self._get_donnees_financieres(symbol)
            if fin_data_word:
                annee_fin_w = fin_data_word.get('annee', 'N/A')

                # ── Détection secteur ────────────────────────────────────────
                is_bank_w = any([
                    fin_data_word.get('caisse_banque_centrale') and float(fin_data_word.get('caisse_banque_centrale') or 0) != 0,
                    fin_data_word.get('produit_net_bancaire')   and float(fin_data_word.get('produit_net_bancaire')   or 0) != 0,
                    fin_data_word.get('dettes_clientele')       and float(fin_data_word.get('dettes_clientele')       or 0) != 0,
                    fin_data_word.get('creances_interbancaires')and float(fin_data_word.get('creances_interbancaires')or 0) != 0,
                ])
                secteur_w = "🏦 BANQUE" if is_bank_w else "🏢 ENTREPRISE"

                # ── Helpers ──────────────────────────────────────────────────
                def _fv(v):
                    """Retourne (float_val, str_formaté) ou (None, None) si absent/zéro."""
                    if v is None or (isinstance(v, str) and v.strip() == ''):
                        return None, None
                    try:
                        f = float(v)
                        if f == 0.0: return None, None
                        s = (f"{f*100:.2f}%" if False  # placeholder
                             else f"{f/1_000_000_000:.3f} Mds FCFA" if abs(f) >= 1_000_000_000
                             else f"{f/1_000_000:.2f} M FCFA"       if abs(f) >= 1_000_000
                             else f"{f:,.2f} FCFA")
                        return f, s
                    except: return None, None

                def _fp(v):
                    """Retourne (float_val, str_pct) ou (None, None)."""
                    if v is None or (isinstance(v, str) and v.strip() == ''):
                        return None, None
                    try:
                        f = float(v)
                        if f == 0.0: return None, None
                        return f, f"{f*100:.2f}%"
                    except: return None, None

                def _fj(v):
                    """Retourne (float_val, str_jours) ou (None, None)."""
                    if v is None: return None, None
                    try:
                        f = float(v)
                        if f == 0.0: return None, None
                        return f, f"{f:.1f} j"
                    except: return None, None

                # Couleurs interprétation
                CLR_OK   = 'C6EFCE'  # vert clair — positif
                CLR_WARN = 'FFEB9C'  # orange clair — attention
                CLR_BAD  = 'FFC7CE'  # rouge clair — négatif
                CLR_INFO = 'DEEAF1'  # bleu clair — informatif
                CLR_HDR  = '1F4E79'  # bleu foncé — entête section

                # ── Construire les sections (label, valeur_str, interprétation, couleur_bg) ──
                sections_data = []

                # ══════════════════════════════════════════════════════════
                # SECTION 1 : BILAN — ACTIF
                # ══════════════════════════════════════════════════════════
                ba_rows = []

                fv, fs = _fv(fin_data_word.get('caisse_banque_centrale'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct_ta = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Représente {pct_ta:.1f}% du bilan. " +
                              ("✅ Réserves confortables, bonne capacité à honorer les réserves obligatoires." if pct_ta > 10
                               else "⚠️ Part faible — vérifier la conformité aux réserves obligatoires BCEAO."))
                    ba_rows.append(('Caisse & Banque Centrale', fs, interp, CLR_OK if pct_ta > 10 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('effets_publics'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = f"Titres d'État détenus ({pct:.1f}% du bilan). Actif très sécurisé, source de liquidité secondaire."
                    ba_rows.append(('Effets publics & val. assimilées', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('creances_interbancaires'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Prêts aux autres banques ({pct:.1f}% du bilan). " +
                              ("✅ Exposition modérée au risque interbancaire." if pct < 15
                               else "⚠️ Exposition significative — risque de contagion en cas de crise bancaire."))
                    ba_rows.append(('Créances interbancaires', fs, interp, CLR_OK if pct < 15 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('creances_clientele'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Portefeuille crédits ({pct:.1f}% du bilan). " +
                              ("✅ Concentration saine sur le métier de crédit." if 30 <= pct <= 70
                               else "⚠️ Concentration excessive (risque crédit élevé)." if pct > 70
                               else "ℹ️ Part faible — modèle orienté autres produits."))
                    ba_rows.append(('Créances sur la clientèle', fs, interp,
                                    CLR_OK if 30 <= pct <= 70 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('creances_clients'))
                if fv:
                    ca_f = float(fin_data_word.get('chiffre_affaires') or 0)
                    pct = (fv / ca_f * 100) if ca_f else 0
                    interp = (f"Créances clients = {pct:.1f}% du CA. " +
                              ("✅ Délais de paiement bien maîtrisés." if pct < 25
                               else "⚠️ Risque de retard de paiement significatif." if pct < 50
                               else "❌ Niveau critique — risque trésorerie élevé."))
                    ba_rows.append(('Créances clients', fs, interp,
                                    CLR_OK if pct < 25 else CLR_WARN if pct < 50 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('stocks'))
                if fv:
                    ca_f = float(fin_data_word.get('chiffre_affaires') or 0)
                    rot = (ca_f / fv) if fv else 0
                    interp = (f"Rotation stocks = {rot:.1f}x/an. " +
                              ("✅ Bonne rotation, stock bien dimensionné." if rot > 4
                               else "⚠️ Rotation lente — risque d'immobilisation de capital." if rot > 1
                               else "❌ Stock potentiellement obsolète ou surdimensionné."))
                    ba_rows.append(('Stocks', fs, interp,
                                    CLR_OK if rot > 4 else CLR_WARN if rot > 1 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('actif_circulant'))
                if fv:
                    pc_f = float(fin_data_word.get('passif_circulant') or 0)
                    ratio = (fv / pc_f) if pc_f else 0
                    interp = (f"Ratio actif circ./passif circ. = {ratio:.2f}x. " +
                              ("✅ Couverture satisfaisante des dettes court terme." if ratio >= 1.2
                               else "⚠️ Couverture limite." if ratio >= 1
                               else "❌ Actif circulant insuffisant pour couvrir les dettes CT."))
                    ba_rows.append(('Actif circulant', fs, interp,
                                    CLR_OK if ratio >= 1.2 else CLR_WARN if ratio >= 1 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('immobilisations_incorporelles'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = f"Brevets, logiciels, fonds commercial ({pct:.1f}% du bilan). Actif intangible à surveiller (risque dépréciation)."
                    ba_rows.append(('Immob. incorporelles', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('immobilisations_corporelles'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Outil industriel ({pct:.1f}% du bilan). " +
                              ("✅ Base productive solide." if pct > 20
                               else "ℹ️ Société à faible intensité capitalistique."))
                    ba_rows.append(('Immob. corporelles', fs, interp, CLR_OK if pct > 20 else CLR_INFO))

                fv, fs = _fv(fin_data_word.get('actif_immobilise_net'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = f"Actif immobilisé = {pct:.1f}% du total. Après amortissements — reflète la valeur nette comptable des actifs durables."
                    ba_rows.append(('Actif immobilisé net', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('tresorerie_actif'))
                if fv:
                    ta_f = float(fin_data_word.get('total_actif') or 0)
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Liquidités disponibles ({pct:.1f}% du bilan). " +
                              ("✅ Trésorerie confortable." if pct > 5
                               else "⚠️ Trésorerie tendue — surveiller les besoins CT."))
                    ba_rows.append(('Trésorerie Actif', fs, interp, CLR_OK if pct > 5 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('total_actif'))
                if fv:
                    interp = "Taille totale du bilan. Indicateur de la puissance financière de l'entreprise."
                    ba_rows.append(('Total Actif / Bilan', fs, interp, CLR_INFO))

                if ba_rows: sections_data.append(("📌 1. BILAN — ACTIF", ba_rows))

                # ══════════════════════════════════════════════════════════
                # SECTION 2 : BILAN — PASSIF
                # ══════════════════════════════════════════════════════════
                bp_rows = []
                ta_f = float(fin_data_word.get('total_actif') or 0)
                cp_f = float(fin_data_word.get('capitaux_propres') or 0)

                fv, fs = _fv(fin_data_word.get('capital_souscrit'))
                if fv:
                    interp = "Capital apporté par les actionnaires à la création ou lors d'augmentations de capital."
                    bp_rows.append(('Capital souscrit', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('reserves'))
                if fv:
                    interp = ("✅ Réserves accumulées significatives — capacité d'autofinancement historique." if fv > 0
                              else "⚠️ Réserves nulles ou négatives — capacité de résistance limitée.")
                    bp_rows.append(('Réserves', fs, interp, CLR_OK if fv > 0 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('capitaux_propres'))
                if fv:
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Fonds propres = {pct:.1f}% du bilan. " +
                              ("✅ Solidité financière satisfaisante." if pct > 20
                               else "⚠️ Fonds propres limités — levier financier élevé." if pct > 8
                               else "❌ Sous-capitalisation critique — risque de solvabilité."))
                    bp_rows.append(('Capitaux propres', fs, interp,
                                    CLR_OK if pct > 20 else CLR_WARN if pct > 8 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('capitaux_permanents'))
                if fv:
                    ai_f = float(fin_data_word.get('actif_immobilise_net') or 0)
                    ratio = (fv / ai_f) if ai_f else 0
                    interp = (f"Capitaux permanents / Actif immobilisé = {ratio:.2f}x. " +
                              ("✅ Les investissements LT sont bien financés par des ressources stables." if ratio >= 1
                               else "⚠️ Une partie des actifs LT est financée par des dettes CT — risque de liquidité."))
                    bp_rows.append(('Capitaux permanents', fs, interp, CLR_OK if ratio >= 1 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('dettes_interbancaires'))
                if fv:
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Refinancement interbancaire ({pct:.1f}% du bilan). " +
                              ("✅ Dépendance modérée au marché monétaire." if pct < 15
                               else "⚠️ Forte dépendance — risque de refinancement en cas de crise."))
                    bp_rows.append(('Dettes interbancaires', fs, interp, CLR_OK if pct < 15 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('dettes_clientele'))
                if fv:
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Dépôts clients = {pct:.1f}% du bilan. " +
                              ("✅ Base de dépôts solide — financement stable et peu coûteux." if pct > 50
                               else "ℹ️ Dépôts modérés — compléter par d'autres ressources."))
                    bp_rows.append(('Dettes clientèle (dépôts)', fs, interp, CLR_OK if pct > 50 else CLR_INFO))

                fv, fs = _fv(fin_data_word.get('dettes_fournisseurs'))
                if fv:
                    ca_f = float(fin_data_word.get('chiffre_affaires') or 0)
                    delai = (fv / ca_f * 365) if ca_f else 0
                    interp = (f"Délai fournisseurs implicite ≈ {delai:.0f} j. " +
                              ("✅ Bon pouvoir de négociation." if delai > 45
                               else "ℹ️ Délais standards de paiement."))
                    bp_rows.append(('Dettes fournisseurs', fs, interp, CLR_OK if delai > 45 else CLR_INFO))

                fv, fs = _fv(fin_data_word.get('dettes_financieres_lt_mt'))
                if fv:
                    pct = (fv / cp_f * 100) if cp_f else 0
                    interp = (f"Dette LT/MT = {pct:.0f}% des fonds propres. " +
                              ("✅ Endettement LT maîtrisé." if pct < 100
                               else "⚠️ Levier financier élevé — surveiller la capacité de remboursement."))
                    bp_rows.append(('Dettes financières LT/MT', fs, interp, CLR_OK if pct < 100 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('dettes_financieres_totales'))
                if fv:
                    pct = (fv / cp_f * 100) if cp_f else 0
                    interp = (f"Dette financière totale = {pct:.0f}% des fonds propres (gearing). " +
                              ("✅ Structure financière saine." if pct < 150
                               else "⚠️ Gearing élevé — risque financier accru." if pct < 300
                               else "❌ Sur-endettement — risque de défaillance."))
                    bp_rows.append(('Dettes financières totales', fs, interp,
                                    CLR_OK if pct < 150 else CLR_WARN if pct < 300 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('dettes_totales'))
                if fv:
                    pct = (fv / ta_f * 100) if ta_f else 0
                    interp = (f"Dettes totales = {pct:.1f}% du bilan. " +
                              ("✅ Taux d'endettement raisonnable." if pct < 70
                               else "⚠️ Endettement dominant — levier élevé." if pct < 85
                               else "❌ Très faible marge de sécurité pour les créanciers."))
                    bp_rows.append(('Dettes totales', fs, interp,
                                    CLR_OK if pct < 70 else CLR_WARN if pct < 85 else CLR_BAD))

                if bp_rows: sections_data.append(("📌 2. BILAN — PASSIF", bp_rows))

                # ══════════════════════════════════════════════════════════
                # SECTION 3 : COMPTE DE RÉSULTAT
                # ══════════════════════════════════════════════════════════
                cr_rows = []
                rn_f = float(fin_data_word.get('resultat_net') or 0)

                fv_pnb, fs_pnb = _fv(fin_data_word.get('produit_net_bancaire'))
                fv_ca,  fs_ca  = _fv(fin_data_word.get('chiffre_affaires'))
                rev_f = fv_pnb or fv_ca or 0

                fv, fs = _fv(fin_data_word.get('ca_pnb'))
                if fv:
                    cr_rows.append(('CA / PNB (agrégé)', fs, "Revenu principal de l'activité — base de tous les calculs de marge.", CLR_INFO))

                if fv_pnb:
                    tx_c = float(fin_data_word.get('taux_croissance_ca') or 0)
                    interp = (f"Revenu bancaire central (≡ CA). Croissance: {tx_c*100:.1f}%. " +
                              ("✅ PNB en progression — dynamisme commercial." if tx_c > 0.05
                               else "⚠️ Croissance faible ou nulle." if tx_c >= 0
                               else "❌ PNB en recul — pression sur la rentabilité."))
                    cr_rows.append(('Produit Net Bancaire (PNB)', fs_pnb, interp,
                                    CLR_OK if tx_c > 0.05 else CLR_WARN if tx_c >= 0 else CLR_BAD))

                if fv_ca:
                    tx_c = float(fin_data_word.get('taux_croissance_ca') or 0)
                    interp = (f"Chiffre d'affaires. Croissance: {tx_c*100:.1f}%. " +
                              ("✅ CA en hausse — bonne dynamique commerciale." if tx_c > 0.05
                               else "⚠️ Croissance limitée." if tx_c >= 0
                               else "❌ CA en baisse — perte de parts de marché ou demande en recul."))
                    cr_rows.append(("Chiffre d'affaires (CA)", fs_ca, interp,
                                    CLR_OK if tx_c > 0.05 else CLR_WARN if tx_c >= 0 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('interets_produits'))
                if fv and is_bank_w:
                    pct = (fv / (fv_pnb or 1)) * 100
                    interp = f"Revenus d'intérêts sur crédits ({pct:.0f}% du PNB). Principal moteur de revenus de la banque."
                    cr_rows.append(('Intérêts & produits assimilés', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('interets_charges'))
                if fv and is_bank_w:
                    pct = (fv / (fv_pnb or 1)) * 100
                    interp = (f"Coût des dépôts collectés ({pct:.0f}% du PNB). " +
                              ("✅ Coût de ressources maîtrisé." if pct < 40
                               else "⚠️ Charges d'intérêts élevées — pression sur les marges."))
                    cr_rows.append(('Intérêts & charges assimilées', fs, interp, CLR_OK if pct < 40 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('commissions_produits'))
                if fv:
                    pct = (fv / (rev_f or 1)) * 100
                    interp = f"Revenus de commissions ({pct:.1f}% des revenus). Diversification bienvenue du modèle."
                    cr_rows.append(('Commissions (produits)', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('charges_generales_exploitation'))
                if fv:
                    pct = (fv / (rev_f or 1)) * 100
                    interp = (f"{pct:.1f}% des revenus. " +
                              ("✅ Frais généraux maîtrisés." if pct < 30
                               else "⚠️ Frais généraux significatifs." if pct < 50
                               else "❌ Structure de coûts lourde — inefficacité opérationnelle."))
                    cr_rows.append(("Charges générales d'exploit.", fs, interp,
                                    CLR_OK if pct < 30 else CLR_WARN if pct < 50 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('dap_immobilisations'))
                if fv:
                    interp = "Dotations aux amortissements — charge non décaissée reflétant l'usure des actifs."
                    cr_rows.append(('DAP immobilisations', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('charges_personnel'))
                if fv:
                    pct = (fv / (rev_f or 1)) * 100
                    interp = (f"Masse salariale = {pct:.1f}% des revenus. " +
                              ("✅ Productivité par employé satisfaisante." if pct < 25
                               else "⚠️ Masse salariale importante." if pct < 45
                               else "❌ Charges salariales très élevées — pression sur la rentabilité."))
                    cr_rows.append(('Charges du personnel', fs, interp,
                                    CLR_OK if pct < 25 else CLR_WARN if pct < 45 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('valeur_ajoutee'))
                if fv:
                    pct = (fv / (fv_ca or 1)) * 100
                    interp = (f"VA = {pct:.1f}% du CA. " +
                              ("✅ Forte valeur ajoutée — activité à haute marge brute." if pct > 40
                               else "ℹ️ VA standard pour le secteur." if pct > 20
                               else "⚠️ Faible VA — modèle à faible marge (distribution, négoce)."))
                    cr_rows.append(('Valeur Ajoutée (VA)', fs, interp,
                                    CLR_OK if pct > 40 else CLR_INFO if pct > 20 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('ebe'))
                if fv:
                    pct = (fv / (rev_f or 1)) * 100
                    interp = (f"EBE = {pct:.1f}% des revenus (marge opérationnelle brute). " +
                              ("✅ Excellente capacité bénéficiaire avant amort." if pct > 25
                               else "✅ Bonne rentabilité opérationnelle." if pct > 10
                               else "⚠️ Marge opérationnelle faible." if pct > 0
                               else "❌ EBE négatif — activité structurellement déficitaire."))
                    cr_rows.append(('EBE', fs, interp,
                                    CLR_OK if pct > 10 else CLR_WARN if pct > 0 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('rbe'))
                if fv and is_bank_w:
                    pct = (fv / (fv_pnb or 1)) * 100
                    interp = (f"RBE = {pct:.1f}% du PNB (PNB - CGE - DAP). " +
                              ("✅ Efficacité opérationnelle bancaire solide." if pct > 50
                               else "⚠️ Efficacité limitée — charges absorb. une large part du PNB."))
                    cr_rows.append(('RBE (Résultat Brut Exploit.)', fs, interp, CLR_OK if pct > 50 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('resultat_exploitation'))
                if fv:
                    pct = (fv / (rev_f or 1)) * 100
                    interp = (f"Profit d'exploitation = {pct:.1f}% des revenus. " +
                              ("✅ Rentabilité opérationnelle solide." if fv > 0 and pct > 5
                               else "⚠️ Bénéfice d'exploitation limité." if fv > 0
                               else "❌ Perte d'exploitation — l'activité principale ne couvre pas ses coûts."))
                    cr_rows.append(("Résultat d'exploitation", fs, interp,
                                    CLR_OK if fv > 0 and pct > 5 else CLR_WARN if fv > 0 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('provisions'))
                if fv:
                    interp = ("Provisions pour risques/créances douteuses. Charge précautionnelle qui réduit le bénéfice imposable." +
                              (" Surveiller la tendance — hausse = dégradation du risque." if is_bank_w else ""))
                    cr_rows.append(('Provisions', fs, interp, CLR_INFO))

                fv, fs = _fv(fin_data_word.get('resultat_net'))
                if fv:
                    pct = (fv / (rev_f or 1)) * 100
                    interp = (f"Marge nette = {pct:.1f}%. " +
                              ("✅ Excellente rentabilité finale." if fv > 0 and pct > 10
                               else "✅ Bénéficiaire." if fv > 0
                               else "❌ Perte nette — impact négatif sur les fonds propres et les dividendes."))
                    cr_rows.append(('Résultat net', fs, interp,
                                    CLR_OK if fv > 0 and pct > 10 else CLR_WARN if fv > 0 else CLR_BAD))

                if cr_rows: sections_data.append(("📌 3. COMPTE DE RÉSULTAT", cr_rows))

                # ══════════════════════════════════════════════════════════
                # SECTION 4 : CASH-FLOWS & TRÉSORERIE
                # ══════════════════════════════════════════════════════════
                cf_rows = []
                rn_f2 = float(fin_data_word.get('resultat_net') or 0)

                fv, fs = _fv(fin_data_word.get('caf'))
                if fv:
                    ratio = (fv / rn_f2) if rn_f2 else 0
                    interp = (f"CAF = {fv/1e9:.2f} Mds FCFA (résultat + amort.). " +
                              ("✅ Forte capacité à financer les investissements et rembourser les dettes." if fv > 0
                               else "❌ CAF négative — la société consomme des liquidités sans en générer."))
                    cf_rows.append(('CAF', fs, interp, CLR_OK if fv > 0 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('flux_operationnel'))
                if fv:
                    interp = (f"Cash généré par l'activité courante. " +
                              ("✅ Activité génératrice de cash — bonne santé opérationnelle." if fv > 0
                               else "❌ Activité consommatrice de cash — modèle sous pression."))
                    cf_rows.append(('Flux opérationnel', fs, interp, CLR_OK if fv > 0 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('flux_investissement'))
                if fv:
                    interp = (f"Sorties nettes pour investissements. " +
                              ("✅ Investissement actif — développement de l'outil productif." if fv < 0
                               else "ℹ️ Cessions nettes — désinvestissement ou rééquilibrage du portefeuille."))
                    cf_rows.append(("Flux d'investissement", fs, interp, CLR_OK if fv < 0 else CLR_INFO))

                fv, fs = _fv(fin_data_word.get('flux_financement'))
                if fv:
                    interp = (f"Flux liés à la dette et aux dividendes. " +
                              ("✅ Remboursements nets — désendettement en cours." if fv < 0
                               else "ℹ️ Nouveaux financements levés — augmentation de capital ou emprunt."))
                    cf_rows.append(("Flux de financement", fs, interp, CLR_OK if fv < 0 else CLR_INFO))

                fv, fs = _fv(fin_data_word.get('free_cash_flow'))
                if fv:
                    interp = (f"FCF = flux opé. - capex. " +
                              ("✅ FCF positif — la société dégage du cash après ses investissements." if fv > 0
                               else "⚠️ FCF négatif — investissements supérieurs au cash opérationnel."))
                    cf_rows.append(('Free Cash Flow', fs, interp, CLR_OK if fv > 0 else CLR_WARN))

                fv, fs = _fv(fin_data_word.get('bfr'))
                if fv:
                    rev_ref = float(fin_data_word.get('chiffre_affaires') or 0)
                    pct = (fv / rev_ref * 100) if rev_ref else 0
                    interp = (f"BFR = {pct:.1f}% du CA. " +
                              ("✅ BFR négatif = la société se finance sur ses clients (modèle distributeur)." if fv < 0
                               else "✅ BFR modéré — cycle d'exploitation bien financé." if pct < 15
                               else "⚠️ BFR élevé — besoin de financement CT important." if pct < 30
                               else "❌ BFR critique — risque de tension de trésorerie."))
                    cf_rows.append(('BFR', fs, interp,
                                    CLR_OK if fv < 0 or pct < 15 else CLR_WARN if pct < 30 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('fonds_roulement'))
                if fv:
                    interp = (f"FR = capitaux permanents - actif immobilisé. " +
                              ("✅ Fonds de roulement positif — les ressources LT financent l'actif LT avec marge." if fv > 0
                               else "❌ FR négatif — une partie des immobilisations est financée par des dettes CT."))
                    cf_rows.append(('Fonds de Roulement', fs, interp, CLR_OK if fv > 0 else CLR_BAD))

                fv, fs = _fv(fin_data_word.get('tresorerie_nette'))
                if fv:
                    interp = (f"Trésorerie nette = FR - BFR. " +
                              ("✅ Position de trésorerie nette excédentaire." if fv > 0
                               else "⚠️ Trésorerie nette négative — recours aux concours bancaires." if fv > -1e9
                               else "❌ Trésorerie structurellement déficitaire — risque de liquidité."))
                    cf_rows.append(('Trésorerie nette', fs, interp,
                                    CLR_OK if fv > 0 else CLR_WARN if fv > -1e9 else CLR_BAD))

                if cf_rows: sections_data.append(("📌 4. CASH-FLOWS & TRÉSORERIE", cf_rows))

                # ══════════════════════════════════════════════════════════
                # SECTION 5 : RATIOS DE RENTABILITÉ
                # ══════════════════════════════════════════════════════════
                rr_rows = []

                fv, fs = _fp(fin_data_word.get('marge_brute'))
                if fv:
                    interp = (f"(CA - coûts directs) / CA. " +
                              ("✅ Excellente marge brute." if fv > 0.40
                               else "✅ Marge brute correcte." if fv > 0.20
                               else "⚠️ Marge brute limitée — activité à faible valeur ajoutée." if fv > 0
                               else "❌ Marge brute négative — coût de revient supérieur au prix de vente."))
                    rr_rows.append(('Marge brute', fs, interp,
                                    CLR_OK if fv > 0.20 else CLR_WARN if fv > 0 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('marge_nette'))
                if fv:
                    interp = (f"Résultat net / revenus = {fv*100:.2f}%. " +
                              ("✅ Excellente profitabilité finale." if fv > 0.10
                               else "✅ Profitabilité satisfaisante." if fv > 0.03
                               else "⚠️ Marge nette très fine — sensible aux aléas." if fv > 0
                               else "❌ Déficitaire — la société détruit de la valeur."))
                    rr_rows.append(('Marge nette', fs, interp,
                                    CLR_OK if fv > 0.03 else CLR_WARN if fv > 0 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('marge_operationnelle'))
                if fv:
                    interp = (f"Rentabilité de l'exploitation pure. " +
                              ("✅ Excellente marge opérationnelle." if fv > 0.15
                               else "✅ Marge opérationnelle correcte." if fv > 0.05
                               else "⚠️ Marge opérationnelle faible." if fv > 0
                               else "❌ Exploitation déficitaire."))
                    rr_rows.append(('Marge opérationnelle', fs, interp,
                                    CLR_OK if fv > 0.05 else CLR_WARN if fv > 0 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('roe'))
                if fv:
                    interp = (f"Résultat net / capitaux propres = {fv*100:.2f}%. " +
                              ("✅ Excellent — au-dessus du coût du capital UEMOA." if fv > 0.15
                               else "✅ Satisfaisant." if fv > 0.08
                               else "⚠️ Faible rémunération des actionnaires." if fv > 0
                               else "❌ ROE négatif — destruction de valeur actionnaire."))
                    rr_rows.append(('ROE (Rentabilité fonds propres)', fs, interp,
                                    CLR_OK if fv > 0.08 else CLR_WARN if fv > 0 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('roa'))
                if fv:
                    bench = "banque: seuil >1%" if is_bank_w else "entreprise: seuil >5%"
                    seuil = 0.01 if is_bank_w else 0.05
                    interp = (f"Résultat net / total actif = {fv*100:.2f}% ({bench}). " +
                              ("✅ Bonne utilisation des actifs." if fv > seuil
                               else "⚠️ Rentabilité des actifs en deçà du seuil sectoriel." if fv > 0
                               else "❌ ROA négatif — actifs non rentables."))
                    rr_rows.append(('ROA (Rentabilité des actifs)', fs, interp,
                                    CLR_OK if fv > seuil else CLR_WARN if fv > 0 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('rotation_actifs'))
                if fv:
                    interp = (f"CA / total actif = {fv*100:.2f}x. " +
                              ("✅ Bonne efficience des actifs." if fv > 0.50
                               else "⚠️ Actifs peu utilisés." if fv > 0.20
                               else "ℹ️ Faible rotation — secteur à forte intensité capitalistique (normal)."))
                    rr_rows.append(('Rotation des actifs', fs, interp,
                                    CLR_OK if fv > 0.50 else CLR_INFO if fv > 0.20 else CLR_WARN))

                fv, fs = _fp(fin_data_word.get('coefficient_exploitation'))
                if fv and is_bank_w:
                    interp = (f"Charges / PNB = {fv*100:.2f}% (cost-to-income). " +
                              ("✅ Excellent — banque très efficiente." if fv < 0.50
                               else "✅ Efficiente (standard UEMOA)." if fv < 0.60
                               else "⚠️ Efficience limitée — charges absorbent trop du PNB." if fv < 0.75
                               else "❌ Coeff. critique — structure de coûts non viable."))
                    rr_rows.append(("Coeff. d'exploitation (banque)", fs, interp,
                                    CLR_OK if fv < 0.60 else CLR_WARN if fv < 0.75 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('cout_risque'))
                if fv is not None:
                    afv = abs(fv)
                    interp = (f"Provisions / charges financières. " +
                              ("✅ Coût du risque sain — portefeuille de qualité." if afv < 0.01
                               else "⚠️ Coût du risque modéré — surveillance requise." if afv < 0.03
                               else "❌ Coût du risque élevé — dégradation significative du portefeuille."))
                    rr_rows.append(('Coût du risque', fs, interp,
                                    CLR_OK if afv < 0.01 else CLR_WARN if afv < 0.03 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('taux_croissance_ca'))
                if fv is not None:
                    lbl = "PNB" if is_bank_w else "CA"
                    interp = (f"Croissance {lbl} vs exercice précédent = {fv*100:.2f}%. " +
                              ("✅ Forte croissance — expansion commerciale." if fv > 0.10
                               else "✅ Croissance positive." if fv > 0
                               else "⚠️ Activité stagnante." if fv > -0.05
                               else "❌ Recul significatif du revenu principal."))
                    rr_rows.append((f'Taux croissance {lbl}', fs, interp,
                                    CLR_OK if fv > 0 else CLR_WARN if fv > -0.05 else CLR_BAD))

                if rr_rows: sections_data.append(("📌 5. RATIOS DE RENTABILITÉ", rr_rows))

                # ══════════════════════════════════════════════════════════
                # SECTION 6 : STRUCTURE & LIQUIDITÉ
                # ══════════════════════════════════════════════════════════
                rs_rows = []

                fv, fs = _fp(fin_data_word.get('autonomie_financiere'))
                if fv:
                    interp = (f"Fonds propres / total bilan = {fv*100:.2f}%. " +
                              ("✅ Solide indépendance financière." if fv > 0.30
                               else "✅ Autonomie correcte." if fv > 0.15
                               else "⚠️ Dépendance aux tiers significative." if fv > 0.08
                               else "❌ Quasi-totalité des actifs financée par des dettes."))
                    rs_rows.append(('Autonomie financière', fs, interp,
                                    CLR_OK if fv > 0.15 else CLR_WARN if fv > 0.08 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('dependance_financiere'))
                if fv:
                    interp = (f"Dettes financières / total bilan = {fv*100:.2f}%. " +
                              ("✅ Dépendance faible." if fv < 0.30
                               else "⚠️ Dépendance notable aux financeurs externes." if fv < 0.60
                               else "❌ Structure très endettée — vulnérable aux chocs de taux."))
                    rs_rows.append(('Dépendance financière', fs, interp,
                                    CLR_OK if fv < 0.30 else CLR_WARN if fv < 0.60 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('ratio_endettement'))
                if fv:
                    interp = (f"Dettes totales / fonds propres = {fv*100:.2f}% (gearing). " +
                              ("✅ Levier financier sain." if fv < 1.0
                               else "⚠️ Levier modéré — à surveiller." if fv < 2.5
                               else "❌ Surendettement — risque financier élevé."))
                    rs_rows.append(("Ratio d'endettement (gearing)", fs, interp,
                                    CLR_OK if fv < 1.0 else CLR_WARN if fv < 2.5 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('solvabilite_generale'))
                if fv:
                    interp = (f"Total actif / dettes totales = {fv*100:.2f}x. " +
                              ("✅ Excellente solvabilité — actifs largement supérieurs aux dettes." if fv > 1.5
                               else "✅ Solvable." if fv > 1.0
                               else "❌ Insolvable — dettes supérieures aux actifs."))
                    rs_rows.append(('Solvabilité générale', fs, interp,
                                    CLR_OK if fv > 1.0 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('liquidite_generale'))
                if fv:
                    interp = (f"Actif circulant / passif circulant = {fv*100:.2f}x. " +
                              ("✅ Excellente liquidité court terme." if fv > 1.5
                               else "✅ Liquidité satisfaisante." if fv > 1.0
                               else "⚠️ Liquidité tendue — risque si remboursement simultané exigé." if fv > 0.7
                               else "❌ Liquidité critique — incapacité probable à honorer les dettes CT."))
                    rs_rows.append(('Liquidité générale', fs, interp,
                                    CLR_OK if fv > 1.0 else CLR_WARN if fv > 0.7 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('liquidite_immediate'))
                if fv:
                    interp = (f"Trésorerie / passif circulant = {fv*100:.2f}x. " +
                              ("✅ Trésorerie immédiate confortable." if fv > 0.30
                               else "⚠️ Trésorerie immédiate limitée." if fv > 0.10
                               else "❌ Cash quasi inexistant face aux dettes CT."))
                    rs_rows.append(('Liquidité immédiate', fs, interp,
                                    CLR_OK if fv > 0.30 else CLR_WARN if fv > 0.10 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('liquidite_reduite'))
                if fv:
                    interp = (f"(Actif circ. - stocks) / passif circ. = {fv*100:.2f}x. " +
                              ("✅ Liquidité réduite solide." if fv > 1.0
                               else "⚠️ Liquidité réduite limite." if fv > 0.7
                               else "❌ Dépendance aux stocks pour couvrir les dettes CT."))
                    rs_rows.append(('Liquidité réduite', fs, interp,
                                    CLR_OK if fv > 1.0 else CLR_WARN if fv > 0.7 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('financement_immobilisations'))
                if fv:
                    interp = (f"Capitaux permanents / actif immobilisé = {fv*100:.2f}x. " +
                              ("✅ Les actifs LT sont intégralement financés par des ressources stables." if fv >= 1.0
                               else "❌ Partie des actifs LT financée par des ressources CT — risque de liquidité."))
                    rs_rows.append(('Financement des immob.', fs, interp, CLR_OK if fv >= 1.0 else CLR_BAD))

                fv, fs = _fj(fin_data_word.get('capacite_remboursement'))
                if fv:
                    interp = (f"Dettes financières / CAF = {fv:.1f} ans. " +
                              ("✅ Dette remboursable rapidement." if fv < 3
                               else "✅ Délai raisonnable." if fv < 5
                               else "⚠️ Remboursement long — endettement fort par rapport au cash." if fv < 8
                               else "❌ Capacité de remboursement très dégradée."))
                    rs_rows.append(('Capacité de remboursement', f"{fv:.1f} ans", interp,
                                    CLR_OK if fv < 5 else CLR_WARN if fv < 8 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('couverture_interets'))
                if fv:
                    interp = (f"Résultat exploit. / charges financières = {fv*100:.2f}x. " +
                              ("✅ Excellente couverture des intérêts." if fv > 3
                               else "✅ Couverture satisfaisante." if fv > 1.5
                               else "⚠️ Couverture limite — sensible aux hausses de taux." if fv > 1
                               else "❌ L'exploitation ne couvre pas les charges financières."))
                    rs_rows.append(("Couverture des intérêts", fs, interp,
                                    CLR_OK if fv > 1.5 else CLR_WARN if fv > 1 else CLR_BAD))

                fv, fs = _fp(fin_data_word.get('couverture_investissements_caf'))
                if fv:
                    interp = (f"CAF / investissements = {fv*100:.2f}x. " +
                              ("✅ La CAF finance largement les investissements." if fv > 1
                               else "⚠️ Financement externe nécessaire pour les investissements."))
                    rs_rows.append(('Couverture invest./CAF', fs, interp, CLR_OK if fv > 1 else CLR_WARN))

                if rs_rows: sections_data.append(("📌 6. STRUCTURE & LIQUIDITÉ", rs_rows))

                # ══════════════════════════════════════════════════════════
                # SECTION 7 : DÉLAIS D'EXPLOITATION
                # ══════════════════════════════════════════════════════════
                dl_rows = []

                fv, fs = _fj(fin_data_word.get('delai_clients'))
                if fv:
                    interp = (f"Les clients paient en moyenne en {fv:.0f} jours. " +
                              ("✅ Recouvrement rapide — trésorerie favorisée." if fv < 30
                               else "✅ Délai standard." if fv < 60
                               else "⚠️ Délai long — risque d'impayés et tension de trésorerie." if fv < 90
                               else "❌ Délai critique — process de recouvrement à revoir."))
                    dl_rows.append(('Délai clients (jours)', fs, interp,
                                    CLR_OK if fv < 60 else CLR_WARN if fv < 90 else CLR_BAD))

                fv, fs = _fj(fin_data_word.get('delai_fournisseurs'))
                if fv:
                    interp = (f"Paiement des fournisseurs en {fv:.0f} jours. " +
                              ("✅ Délai favorable — la société conserve sa trésorerie." if fv > 60
                               else "✅ Délai standard." if fv > 30
                               else "⚠️ Paiement rapide — pouvoir de négociation limité."))
                    dl_rows.append(('Délai fournisseurs (jours)', fs, interp,
                                    CLR_OK if fv > 30 else CLR_WARN))

                fv, fs = _fj(fin_data_word.get('duree_stockage'))
                if fv:
                    interp = (f"Stock écoulé en {fv:.0f} jours. " +
                              ("✅ Rotation rapide — faible immobilisation de capital." if fv < 30
                               else "✅ Stockage standard." if fv < 60
                               else "⚠️ Durée de stockage longue — capital immobilisé." if fv < 120
                               else "❌ Stock quasi-bloqué — risque d'obsolescence ou de surstock."))
                    dl_rows.append(('Durée de stockage (jours)', fs, interp,
                                    CLR_OK if fv < 60 else CLR_WARN if fv < 120 else CLR_BAD))

                if dl_rows: sections_data.append(("📌 7. DÉLAIS D'EXPLOITATION", dl_rows))

                # ══════════════════════════════════════════════════════════
                # GÉNÉRATION DU TABLEAU WORD À 3 COLONNES
                # ══════════════════════════════════════════════════════════
                if sections_data:
                    doc.add_paragraph()
                    h_fin = doc.add_heading(
                        f"📊 DONNÉES FINANCIÈRES STRUCTURÉES & INTERPRÉTÉES — Exercice {annee_fin_w} [{secteur_w}]",
                        level=3
                    )
                    h_fin.runs[0].font.color.rgb = RGBColor(0, 70, 127)
                    h_fin.runs[0].font.size = Pt(11)

                    note_p = doc.add_paragraph()
                    note_r = note_p.add_run(
                        "Source : brvm_donnees_financieres — chiffres officiels. "
                        "🟢 = signal positif  🟡 = vigilance  🔴 = signal négatif  🔵 = informatif."
                    )
                    note_r.font.size = Pt(8)
                    note_r.font.italic = True
                    note_r.font.color.rgb = RGBColor(80, 80, 80)

                    for section_title, rows in sections_data:
                        # Titre de section
                        sec_p = doc.add_paragraph()
                        sec_p.paragraph_format.space_before = Pt(10)
                        sec_p.paragraph_format.space_after  = Pt(2)
                        sec_r = sec_p.add_run(section_title)
                        sec_r.bold = True
                        sec_r.font.size = Pt(10)
                        sec_r.font.color.rgb = RGBColor(0, 51, 102)

                        # Tableau 3 colonnes : Indicateur | Valeur | Interprétation
                        tbl_fin = doc.add_table(rows=1, cols=3)
                        tbl_fin.style = 'Table Grid'

                        # En-tête
                        hdr = tbl_fin.rows[0].cells
                        for ci, htxt in enumerate(['Indicateur', 'Valeur', 'Interprétation']):
                            hdr[ci].text = htxt
                            r = hdr[ci].paragraphs[0].runs[0] if hdr[ci].paragraphs[0].runs else hdr[ci].paragraphs[0].add_run(htxt)
                            r.bold = True
                            r.font.size = Pt(8.5)
                            r.font.color.rgb = RGBColor(255, 255, 255)
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:fill'), CLR_HDR)
                            shd.set(qn('w:val'), 'clear')
                            hdr[ci]._element.get_or_add_tcPr().append(shd)

                        # Lignes de données
                        for lbl, val_s, interp_s, bg_color in rows:
                            tr = tbl_fin.add_row().cells

                            # Col 0 : Indicateur
                            tr[0].text = lbl
                            r0 = tr[0].paragraphs[0].runs[0] if tr[0].paragraphs[0].runs else tr[0].paragraphs[0].add_run(lbl)
                            r0.font.size = Pt(8)
                            r0.bold = True
                            r0.font.color.rgb = RGBColor(30, 30, 60)
                            shd0 = OxmlElement('w:shd')
                            shd0.set(qn('w:fill'), 'EEF2F7')
                            shd0.set(qn('w:val'), 'clear')
                            tr[0]._element.get_or_add_tcPr().append(shd0)

                            # Col 1 : Valeur — fond coloré selon signal
                            tr[1].text = val_s
                            r1 = tr[1].paragraphs[0].runs[0] if tr[1].paragraphs[0].runs else tr[1].paragraphs[0].add_run(val_s)
                            r1.font.size = Pt(8.5)
                            r1.bold = True
                            # Couleur texte valeur
                            try:
                                num_chk = float(val_s.replace('%','').replace(' Mds FCFA','').replace(' M FCFA','').replace(' FCFA','').replace(' j','').replace(' ans','').replace(',','').strip())
                                r1.font.color.rgb = RGBColor(0,100,0) if num_chk > 0 else RGBColor(180,0,0) if num_chk < 0 else RGBColor(80,80,80)
                            except:
                                r1.font.color.rgb = RGBColor(0,0,0)
                            shd1 = OxmlElement('w:shd')
                            shd1.set(qn('w:fill'), bg_color)
                            shd1.set(qn('w:val'), 'clear')
                            tr[1]._element.get_or_add_tcPr().append(shd1)

                            # Col 2 : Interprétation
                            tr[2].text = interp_s
                            r2 = tr[2].paragraphs[0].runs[0] if tr[2].paragraphs[0].runs else tr[2].paragraphs[0].add_run(interp_s)
                            r2.font.size = Pt(7.5)
                            r2.font.color.rgb = RGBColor(40, 40, 40)
                            shd2 = OxmlElement('w:shd')
                            shd2.set(qn('w:fill'), 'FFFFFF')
                            shd2.set(qn('w:val'), 'clear')
                            tr[2]._element.get_or_add_tcPr().append(shd2)

                        # Largeurs colonnes : 4cm | 3.5cm | 10cm
                        try:
                            col_widths = [Cm(4.0), Cm(3.5), Cm(10.0)]
                            for col_i, col in enumerate(tbl_fin.columns):
                                for cell in col.cells:
                                    cell.width = col_widths[col_i]
                        except: pass

                    doc.add_paragraph()

            paragraphs = analysis.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    
                    if para_text.startswith('**PARTIE') or para_text.startswith('**CONCLUSION'):
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(12)
                        p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                        p.paragraph_format.space_before = Pt(12)
            
            doc.add_paragraph()

            # ═══════ SECTION brvm_documents ═══════════════════════════════════════
            brvm_docs_raw = company_data.get('brvm_docs_raw', [])
            if not brvm_docs_raw:
                # Fallback: chercher dans all_analyses via data_dict (cas où non propagé)
                pass
            if brvm_docs_raw:
                doc.add_paragraph()
                hdg = doc.add_heading('📎 DOCUMENTS & COMMUNICATIONS OFFICIELS BRVM', level=3)
                hdg.runs[0].font.color.rgb = RGBColor(0, 70, 127)

                nb_docs_display = min(len(brvm_docs_raw), 5)
                intro_p = doc.add_paragraph()
                intro_run = intro_p.add_run(
                    f"{len(brvm_docs_raw)} document(s) disponible(s) pour cette société "
                    f"(affichage des {nb_docs_display} plus récents)."
                )
                intro_run.font.size = Pt(9)
                intro_run.font.italic = True
                intro_run.font.color.rgb = RGBColor(100, 100, 100)

                for doc_i, raw_doc in enumerate(brvm_docs_raw[:5], 1):
                    d = self._format_brvm_documents_for_word(raw_doc, doc_i)

                    # En-tête du document
                    doc_hdr = doc.add_paragraph()
                    doc_hdr.paragraph_format.space_before = Pt(8)
                    run_num = doc_hdr.add_run(f"[{d['num']}] ")
                    run_num.bold = True
                    run_num.font.size = Pt(10)
                    run_num.font.color.rgb = RGBColor(0, 102, 204)
                    run_titre = doc_hdr.add_run(d['titre'])
                    run_titre.bold = True
                    run_titre.font.size = Pt(10)

                    # Méta-ligne : date + catégorie + impact
                    meta_p = doc.add_paragraph()
                    meta_p.paragraph_format.left_indent = Pt(20)
                    meta_p.paragraph_format.space_before = Pt(1)
                    meta_run = meta_p.add_run(
                        f"Date : {d['date_doc']}  |  Catégorie : {d['categorie']}  |  Impact : "
                    )
                    meta_run.font.size = Pt(8.5)
                    meta_run.font.color.rgb = RGBColor(80, 80, 80)
                    impact_run = meta_p.add_run(d['impact'])
                    impact_run.font.size = Pt(8.5)
                    impact_run.bold = True
                    impact_run.font.color.rgb = d['impact_color']

                    # Résumé
                    if d['resume']:
                        resume_p = doc.add_paragraph()
                        resume_p.paragraph_format.left_indent = Pt(20)
                        resume_p.paragraph_format.space_before = Pt(2)
                        resume_run = resume_p.add_run(d['resume'][:700])
                        resume_run.font.size = Pt(9)
                        if len(d['resume']) > 700:
                            resume_p.add_run("…").font.size = Pt(9)

                    # Points clés (si disponibles)
                    if d['points']:
                        pk_hdr = doc.add_paragraph()
                        pk_hdr.paragraph_format.left_indent = Pt(20)
                        pk_hdr.paragraph_format.space_before = Pt(3)
                        pk_hdr.add_run("Points clés :").bold = True
                        pk_hdr.runs[-1].font.size = Pt(9)
                        for pt in d['points']:
                            pk_p = doc.add_paragraph(style='List Bullet')
                            pk_p.paragraph_format.left_indent = Pt(35)
                            pk_p.paragraph_format.space_before = Pt(1)
                            pk_p.paragraph_format.space_after = Pt(1)
                            pk_run = pk_p.add_run(str(pt)[:200])
                            pk_run.font.size = Pt(9)

                    # Séparateur léger entre documents
                    if doc_i < nb_docs_display:
                        sep_p = doc.add_paragraph()
                        sep_p.paragraph_format.left_indent = Pt(20)
                        sep_run = sep_p.add_run("· " * 30)
                        sep_run.font.size = Pt(7)
                        sep_run.font.color.rgb = RGBColor(180, 180, 180)


            # ═══════ SECTION brvm_rapports_societes ═══════════════════════════════
            brvm_rapports_raw = company_data.get('brvm_rapports_raw', [])
            if brvm_rapports_raw:
                doc.add_paragraph()
                hdg2 = doc.add_heading('📋 RAPPORTS & COMMUNIQUÉS OFFICIELS', level=3)
                hdg2.runs[0].font.color.rgb = RGBColor(0, 90, 50)

                nb_rap_display = min(len(brvm_rapports_raw), 5)
                intro2 = doc.add_paragraph()
                ir = intro2.add_run(
                    f"{len(brvm_rapports_raw)} rapport(s)/communiqué(s) disponible(s) "
                    f"(affichage des {nb_rap_display} plus récents)."
                )
                ir.font.size = Pt(9)
                ir.font.italic = True
                ir.font.color.rgb = RGBColor(100, 100, 100)

                for ri, raw_rap in enumerate(brvm_rapports_raw[:5], 1):
                    r = self._format_rapports_societes_for_word(raw_rap, ri)

                    # ── En-tête ──────────────────────────────────────────────────
                    rhdr = doc.add_paragraph()
                    rhdr.paragraph_format.space_before = Pt(8)
                    run_ri = rhdr.add_run(f"[{r['idx']}] ")
                    run_ri.bold = True
                    run_ri.font.size = Pt(10)
                    run_ri.font.color.rgb = RGBColor(0, 120, 60)
                    run_rt = rhdr.add_run(r['doc_titre'])
                    run_rt.bold = True
                    run_rt.font.size = Pt(10)

                    # ── Méta ─────────────────────────────────────────────────────
                    rmeta = doc.add_paragraph()
                    rmeta.paragraph_format.left_indent = Pt(20)
                    rmeta.paragraph_format.space_before = Pt(1)
                    meta_txt = f"Année : {r['annee']}"
                    if r['type_rapport']:
                        meta_txt += f"  |  Type : {r['type_rapport']}"
                    if r['date_rapport']:
                        meta_txt += f"  |  Date : {r['date_rapport']}"
                    rmr = rmeta.add_run(meta_txt)
                    rmr.font.size = Pt(8.5)
                    rmr.font.color.rgb = RGBColor(80, 80, 80)

                    # ── Recommandation source ────────────────────────────────────
                    if r['recommandation']:
                        rrec = doc.add_paragraph()
                        rrec.paragraph_format.left_indent = Pt(20)
                        rrec.paragraph_format.space_before = Pt(1)
                        rrec.add_run("Recommandation source : ").font.size = Pt(9)
                        rrec.runs[-1].bold = True
                        rec_run = rrec.add_run(r['recommandation'])
                        rec_run.font.size = Pt(9)
                        rec_run.bold = True
                        rec_run.font.color.rgb = r['rec_color']

                    # ── Résumé ───────────────────────────────────────────────────
                    if r['resume']:
                        rres = doc.add_paragraph()
                        rres.paragraph_format.left_indent = Pt(20)
                        rres.paragraph_format.space_before = Pt(2)
                        rres_run = rres.add_run(r['resume'][:700])
                        rres_run.font.size = Pt(9)
                        if len(r['resume']) > 700:
                            rres.add_run("…").font.size = Pt(9)

                    # ── Indicateurs financiers ───────────────────────────────────
                    if r['indicateurs']:
                        rind_hdr = doc.add_paragraph()
                        rind_hdr.paragraph_format.left_indent = Pt(20)
                        rind_hdr.paragraph_format.space_before = Pt(3)
                        rind_hdr_run = rind_hdr.add_run("Indicateurs financiers :")
                        rind_hdr_run.bold = True
                        rind_hdr_run.font.size = Pt(9)

                        # Tableau compact indicateurs
                        ind_items = list(r['indicateurs'].items())
                        if ind_items:
                            ind_tbl = doc.add_table(rows=1, cols=min(len(ind_items), 4))
                            ind_tbl.style = 'Light Grid Accent 1'
                            ind_tbl.paragraph_format = None
                            hcells = ind_tbl.rows[0].cells
                            for ci, (k, v) in enumerate(ind_items[:4]):
                                hcells[ci].text = ''
                                ph = hcells[ci].paragraphs[0]
                                ph.paragraph_format.left_indent = Pt(5)
                                rk = ph.add_run(f"{k}\n")
                                rk.bold = True
                                rk.font.size = Pt(8)
                                rv = ph.add_run(str(v))
                                rv.font.size = Pt(9)
                            # Deuxième ligne si > 4 indicateurs
                            if len(ind_items) > 4:
                                row2 = ind_tbl.add_row().cells
                                for ci, (k, v) in enumerate(ind_items[4:8]):
                                    row2[ci].text = ''
                                    ph2 = row2[ci].paragraphs[0]
                                    rk2 = ph2.add_run(f"{k}\n")
                                    rk2.bold = True
                                    rk2.font.size = Pt(8)
                                    rv2 = ph2.add_run(str(v))
                                    rv2.font.size = Pt(9)
                            doc.add_paragraph()

                    # ── Points clés ──────────────────────────────────────────────
                    if r['points']:
                        rpk_hdr = doc.add_paragraph()
                        rpk_hdr.paragraph_format.left_indent = Pt(20)
                        rpk_hdr.paragraph_format.space_before = Pt(3)
                        rpk_hdr.add_run("Points clés :").bold = True
                        rpk_hdr.runs[-1].font.size = Pt(9)
                        for pt in r['points']:
                            rpk_p = doc.add_paragraph(style='List Bullet')
                            rpk_p.paragraph_format.left_indent = Pt(35)
                            rpk_p.paragraph_format.space_before = Pt(1)
                            rpk_p.paragraph_format.space_after = Pt(1)
                            rpk_p.add_run(str(pt)[:200]).font.size = Pt(9)

                    # ── Risques ──────────────────────────────────────────────────
                    if r['risques']:
                        rrisk = doc.add_paragraph()
                        rrisk.paragraph_format.left_indent = Pt(20)
                        rrisk.paragraph_format.space_before = Pt(2)
                        rrisk.add_run("⚠️ Risques : ").bold = True
                        rrisk.runs[-1].font.size = Pt(9)
                        rrisk.runs[-1].font.color.rgb = RGBColor(192, 80, 0)
                        rrisk.add_run(r['risques'][:300]).font.size = Pt(9)

                    # ── Perspectives ─────────────────────────────────────────────
                    if r['perspectives']:
                        rpersp = doc.add_paragraph()
                        rpersp.paragraph_format.left_indent = Pt(20)
                        rpersp.paragraph_format.space_before = Pt(2)
                        rpersp.add_run("🔮 Perspectives : ").bold = True
                        rpersp.runs[-1].font.size = Pt(9)
                        rpersp.runs[-1].font.color.rgb = RGBColor(0, 100, 160)
                        rpersp.add_run(r['perspectives'][:300]).font.size = Pt(9)

                    # ── Séparateur ───────────────────────────────────────────────
                    if ri < nb_rap_display:
                        sep_rp = doc.add_paragraph()
                        sep_rp.paragraph_format.left_indent = Pt(20)
                        sep_rp.add_run("· " * 30).font.size = Pt(7)
                        sep_rp.runs[-1].font.color.rgb = RGBColor(180, 180, 180)

            # ═══════════════════════════════════════════════════════════════════════
            # ═══════════════════════════════════════════════════════════════════════
            doc.add_paragraph()
            doc.add_paragraph("═" * 80)

            if idx < len(all_analyses):
                doc.add_page_break()
        
        # ========== PIED DE PAGE ==========
        doc.add_page_break()
        footer = doc.add_heading('NOTES IMPORTANTES', level=1)
        footer_text = doc.add_paragraph(
            "1. Les analyses techniques sont basées sur les 5 indicateurs classiques (MM, Bollinger, MACD, RSI, Stochastique).\n"
            "2. Les analyses fondamentales proviennent des rapports financiers officiels de la BRVM.\n"
            "3. Les recommandations sont générées par intelligence artificielle Multi-AI (DeepSeek, Gemini, Mistral) avec rotation automatique.\n"
            "4. Tous les cours et valeurs sont exprimés en FCFA (Francs CFA).\n"
            "5. Les prédictions sont des estimations statistiques et ne garantissent pas les performances futures.\n"
            "6. L'analyse de liquidité est basée sur les 30 derniers jours de transactions.\n"
            "7. Les matrices de convergence et de divergence sont des outils d'aide à la décision complémentaires.\n"
            "8. Ce document est strictement confidentiel et destiné à l'usage professionnel uniquement.\n"
            "9. Consultez toujours un conseiller financier agréé avant toute décision d'investissement.\n"
            f"10. Rapport généré automatiquement - Version 30.2 Ultimate - {len(all_analyses)} sociétés analysées."
        )
        
        # Signature IA
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.add_run("Analyse Multi-AI - ").italic = True
        sig.add_run(f"DeepSeek: {self.request_count['deepseek']} req | ").italic = True
        sig.add_run(f"Gemini: {self.request_count['gemini']} req | ").italic = True
        sig.add_run(f"Mistral: {self.request_count['mistral']} req").italic = True
        sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig.runs[0].font.size = Pt(8)
        sig.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Sauvegarde
        filename = f"Rapport_Ultimate_BRVM_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(filename)
        
        logging.info(f"   ✅ Document créé: {filename}")
        
        # Préparer la synthèse pour la DB
        if market_indicators and market_indicators.get('composite'):
            cap_val = market_indicators.get('capitalisation', 0)
            cap_mds = cap_val / 1e9
            cap_mds_entier = int(cap_mds)
            cap_mds_dec = round((cap_mds - cap_mds_entier) * 1000)
            cap_mds_display = f"{cap_mds_entier:,}".replace(",", " ") + f",{cap_mds_dec:03d}"
            synthesis_text = (
                f"Analyse ULTIMATE de {len(all_analyses)} sociétés. "
                f"Indices: BRVM Composite {market_indicators['composite']:.2f} pts. "
                f"Capitalisation: {cap_mds_display} Mds FCFA. "
                f"Multi-AI: DeepSeek ({self.request_count['deepseek']}), "
                f"Gemini ({self.request_count['gemini']}), Mistral ({self.request_count['mistral']})."
            )
        else:
            synthesis_text = f"Analyse ULTIMATE de {len(all_analyses)} sociétés de la BRVM."
        
        # market_events : utiliser les alertes Google déjà chargées ou chaîne vide
        market_events = ""
        try:
            _me_df = self._get_google_alerts_events()
            if not _me_df.empty:
                market_events = " | ".join(
                    str(r.get('titre') or r.get('resume',''))[:80]
                    for _, r in _me_df.head(5).iterrows()
                )
        except Exception:
            pass

        self._save_to_database(
            datetime.now().date(),
            synthesis_text,
            top_10_list,
            flop_10_list,
            market_events,
            all_company_data,
            filename
        )
        
        return filename

    def generate_all_reports(self, new_fundamental_analyses):
        """Génération du rapport ULTIMATE complet avec toutes les analyses"""
        logging.info("="*80)
        logging.info("📝 ÉTAPE 5: GÉNÉRATION RAPPORTS (V30.3 — FIX GRAPHS 2025)")
        logging.info("🤖 Multi-AI: DeepSeek → Gemini → Mistral")
        logging.info("📊 Analyses: Sectorielles + Convergence + Liquidité + Divergences + Risque/Horizon")
        logging.info("="*80)
        
        # Vérifier qu'au moins une clé API est disponible
        if not any([DEEPSEEK_API_KEY, GEMINI_API_KEY, MISTRAL_API_KEY, ANTHROPIC_API_KEY]):
            logging.error("❌ Aucune clé API configurée!")
            return
        
        available_apis = []
        missing_apis = []
        if DEEPSEEK_API_KEY:
            available_apis.append("DeepSeek")
        else:
            missing_apis.append("DeepSeek")
        if GEMINI_API_KEY:
            available_apis.append("Gemini")
        else:
            missing_apis.append("Gemini")
        if MISTRAL_API_KEY:
            available_apis.append("Mistral")
        else:
            missing_apis.append("Mistral")
        
        logging.info(f"✅ API disponibles: {', '.join(available_apis)}")
        if missing_apis:
            logging.warning(f"⚠️  API non configurées (ajouter dans GitHub Secrets): {', '.join(missing_apis)}")
        
        df = self._get_all_data_from_db()
        
        if df.empty:
            logging.error("❌ Aucune donnée disponible")
            return
        
        predictions_df = self._get_predictions_from_db()
        brvm_docs_by_symbol     = self._get_brvm_documents()
        brvm_rapports_by_symbol = self._get_brvm_rapports_societes()
        
        logging.info(f"🤖 Génération de {len(df)} analyse(s) avec rotation Multi-AI...")

        # ── Fonction locale : confiance dynamique ─────────────────────────────
        def _dynamic_confidence(tech_dec, fund_dec, rec_score):
            """
            Calcule un niveau de confiance selon la convergence tech / fondamental.
            Élevée  : tech et fond convergent ET score ≥ 4
            Faible  : forte divergence (ACHAT vs VENTE)
            Moyen   : tous les autres cas
            """
            td = str(tech_dec).upper()
            fd = str(fund_dec).upper()
            # Convergence forte
            if td == fd and td != 'NEUTRE' and rec_score >= 4:
                return 'Élevée'
            # Divergence directe ACHAT ↔ VENTE
            if (('ACHAT' in td and 'VENTE' in fd) or
                ('VENTE' in td and 'ACHAT' in fd)):
                return 'Faible'
            # Score extrême même sans convergence parfaite
            if rec_score == 5:
                return 'Élevée'
            if rec_score == 1:
                return 'Faible'
            return 'Moyen'

        all_analyses = {}
        all_company_data = {}
        
        for idx, row in df.iterrows():
            symbol = row['symbol']
            company_id = row['company_id']
            company_name = row.get('company_name', 'N/A')
            
            hist_df = self._get_historical_data_100days(company_id)
            
            historical_summary = "Données historiques non disponibles."
            price_evolution_100d = None
            highest_price = None
            lowest_price = None
            
            if not hist_df.empty and len(hist_df) > 1:
                import numpy as np
                from scipy import stats as scipy_stats

                prices_s = hist_df['price'].astype(float)
                prix_debut = float(hist_df.iloc[0]['price'])
                prix_fin   = float(hist_df.iloc[-1]['price'])
                prix_max   = float(prices_s.max())
                prix_min   = float(prices_s.min())
                evolution_pct = ((prix_fin - prix_debut) / prix_debut * 100) if prix_debut > 0 else 0

                price_evolution_100d = evolution_pct
                highest_price        = prix_max
                lowest_price         = prix_min

                # ── Variation J-1 (cours actuel vs avant-dernier cours) ──────────
                var_j1 = None
                var_j1_txt = "N/A"
                if len(hist_df) >= 2:
                    p_curr = float(hist_df.iloc[-1]['price'])
                    p_prev = float(hist_df.iloc[-2]['price'])
                    if p_prev > 0:
                        var_j1 = ((p_curr - p_prev) / p_prev) * 100
                        sign_j1 = "+" if var_j1 >= 0 else ""
                        var_j1_txt = f"{sign_j1}{var_j1:.2f}%"

                # ── Capitalisation boursière (dernière valeur disponible) ────────
                capit = None
                capit_txt = "N/D"
                if 'company_capitalization' in hist_df.columns:
                    cap_series = hist_df['company_capitalization'].dropna()
                    if not cap_series.empty:
                        capit = float(cap_series.iloc[-1])
                        if capit >= 1e9:
                            capit_txt = f"{capit/1e9:.2f} Mds FCFA"
                        elif capit >= 1e6:
                            capit_txt = f"{capit/1e6:.0f} M FCFA"
                        else:
                            capit_txt = f"{capit:,.0f} FCFA"

                # ── Statistiques descriptives du cours sur 100j ──────────────────
                p_mean   = float(prices_s.mean())
                p_median = float(prices_s.median())
                p_std    = float(prices_s.std())
                p_cv     = (p_std / p_mean * 100) if p_mean > 0 else 0
                try:
                    p_mode = float(prices_s.mode().iloc[0])
                except Exception:
                    p_mode = p_mean
                try:
                    p_kurt = float(scipy_stats.kurtosis(prices_s, fisher=True))
                except Exception:
                    p_kurt = 0.0
                try:
                    p_skew = float(scipy_stats.skew(prices_s))
                except Exception:
                    p_skew = 0.0

                # Interprétation kurtosis et skewness
                if p_kurt > 1:
                    kurt_interp = "leptokurtique (queues épaisses, risque de pics)"
                elif p_kurt < -1:
                    kurt_interp = "platykurtique (distribution aplatie, faible concentration)"
                else:
                    kurt_interp = "mésokurtique (proche d'une distribution normale)"

                if p_skew > 0.5:
                    skew_interp = "asymétrie positive (plus de journées sous la moyenne)"
                elif p_skew < -0.5:
                    skew_interp = "asymétrie négative (plus de journées au-dessus de la moyenne)"
                else:
                    skew_interp = "distribution symétrique"

                # ── Ratios de valorisation (si fin_data disponible plus tard) ────
                # Calculés ici, injectés dans data_dict pour usage dans le prompt
                returns_s  = prices_s.pct_change().dropna()
                vol_annualisee = float(returns_s.std() * np.sqrt(252) * 100) if len(returns_s) > 1 else 0

                historical_summary = (
                    f"╔══ ANALYSE DU COURS — 100 DERNIERS JOURS ══╗\n"
                    f"Variation totale        : {evolution_pct:+.2f}%  |  "
                    f"Variation J-1           : {var_j1_txt}\n"
                    f"Cours actuel            : {prix_fin:,.0f} FCFA  |  "
                    f"Capitalisation boursière: {capit_txt}\n"
                    f"Plus haut               : {prix_max:,.0f} FCFA  |  "
                    f"Plus bas                : {prix_min:,.0f} FCFA\n"
                    f"Volume moyen / jour     : {hist_df['volume'].mean():,.0f} titres\n"
                    f"\n── Statistiques descriptives du cours (100j) ──\n"
                    f"Moyenne                 : {p_mean:,.0f} FCFA\n"
                    f"Médiane                 : {p_median:,.0f} FCFA\n"
                    f"Mode                    : {p_mode:,.0f} FCFA\n"
                    f"Écart-type              : {p_std:,.0f} FCFA  (CV={p_cv:.1f}%)\n"
                    f"Kurtosis (excès)        : {p_kurt:.3f} → {kurt_interp}\n"
                    f"Skewness                : {p_skew:.3f} → {skew_interp}\n"
                    f"Volatilité annualisée   : {vol_annualisee:.1f}%\n"
                    f"╚══════════════════════════════════════════╝"
                )

                # Sauvegarder les valeurs clés pour ratios de valorisation
                _hist_stats = {
                    'prix_actuel':        prix_fin,
                    'capitalisation':     capit,
                    'var_j1':             var_j1,
                    'var_j1_txt':         var_j1_txt,
                    'p_mean':             p_mean,
                    'p_std':              p_std,
                    'p_kurt':             p_kurt,
                    'p_skew':             p_skew,
                    'vol_annualisee':     vol_annualisee,
                    'kurt_interp':        kurt_interp,
                    'skew_interp':        skew_interp,
                }
            
            fundamental_text = ""
            raw_summaries = row.get('fundamental_summaries')
            nb_rapports_db = int(row.get('nb_rapports_fondamentaux', 0)) if pd.notna(row.get('nb_rapports_fondamentaux', None)) else 0
            
            if raw_summaries and pd.notna(raw_summaries) and str(raw_summaries).strip():
                reports = str(raw_summaries).split('###SEP_REPORT###')
                fundamental_parts = []
                for report in reports:
                    report = report.strip()
                    if not report:
                        continue
                    parts = report.split('###SEP_FIELD###')
                    if len(parts) >= 3:
                        title           = parts[0].strip()
                        report_date_str = parts[1].strip()
                        summary         = '###SEP_FIELD###'.join(parts[2:]).strip()
                        if not summary:
                            continue
                        try:
                            from datetime import date as _date
                            report_year = int(report_date_str[:4])
                            report_month = int(report_date_str[5:7]) if len(report_date_str) >= 7 else 1
                            today = _date.today()
                            months_ago = (today.year - report_year) * 12 + (today.month - report_month)
                            if months_ago <= 3:
                                freshness = "⭐ RÉCENT"
                            elif months_ago <= 12:
                                freshness = "📅 À JOUR"
                            else:
                                freshness = "⚠️ ANCIEN"
                            date_label = f"{report_date_str} {freshness}"
                        except Exception:
                            date_label = report_date_str
                        fundamental_parts.append(
                            f"--- RAPPORT: {title} | Date: {date_label} ---\n{summary}"
                        )
                    elif len(parts) == 1 and parts[0]:
                        fundamental_parts.append(parts[0])
                
                if fundamental_parts:
                    fundamental_text = "\n\n".join(fundamental_parts)
                    logging.info(f"   📄 {symbol}: {len(fundamental_parts)}/{nb_rapports_db} rapport(s) parsé(s) | {len(fundamental_text)} caractères")
                else:
                    logging.warning(f"   ⚠️ {symbol}: fundamental_summaries présent ({nb_rapports_db} en DB) mais parsing échoué")
                    logging.warning(f"   🔍 {symbol}: début={str(raw_summaries)[:200]}")
            else:
                if nb_rapports_db > 0:
                    logging.warning(f"   ⚠️ {symbol}: {nb_rapports_db} rapport(s) en DB mais fundamental_summaries vide — vérifier la requête SQL")
                else:
                    logging.info(f"   ℹ️ {symbol}: Aucune analyse fondamentale en base")

            # ── Ratios de valorisation (PER, P/B, EV/EBITDA, BPA) ─────────────────
            val_ratios = {}
            try:
                fin_d = self._get_donnees_financieres(symbol)
                hs    = _hist_stats if not hist_df.empty and len(hist_df) > 1 else {}
                if fin_d and hs:
                    px       = hs.get('prix_actuel')
                    mkt_cap  = hs.get('capitalisation')
                    rn       = float(fin_d.get('resultat_net') or 0)
                    cp       = float(fin_d.get('capitaux_propres') or 0)
                    cs       = float(fin_d.get('capital_souscrit') or 0)
                    ebitda   = float(fin_d.get('ebe') or fin_d.get('rbe') or 0)
                    dettes   = float(fin_d.get('dettes_financieres_totales') or fin_d.get('dettes_totales') or 0)
                    tresor   = float(fin_d.get('tresorerie_nette') or fin_d.get('tresorerie_actif') or 0)
                    annee_v  = fin_d.get('annee', 'N/A')

                    # Nombre d'actions estimé (capital souscrit / valeur nominale estimée)
                    # Si capital_souscrit et cours connus, nb_actions ~ cap_bours / cours
                    nb_actions = (mkt_cap / px) if (mkt_cap and px and px > 0) else None

                    # BPA / EPS
                    if rn and nb_actions and nb_actions > 0:
                        bpa = rn / nb_actions
                        val_ratios['bpa'] = bpa
                        val_ratios['bpa_txt'] = f"{bpa:,.0f} FCFA/action"

                    # PER
                    if rn and nb_actions and nb_actions > 0 and px:
                        per = px / (rn / nb_actions)
                        val_ratios['per'] = per
                        if per < 0:
                            per_interp = "négatif (société déficitaire)"
                        elif per < 10:
                            per_interp = "faible (potentiellement sous-évalué)"
                        elif per < 20:
                            per_interp = "modéré (valorisation raisonnable)"
                        elif per < 35:
                            per_interp = "élevé (croissance attendue ou surévaluation)"
                        else:
                            per_interp = "très élevé (prudence requise)"
                        val_ratios['per_txt'] = f"{per:.1f}x → {per_interp}"

                    # P/B (Price-to-Book)
                    if cp and cp > 0 and mkt_cap:
                        pb = mkt_cap / cp
                        val_ratios['pb'] = pb
                        if pb < 1:
                            pb_interp = "< 1 (marché valorise sous la valeur comptable — opportunité potentielle)"
                        elif pb < 2:
                            pb_interp = "1-2x (valorisation raisonnable)"
                        elif pb < 4:
                            pb_interp = "2-4x (prime de croissance accordée)"
                        else:
                            pb_interp = f"> 4x (forte prime — valorisation exigeante)"
                        val_ratios['pb_txt'] = f"{pb:.2f}x → {pb_interp}"

                    # EV/EBITDA
                    if ebitda and ebitda > 0 and mkt_cap:
                        ev = mkt_cap + dettes - tresor
                        ev_ebitda = ev / ebitda
                        val_ratios['ev_ebitda'] = ev_ebitda
                        if ev_ebitda < 5:
                            ev_interp = "< 5x (très bon marché)"
                        elif ev_ebitda < 10:
                            ev_interp = "5-10x (raisonnable)"
                        elif ev_ebitda < 15:
                            ev_interp = "10-15x (valorisation élevée)"
                        else:
                            ev_interp = f"> 15x (très exigeant)"
                        val_ratios['ev_ebitda_txt'] = f"{ev_ebitda:.1f}x → {ev_interp}"
                        val_ratios['ev'] = ev

                    val_ratios['annee_fin'] = annee_v
                    val_ratios['mkt_cap'] = mkt_cap
                    val_ratios['mkt_cap_txt'] = hs.get('capitalisation') and (
                        f"{mkt_cap/1e9:.2f} Mds FCFA" if mkt_cap >= 1e9 else f"{mkt_cap/1e6:.0f} M FCFA"
                    ) or "N/D"

            except Exception as e_val:
                logging.warning(f"   ⚠️ {symbol}: Calcul ratios valorisation échoué: {e_val}")

            data_dict = {
                'price': row.get('price'),
                'volume': row.get('volume'),
                'historical_summary': historical_summary,
                'hist_stats': _hist_stats if not hist_df.empty and len(hist_df) > 1 else {},
                'val_ratios': val_ratios,
                'mm_20': row.get('mm20'),
                'mm_50': row.get('mm50'),
                'mm_decision': row.get('mm_decision'),
                'bollinger_upper': row.get('bollinger_superior'),
                'bollinger_lower': row.get('bollinger_inferior'),
                'bollinger_decision': row.get('bollinger_decision'),
                'macd_value': row.get('macd_line'),
                'macd_signal': row.get('signal_line'),
                'macd_decision': row.get('macd_decision'),
                'rsi_value': row.get('rsi'),
                'rsi_decision': row.get('rsi_decision'),
                'stochastic_k': row.get('stochastic_k'),
                'stochastic_d': row.get('stochastic_d'),
                'stochastic_decision': row.get('stochastic_decision'),
                'fundamental_analyses': fundamental_text if fundamental_text else "Aucun rapport financier enregistré en base pour cette société.",
                'predictions': [],
                'brvm_docs_raw': []
            }

            # ── Enrichissement avec brvm_documents ─────────────────────────────
            symbol_brvm_docs = brvm_docs_by_symbol.get(symbol, [])
            brvm_docs_text = self._format_brvm_documents_for_ai(symbol_brvm_docs)
            if brvm_docs_text:
                nb_docs = len(symbol_brvm_docs)
                logging.info(f'   📎 {symbol}: {nb_docs} document(s) brvm_documents disponibles')
                separator = '\n\n' + '─'*60 + '\n📎 DOCUMENTS OFFICIELS BRVM (brvm_documents):\n' + '─'*60 + '\n'
                existing = data_dict['fundamental_analyses']
                if existing and 'Aucun rapport' not in existing:
                    data_dict['fundamental_analyses'] = existing + separator + brvm_docs_text
                else:
                    data_dict['fundamental_analyses'] = separator.lstrip() + brvm_docs_text
            data_dict['brvm_docs_raw'] = symbol_brvm_docs

            # ── Enrichissement avec brvm_rapports_societes ──────────────────────
            symbol_rapports = brvm_rapports_by_symbol.get(symbol, [])
            rapports_text   = self._format_rapports_societes_for_ai(symbol_rapports)
            if rapports_text:
                nb_rap = len(symbol_rapports)
                logging.info(f'   📋 {symbol}: {nb_rap} rapport(s) brvm_rapports_societes disponibles')
                sep2 = ('\n\n' + '─'*60 +
                        '\n📋 RAPPORTS & COMMUNIQUÉS (brvm_rapports_societes):\n' +
                        '─'*60 + '\n')
                existing2 = data_dict['fundamental_analyses']
                if existing2 and 'Aucun rapport' not in existing2:
                    data_dict['fundamental_analyses'] = existing2 + sep2 + rapports_text
                else:
                    data_dict['fundamental_analyses'] = sep2.lstrip() + rapports_text
            data_dict['brvm_rapports_raw'] = symbol_rapports

            symbol_predictions = predictions_df[predictions_df['symbol'] == symbol]
            if not symbol_predictions.empty:
                # Récupérer prédictions complètes avec IC
                data_dict['predictions'] = [
                    {
                        'date':       str(r['prediction_date']),
                        'price':      float(r['predicted_price'])      if pd.notna(r['predicted_price'])  else None,
                        'lower':      float(r['lower_bound'])          if pd.notna(r.get('lower_bound'))  else None,
                        'upper':      float(r['upper_bound'])          if pd.notna(r.get('upper_bound'))  else None,
                        'confidence': str(r.get('confidence_level','')) if pd.notna(r.get('confidence_level')) else '',
                    }
                    for _, r in symbol_predictions.head(10).iterrows()
                    if pd.notna(r['predicted_price'])
                ]
                if data_dict['predictions']:
                    pred_list = [
                        f"{p['date']}: {p['price']:.0f} FCFA"
                        for p in data_dict['predictions'][:5]
                        if p['price'] is not None
                    ]
                    data_dict['predictions_text'] = ", ".join(pred_list) if pred_list else "Aucune prédiction disponible"
                else:
                    data_dict['predictions_text'] = "Aucune prédiction disponible"
            else:
                data_dict['predictions_text'] = "Aucune prédiction disponible"
            
            analysis = self._generate_professional_analysis(symbol, data_dict)
            all_analyses[symbol] = analysis
            
            # ── Décision technique préliminaire (avant appel recommendation) ──
            _tech_sigs_pre = []
            for _k in ['mm_decision','bollinger_decision','macd_decision',
                       'rsi_decision','stochastic_decision']:
                _v = row.get(_k)
                if _v and pd.notna(_v):
                    _tech_sigs_pre.append(str(_v))
            _buy_pre  = sum(1 for s in _tech_sigs_pre if 'Achat' in s)
            _sell_pre = sum(1 for s in _tech_sigs_pre if 'Vente' in s)
            _td_pre   = 'ACHAT' if _buy_pre > _sell_pre else ('VENTE' if _sell_pre > _buy_pre else 'NEUTRE')

            # ── Décision fondamentale préliminaire (depuis texte IA brut) ──────
            _al_pre = analysis.lower()
            if 'achat' in _al_pre:   _fd_pre = 'ACHAT'
            elif 'vente' in _al_pre: _fd_pre = 'VENTE'
            else:                    _fd_pre = 'NEUTRE'

            # Délai entre sociétés pour éviter le rate limit Mistral free tier
            time.sleep(1)

            recommendation, rec_score = self._extract_recommendation_from_analysis(
                analysis, tech_decision=_td_pre, fund_decision=_fd_pre
            )
            
            # ✅ V30: Calcul de la décision technique (majorité des indicateurs)
            tech_signals = []
            for key in ['mm_decision', 'bollinger_decision', 'macd_decision', 'rsi_decision', 'stochastic_decision']:
                val = row.get(key)
                if val and pd.notna(val):
                    tech_signals.append(str(val))
            
            buy_count = sum(1 for s in tech_signals if 'Achat' in s)
            sell_count = sum(1 for s in tech_signals if 'Vente' in s)
            
            if buy_count > sell_count:
                technical_decision = "ACHAT"
            elif sell_count > buy_count:
                technical_decision = "VENTE"
            else:
                technical_decision = "NEUTRE"
            
            # ✅ V30: Décision fondamentale (basée sur recommandation finale)
            if 'ACHAT' in recommendation.upper():
                fundamental_decision = "ACHAT"
            elif 'VENTE' in recommendation.upper():
                fundamental_decision = "VENTE"
            else:
                fundamental_decision = "NEUTRE"
            
            # ✅ V30: Calcul du risque chiffré
            # Créer un dict temporaire pour _calculate_risk_score
            temp_data = {
                'company_id': company_id,
                'mm_decision': row.get('mm_decision'),
                'bollinger_decision': row.get('bollinger_decision'),
                'macd_decision': row.get('macd_decision'),
                'rsi_decision': row.get('rsi_decision'),
                'stochastic_decision': row.get('stochastic_decision')
            }
            
            risk_data = self._calculate_risk_score(temp_data)
            
            all_company_data[symbol] = {
                'company_id': company_id,
                'company_name': company_name,
                'sector': row.get('sector'),
                'current_price': float(row.get('price', 0)) if pd.notna(row.get('price')) else None,
                'capitalisation':     capit      if capit is not None else None,
                'capitalisation_txt': capit_txt,
                'volume_moyen_jour':  float(hist_df['volume'].mean()) if not hist_df.empty and 'volume' in hist_df.columns else None,
                'vol_annualisee':     vol_annualisee if 'vol_annualisee' in dir() else None,
                'price_evolution_100d': price_evolution_100d,
                'highest_price_100d': highest_price,
                'lowest_price_100d': lowest_price,
                'mm20': float(row.get('mm20', 0)) if pd.notna(row.get('mm20')) else None,
                'mm50': float(row.get('mm50', 0)) if pd.notna(row.get('mm50')) else None,
                'mm_decision': row.get('mm_decision'),
                'bollinger_upper': float(row.get('bollinger_superior', 0)) if pd.notna(row.get('bollinger_superior')) else None,
                'bollinger_lower': float(row.get('bollinger_inferior', 0)) if pd.notna(row.get('bollinger_inferior')) else None,
                'bollinger_decision': row.get('bollinger_decision'),
                'macd_value': float(row.get('macd_line', 0)) if pd.notna(row.get('macd_line')) else None,
                'macd_signal': float(row.get('signal_line', 0)) if pd.notna(row.get('signal_line')) else None,
                'macd_decision': row.get('macd_decision'),
                'rsi_value': float(row.get('rsi', 0)) if pd.notna(row.get('rsi')) else None,
                'rsi_decision': row.get('rsi_decision'),
                'stochastic_k': float(row.get('stochastic_k', 0)) if pd.notna(row.get('stochastic_k')) else None,
                'stochastic_d': float(row.get('stochastic_d', 0)) if pd.notna(row.get('stochastic_d')) else None,
                'stochastic_decision': row.get('stochastic_decision'),
                'full_analysis': analysis,
                'technical_conclusion': '',
                'fundamental_analysis': fundamental_text,
                'investment_conclusion': '',
                'recommendation': recommendation,
                'recommendation_score': rec_score,
                'confidence_level': _dynamic_confidence(technical_decision, fundamental_decision, rec_score),
                'risk_level': risk_data['level'],  # ✅ V30: Niveau de risque calculé
                'investment_horizon': 'Moyen terme',
                # ✅ V30: Nouvelles colonnes
                'technical_decision': technical_decision,
                'fundamental_decision': fundamental_decision,
                'risk_score': risk_data['score'],
                'risk_details': json.dumps(risk_data['details'], ensure_ascii=False),
                # ✅ Sources documentaires enrichies
                'brvm_docs_raw':     data_dict.get('brvm_docs_raw', []),
                'brvm_rapports_raw': data_dict.get('brvm_rapports_raw', []),
                'predictions_full':  data_dict.get('predictions', []),
            }
            # ── Score composite d'investissement ────────────────────────────────
            inv_score, inv_label = self._compute_investment_score(all_company_data[symbol])
            all_company_data[symbol]['investment_score'] = inv_score
            all_company_data[symbol]['investment_label'] = inv_label
        
        filename = self._create_word_document(all_analyses, all_company_data)
        
        logging.info(f"\n✅ Rapport ULTIMATE généré: {filename}")
        logging.info(f"📊 Statistiques requêtes Multi-AI:")
        logging.info(f"   - DeepSeek: {self.request_count['deepseek']}")
        logging.info(f"   - Gemini: {self.request_count['gemini']}")
        logging.info(f"   - Mistral: {self.request_count['mistral']}")
        logging.info(f"   - Claude:  {self.request_count['claude']}")
        logging.info(f"   - TOTAL: {self.request_count['total']}")
        logging.info(f"\n📋 Analyses incluses:")
        logging.info(f"   ✅ Analyse par secteur")
        logging.info(f"   ✅ Matrice de convergence des signaux")
        logging.info(f"   ✅ Analyse de liquidité")
        logging.info(f"   ✅ Top 10 divergences majeures")
        logging.info(f"   ✅ Matrice Risque vs Horizon")

    def __del__(self):
        if self.db_conn and not self.db_conn.closed:
            self.db_conn.close()


if __name__ == "__main__":
    try:
        report_generator = BRVMReportGenerator()
        report_generator.generate_all_reports([])
    except Exception as e:
        logging.critical(f"❌ Erreur: {e}", exc_info=True)
